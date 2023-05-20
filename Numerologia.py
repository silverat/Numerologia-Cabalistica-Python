# Neste projeto trabalhamos com tratamento de data, pois precisamos dividir o arquivo 
# Após as funções que calculam as datas, utilizamos as bibliotecas para gerar um arquivo em doc.
import re
import datetime
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, Inches
import sys
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


anoatual = date.today().year
mesatual = date.today().month
diaatual = date.today().day

Descmes = {
    1: 'Janeiro',
    2: 'Fevereiro',
    3: 'Março',
    4: 'Abril',
    5: 'Maio',
    6: 'Junho',
    7: 'Julho',
    8: 'Agosto',
    9: 'Setembro',
    10:'Outubro',
    11:'Novembro',
    12:'Dezembro'
}


cabalacod = {'A': 1, 'I': 1, 'Q': 1, 'J': 1, 'Y': 1, 'Õ': 1, 'B': 2, 'K': 2, 'R': 2, 'C': 3,
             'G': 3, 'L': 3, 'S': 3, 'Ê': 3, 'Á': 3, 'Í': 3, 'D': 4, 'M': 4, 'T': 4, 'Û': 4, 'Ã': 4,
             'E': 5, 'H': 5, 'N': 5, 'Ô': 5, 'U': 6, 'V': 6, 'W': 6, 'X': 6, 'Ç': 6, 'O': 7,
             'Z': 7, 'É': 7, 'F': 8, 'P': 8, 'Â': 8, 'Î': 8, 'Ú': 8, 'Ó': 9, '’': 2, 'À': 2,
             'È': 1, 'Ì': 2, 'Ò': 5, 'Ù': 3, 'Ä': 2, 'Ë': 1, 'Ï': 2, 'Ö': 5, 'Ü': 3}
harmoniaconjugal = [[1, 9, [4, 8], [6, 7], [2, 3, 5]], [2, 8, [7, 9], 5, [1, 3, 4, 6]],
                    [3, 7, [5, 6, 9], [4, 8], [1, 2]], [4, 6, [1, 8], [3, 5], [2, 7, 9]],
                    [5, 5, [3, 9], [2, 4, 66], [1, 7, 8]], [6, 4, [3, 7, 9], [1, 8, 55], 2],
                    [7, 3, [2, 6], [1, 9], [4, 5, 8]], [8, 2, [1, 4], [3, 6], [5, 7, 9]],
                    [9, 1, [2, 3, 5, 6], 7, [4, 8]]]
numerosfavoraveis = [
    ['JANEIRO', [1, [1, 5]], [2, [1, 6]], [3, [3, 6]], [4, [1, 5]], [5, [5, 6]], [6, [5, 6]], [7, [1, 7]], [8, [1, 3]],
     [9, [6, 9]], [10, [1, 5]], [11, [1, 6]], [12, [6, 9]], [13, [1, 5]], [14, [5, 6]], [15, [5, 6]],
     [16, [1, 5]], [17, [1, 3]], [18, [5, 6]], [19, [1, 5]], [20, [1, 6]], [21, [3, 6]], [22, [1, 5]], [23, [5, 6]],
     [24, [5, 6]], [25, [1, 5]], [26, [2, 3]], [27, [6, 9]], [28, [2, 7]], [29, [5, 7]], [30, [2, 3]], [31, [2, 7]], ],
    ['FEVEREIRO', [1, [2, 7]], [2, [2, 7]], [3, [3, 6]], [4, [2, 7]], [5, [5, 6]], [6, [3, 6]], [7, [2, 7]],
     [8, [2, 3]], [9, [3, 6]], [10, [2, 7]], [11, [5, 7]], [12, [5, 6]], [13, [2, 7]], [14, [5, 6]], [15, [3, 6]],
     [16, [2, 5]], [17, [2, 3]], [18, [3, 6]], [19, [2, 7]], [20, [2, 7]], [21, [3, 6]], [22, [2, 7]], [23, [5, 6]],
     [24, [5, 6]], [25, [2, 7]], [26, [2, 3]], [27, [6, 9]], [28, [2, 7]], [29, [6, 7]], [30, [3, 9]], [31, [1, 7]], ],
    ['MARÇO', [1, [1, 7]], [2, [2, 7]], [3, [3, 6]], [4, [1, 7]], [5, [5, 7]], [6, [3, 6]], [7, [2, 7]], [8, [3, 6]],
     [9, [6, 9]], [10, [1, 7]], [11, [1, 7]], [12, [6, 7]], [13, [1, 5]], [14, [5, 7]], [15, [3, 6]],
     [16, [1, 2]], [17, [3, 6]], [18, [3, 6]], [19, [1, 7]], [20, [2, 7]], [21, [3, 6]], [22, [1, 7]], [23, [6, 7]],
     [24, [3, 6]], [25, [2, 7]], [26, [1, 3]], [27, [1, 9]], [28, [5, 9]], [29, [1, 7]], [30, [3, 6]], [31, [1, 5]], ],
    ['ABRIL ', [1, [1, 7]], [2, [1, 7]], [3, [3, 9]], [4, [1, 7]], [5, [5, 7]], [6, [3, 6]], [7, [5, 7]], [8, [1, 3]],
     [9, [3, 9]], [10, [1, 7]], [11, [1, 7]], [12, [1, 9]], [13, [1, 7]], [14, [5, 7]], [15, [3, 6]],
     [16, [1, 2]], [17, [1, 3]], [18, [1, 3]], [19, [1, 7]], [20, [2, 7]], [21, [1, 3]], [22, [1, 7]], [23, [5, 7]],
     [24, [3, 5]], [25, [5, 7]], [26, [2, 3]], [27, [3, 6]], [28, [2, 7]], [29, [1, 7]], [30, [5, 6]], ],
    ['MAIO ', [1, [1, 2]], [2, [2, 7]], [3, [3, 6]], [4, [1, 7]], [5, [5, 6]], [6, [5, 6]], [7, [2, 7]], [8, [2, 5]],
     [9, [5, 9]], [10, [1, 5]], [11, [1, 7]], [12, [2, 6]], [13, [1, 7]], [14, [5, 6]], [15, [5, 6]],
     [16, [2, 5]], [17, [2, 3]], [18, [5, 6]], [19, [1, 2]], [20, [2, 7]], [21, [3, 6]], [22, [1, 7]], [23, [5, 6]],
     [24, [5, 6]], [25, [2, 7]], [26, [2, 5]], [27, [5, 9]], [28, [2, 7]], [29, [5, 7]], [30, [2, 3]], ],
    ['JUNHO', [1, [1, 5]], [2, [2, 7]], [3, [5, 6]], [4, [1, 5]], [5, [5, 6]], [6, [5, 6]], [7, [2, 7]], [8, [3, 7]],
     [9, [5, 9]], [10, [1, 5]], [11, [5, 7]], [12, [5, 6]], [13, [1, 5]], [14, [5, 6]], [15, [5, 6]],
     [16, [2, 5]], [17, [2, 5]], [18, [5, 6]], [19, [1, 5]], [20, [2, 7]], [21, [5, 6]], [22, [1, 5]], [23, [5, 6]],
     [24, [5, 6]], [25, [2, 7]], [26, [2, 5]], [27, [5, 6]], [28, [2, 7]], [29, [1, 7]], ],
    ['JULHO ', [1, [1, 2]], [2, [2, 7]], [3, [2, 3]], [4, [1, 7]], [5, [5, 7]], [6, [2, 6]], [7, [2, 7]], [8, [2, 3]],
     [9, [2, 3]], [10, [1, 2]], [11, [1, 7]], [12, [2, 6]], [13, [1, 2]], [14, [5, 7]], [15, [6, 7]],
     [16, [1, 2]], [17, [2, 3]], [18, [2, 3]], [19, [1, 2]], [20, [2, 7]], [21, [3, 6]], [22, [1, 2]], [23, [5, 7]],
     [24, [6, 7]], [25, [2, 7]], [26, [2, 3]], [27, [1, 9]], [28, [2, 7]], [29, [1, 7]], [30, [3, 6]], [31, [1, 7]], ],
    ['AGOSTO', [1, [1, 2]], [2, [1, 5]], [3, [3, 6]], [4, [1, 2]], [5, [1, 5]], [6, [3, 6]], [7, [2, 7]], [8, [2, 3]],
     [9, [3, 6]], [10, [1, 2]], [11, [1, 7]], [12, [1, 6]], [13, [1, 2]], [14, [1, 5]], [15, [1, 6]],
     [16, [1, 2]], [17, [1, 3]], [18, [1, 3]], [19, [1, 2]], [20, [2, 7]], [21, [3, 6]], [22, [1, 2]], [23, [1, 5]],
     [24, [3, 6]], [25, [2, 7]], [26, [2, 3]], [27, [3, 6]], [28, [2, 5]], [29, [1, 5]], [30, [3, 6]], [31, [1, 5]], ],
    ['SETEMBRO', [1, [1, 5]], [2, [2, 5]], [3, [3, 6]], [4, [1, 5]], [5, [5, 6]], [6, [5, 6]], [7, [2, 5]], [8, [2, 3]],
     [9, [3, 6]], [10, [1, 2]], [11, [1, 5]], [12, [3, 6]], [13, [1, 5]], [14, [5, 6]], [15, [5, 6]],
     [16, [2, 5]], [17, [2, 3]], [18, [3, 6]], [19, [1, 5]], [20, [2, 7]], [21, [3, 6]], [22, [1, 7]], [23, [5, 6]],
     [24, [3, 6]], [25, [2, 7]], [26, [3, 6]], [27, [6, 9]], [28, [2, 7]], [29, [1, 7]], [30, [3, 6]], [31, [1, 7]], ],
    ['OUTUBRO', [1, [2, 7]], [2, [2, 7]], [3, [3, 6]], [4, [1, 7]], [5, [5, 6]], [6, [3, 6]], [7, [2, 7]], [8, [3, 6]],
     [9, [3, 6]], [10, [1, 5]], [11, [1, 6]], [12, [2, 6]], [13, [1, 7]], [14, [5, 6]], [15, [3, 6]],
     [16, [1, 2]], [17, [3, 6]], [18, [3, 6]], [19, [2, 7]], [20, [2, 7]], [21, [3, 6]], [22, [1, 7]], [23, [5, 6]],
     [24, [3, 6]], [25, [2, 7]], [26, [3, 6]], [27, [6, 9]], [28, [2, 7]], [29, [1, 7]], [30, [3, 6]], [31, [1, 3]], ],
    ['NOVEMBRO', [1, [1, 7]], [2, [1, 7]], [3, [3, 9]], [4, [1, 7]], [5, [5, 7]], [6, [3, 5]], [7, [1, 7]], [8, [3, 9]],
     [9, [3, 9]], [10, [2, 7]], [11, [1, 7]], [12, [1, 9]], [13, [1, 7]], [14, [5, 7]], [15, [3, 5]],
     [16, [1, 5]], [17, [3, 9]], [18, [3, 9]], [19, [1, 7]], [20, [2, 7]], [21, [3, 9]], [22, [1, 7]], [23, [5, 7]],
     [24, [3, 5]], [25, [1, 7]], [26, [3, 9]], [27, [3, 9]], [28, [2, 7]], [29, [1, 7]], [30, [3, 6]], ],
    ['DEZEMBRO', [1, [1, 7]], [2, [2, 7]], [3, [3, 6]], [4, [1, 7]], [5, [3, 6]], [6, [3, 6]], [7, [2, 7]], [8, [2, 3]],
     [9, [3, 9]], [10, [1, 7]], [11, [1, 7]], [12, [6, 9]], [13, [1, 3]], [14, [5, 6]], [15, [3, 6]],
     [16, [1, 2]], [17, [2, 3]], [18, [3, 6]], [19, [1, 7]], [20, [2, 7]], [21, [3, 6]], [22, [1, 7]], [23, [5, 6]],
     [24, [3, 6]], [25, [3, 7]], [26, [3, 6]], [27, [6, 9]], [28, [5, 6]], [29, [1, 6]], [30, [3, 6]], ]]
harmoniadatanasc = [[1, [2, 4, 9]], [2, [1, 2, 3, 4, 5, 6, 7, 8, 9]], [3, [2, 3, 6, 8, 9]], [4, [1, 2, 6, 7]],
                    [5, [2, 5, 6, 7, 9]], [6, [2, 3, 4, 5, 6, 9]], [7, [2, 4, 5, 7]], [8, [2, 3, 9]],
                    [9, [1, 2, 3, 5, 6, 8, 9]]]
neutronasc = [[1, [1, 5, 6, 8]], [2, 0], [3, [7, 4]], [4, [3, 5, 9]], [5, [1, 4]], [6, 1], [7, [3, 9]], [8, [1, 6]],
              [9, [4, 7]]]
incompativelnasc = [[1, [3, 7]], [2, 0], [3, [1, 5]], [4, [4, 8]], [5, [3, 8]], [6, [7, 8]], [7, [1, 6, 8]],
                    [8, [4, 5, 7, 8]], [9, [0]]]
cores = {1: 'Todos os tons de amarelo e laranja, castanho, dourado, verde, creme e branco.',
         2: 'Todos os tons de verde, creme, branco e cinza.',
         3: 'Violeta, vinho, púrpura e vermelha.',
         4: 'Azul, cinza, púrpura e ouro.',
         5: 'Todas as cores claras, cinza e prateado.',
         6: 'Rosa, azul e verde.',
         7: 'Verde, amarelo, branco, cinza e azul-claro.',
         8: 'Púrpura, cinza, azul, preto e castanho.',
         9: 'Vermelho, rosa, coral e vinho.',
         11: 'Branco, violeta e cores claras.',
         22: 'Violeta, branco e cores claras.'}
saude = {1: 'coração, cabeça, emocional.',
         2: 'rins, estômago, nervos.',
         3: 'garganta e fígado.',
         4: 'dentes, ossos, circulação',
         5: 'órgãos sexuais e nervos.',
         6: 'coração e pescoço',
         7: 'glândulas e nervos',
         8: 'estômago e nervos.'}

Motivação = {
    0: 'O número de motivação descreve os motivos que estão por trás das decisões que uma pessoa toma e do seu modo de proceder. É o número que corresponde à ação e a maneira que essa ação é desenvolvida.',
    1: 'É normalmente ambicioso, criativo, intelectual, original, inventivo, que não gosta de detalhes; quer liderar, dirigir, dominar, ser elogiado e, por vezes, é obstinado e autoritário.'
       ' O número um tem espírito pioneiro, gostando de explorar, não se importando com os riscos que possa ter na busca da sua realização. É movido pela iniciativa e pela criatividade.'
       ' Não gosta muito de receber ordens de quem quer que seja e trabalha melhor só ou em cargo de chefia. Ação, honestidade e lealdade, fazem parte desta excelente energia. Às vezes é incompreensível e se recusa a aceitar conselhos, fatos esses que com certeza causarão transtornos à sua carreira e aos seus planos profissionais.'
       ' Quando não desenvolvido espiritualmente, pode inclinar se para o egoísmo, vaidade e arrogância excessivas. Torna-se impaciente e dotado de pouca diplomacia e tato para tratar com as pessoas. Por esse motivo, o possuidor deste número enfrenta mais dificuldades do que os outros, principalmente no seu meio profissional, pois o seu jeito exigente, arrogante e ditatorial, vai fazendo inimizades por onde passa. Em alguns casos (não raros), essa dificuldade de convívio se estende aos familiares, amigos e companheiros afetivos.'
       '\nORIENTAÇÃO: Cultura, educação e refinamento pessoal são características indispensáveis que devem ser adquiridas para o seu triunfo pessoal, profissional e principalmente afetivo.',
    2: 'O dois anseia por amor e compreensão e profissionalmente sente-se mais à vontade em trabalhos cooperativos, na retaguarda, sem muito aparecer. Ele quer casamento, companheirismo, paz, harmonia e conforto. Manifesta a sua natureza sensível através da suavidade, cordialidade e prestatividade, sendo a sua principal característica a cooperação.'
       ' Pela sua “aparente”  passividade,  carência,  vulnerabilidade e calma natural, normalmente as pessoas que com ele convivem, quase sempre se aproveitam dessa sua natureza gentil e o usam para proveito próprio.'
       ' Normalmente não procura impor suas ideias (por ser conservador) preferindo escutar as dos outros antes de expor as próprias. Está sempre procurando reunir conhecimentos sobre todos os assuntos, e relaciona-se com todas as pessoas sem discriminar raça, credo, classe ou fortuna, numa só amizade e dedicação.'
       '\nORIENTAÇÃO: É muito importante para o seu desenvolvimento profissional e pessoal, que aprenda a conviver com o público (ser mais comunicativo) e espalhar conhecimentos para todos, levando-lhes sua mensagem de harmonia, paz e verdade.',
    3: 'Adora uma plateia; quer ser popular, ter muitos amigos e viver rodeado de beleza. Tem natureza impetuosa entusiástica e sonhadora. É um ser de alta capacidade artística, espírito humano elevado e está sempre procurando levar alegria e prazer aos que o rodeiam. É por demais sincero e tem espírito de liberdade e sociabilidade.'
       ' O número três sugere uma pessoa entusiástica, que gosta de se divertir, é sociável e amigável. Contudo, fazer escolhas   e amizades é um dos seus grandes desafios, pois como gosta  de interagir com os outros através de diversões, reuniões sociais e conversas íntimas, às vezes esses não estão ou não são receptivos aos seus interesses e, assim, se afastam desse tipo “aparecido”.'
       ' O 3 tem poderes psíquicos sobre os demais, principalmente no plano espiritual e, em vista desse poder, deve tomar  o máximo cuidado com o uso que faz com os mesmos, pois pode comprometer sua paz e seu viver harmônico, quase  sempre inadvertidamente, pois mesmo sendo espirituoso e sabendo lidar com as palavras, se sente tímido e inseguro, podendo se isolar ou ficar inibido, achando difícil expressar seus sentimentos.'
       ' A realização e a satisfação emocional normalmente derivam da alegria de compartilhar e aprender sobre o amor e a compaixão.'
       '\nORIENTAÇÃO: Pelo seu lado altruísta que não sabe dizer não, assumindo mais compromissos do que consegue cumprir, deve evitar os afazeres corriqueiros, estudar filosofia humana e fazer uma coisa de cada vez, indo até o fim, ou seja, terminar o que começa.',
    4: 'Precisa de ordem e de normas tanto em casa como no trabalho. O 4 quer trabalhar metodicamente e com afinco em favor dos outros e não gosta muito de inovações. É um conservador nato.'
       ' É realista e equilibrado e sempre podemos contar com ele. Profissional de alto gabarito é tremendamente dedicado à profissão, sendo de certa maneira perfeccionista nos detalhes e na conclusão de um trabalho bem feito. O ser número quatro gosta de praticidade, de coisas que pode ver, tocar e, principalmente, que se desenvolvem e que protegem. Num nível mais elevado, ele é o ser maduro, sério, firme como uma rocha, gostando de estar protegido contra os embates da vida, ou seja, não gosta de ser pego de surpresa em coisa alguma.'
       ' É tremendamente autodisciplinado, trabalhador incansável, obediente, perseverante, sincero, honesto, paciente, obstinado, conservador e deseja a todo custo estabilidade. Normalmente o quatro é astuto, preocupado com a própria segurança e de certa forma atraído para o mundo financeiro, industrial ou negócios em geral. Embora sejam indivíduos um tanto severos e que dificilmente demonstram emoções, são tremendamente leais, honestos e confiáveis.'
       '\nORIENTAÇÃO: A ambição material deve vir através dos seus talentos profissionais e esforços continuados. Através desse método pode conseguir tudo o que desejar, mas também deve ter planos que visem beneficiar toda a humanidade. Um outro fator de grande importância é que tendo tendência à teimosia e à rigidez (de ideias e ações), deve aprender a ser mais flexível e adaptável.',
    5: 'É um ser “mutável”, que gosta de variedade, de experiências incomuns e está sempre à procura de novas oportunidades. Tem percepção arguta, perspicaz e natural curiosidade, o que o leva a querer investigar e elucidar seja que problema for. Amante de novidades, está sempre entusiasmado com o que é novo, moderno, atrativo e incomum.'
       ' Poucos o acompanham o ser número cinco em matéria  de raciocínio rápido e perspicácia, pois ele se adapta a qualquer meio ou situação. No que tange à sociabilidade, o cinco   é versátil, bom falador, amigo, festeiro, fazendo amizades com grande facilidade e, desta maneira, sendo sempre bem-vindo em qualquer roda ou ocasião.'
       ' Como possui grande talento e versatilidade, sai-se sempre bem em qualquer função ou atividade, trabalhando somente naquilo que gosta e, assim, jamais se cansando ou mesmo se aborrecendo em suas atividades “corriqueiras”.'
       ' Liberal e de mente aberta, pode facilmente se adaptar a novas situações ou ter uma atitude progressista. Como é muito habilidoso e eficiente, pode aprender rapidamente e captar as situações com a sua abordagem realista. As viagens e mudanças são partes inerentes do seu estilo de vida.'
       '\nORIENTAÇÃO: Como tem fortes inclinações  psicológicas e filosóficas, deve adquirir conhecimentos nessas áreas través da dedicação aos estudos experimentais, ser perseverante em seus projetos e objetivos, terminando o que começa.',
    6: 'É, entre todos os números, o que procura sempre a harmonia pessoal e social. Os atributos principais deste número são o idealismo, a criatividade, o humanitarismo, a   compaixão e a visão. Como é idealista, tem como princípio orientar e consertar tudo o que está errado no mundo. Quer criar raízes e fazer com que sua vida gire em torno do seu lar e das pessoas queridas.'
       ' É conciliador, aprecia a fertilidade (adora família grande), tudo que cresce, produz e se expande, entusiasma-se com o milagre da vida, sentindo em tudo a presença do amor. Em vista desse predicado, age como conselheiro confiável e prático e tem amigos que sempre o procuram em épocas difíceis. Por índole, é simpático, intuitivo e persistente. Possui temperamento equilibrado, é muito responsável e inclinado aos problemas domésticos (adora uma casa bonita com muitas plantas, animais, bons móveis e tudo na mais perfeita ordem). Tem gostos refinados e se sente bem no conforto e na elegância. Conquista facilmente a popularidade e o êxito social.'
       '\nORIENTAÇÃO: Como tem tendência à vaidade e ao egoísmo, deve trabalhar esses atributos negativos através de estudos holísticos e metafísicos. Se aceitar suas limitações e conscientizar-se delas, equilibrará seus sentimentos e pensamentos e poderá ter a harmonia, a paz e a prosperidade tão desejadas.',
    7: 'Detesta ser mandado, de desconforto físico, barulho e confusão. É observador e pesquisador, gostando de descobrir o porquê de tudo. Não gosta de ter a liberdade tolhida, quer paz e tranquilidade para viver consigo mesmo, para poder sonhar e meditar.'
       ' O sete possui intuição, capacidade mental e racional. É considerado o número da espiritualidade e, por isso, normalmente os seus possuidores são introspectivos, silenciosos, perfeccionistas, arredios, estudiosos, possuindo uma sabedoria além do normal.'
       ' É por demais íntegro e tem o senso de justiça e clemência muito elevado. Está sempre querendo aprender e, principalmente, entender o que quer que seja. Resumindo, está sempre em busca de sabedoria.'
       ' Quanto a relacionamentos conjugais ou somente uniões (não oficiais), o número sete deve ter prudência e se unir somente após ponderar os prós e contras, pois se o parceiro não for totalmente compatível com ele, com certeza a união não durará muito, e isso o faz sofrer em demasia. Tem tendência ao ciúme infundado e exigir demais dos parceiros.'
       ' Como é muito íntegro, não tolera injustiças e está sempre pronto a defender os fracos e oprimidos.'
       '\nORIENTAÇÃO: As bebidas alcoólicas, cigarros e drogas, em geral, são venenos para o seu organismo; evite-os. Será muito mais feliz se viver próximo à água, seja ela de rio, lagoa ou mar. Quanto as relações conjugais, o sete deve buscar autoconhecimento nesse campo, pois como está sempre em busca de sabedoria e compreensão, esquece-se de que o parceiro também anseia pelos mesmos predicados.',
    8: 'Nasceu para o mundo dos grandes negócios e adora lutar contra seus opositores.  Normalmente é ambicioso, quer poder, riqueza e sucesso. Em virtude dessa ambição, está sempre motivado e determinado a seguir em frente, em busca dos seus objetivos.'
       ' O oito é arguto, inteligente, observador, conservador e sabe por instinto enfrentar os embates da vida sem muitos alardes ou desesperos. Não é chegado a ter muitos amigos, mas os que tem, lhe são tremendamente caros, defendendo-os e orientando-os em ocasiões de crises e perturbações'
       ' A sua normal ambição material o torna autoritário e com desejo de dominação. Quando culturalmente desenvolvido, naturalmente adquire espírito intelectual, analítico, bem equilibrado e se torna muito eficiente naquilo que se propõe a fazer. Como tem grande senso natural para os negócios, deve procurar desenvolver suas habilidades inatas de organização e execução.'
       '\nORIENTAÇÃO: Deve cultivar a persistência, o senso analítico e agir com diplomacia, paciência e tolerância para conseguir tudo o que deseja. Deve, ainda, aprender a perdoar, compreender e considerar as fraquezas das pessoas quando fizer um julgamento.',
    9: 'É o ápice da realização intelectual e espiritual. Busca sempre o conhecimento, quer ensinar, aconselhar e servir à humanidade. Tem um saber subconsciente que, se for desenvolvido, pode-se revelar genial. É magnético e carismático, mente perceptiva e habilidades psíquicas que apontam para uma receptividade universal.'
       ' O nove é impessoal, revestido de desinteresse material mas, quando não ciente desses predicados, ou seja, quando não desenvolvido espiritualmente, torna-se vaidoso, gostando de ser elogiado pelas pessoas que o rodeiam.'
       ' Embora seja um ser generoso e compassivo, com grande imaginação, pode se iludir ou ficar emocionalmente frustrado quando percebe que as outras pessoas não compreendem as suas perspectivas elevadas. Isso também se dá no âmbito profissional, pois deseja que os colegas ou superiores o compreendam (o que não é muito fácil!) e, assim, frustrando-se, acabando por desistir de suas excelentes ideias. Não tem medo de nada  nem de ninguém, e quando aparecem obstáculos em sua vida, sabe muito bem se desvencilhar deles com maestria e energia.'
       ' A combinação da inspiração e do idealismo com uma vida interior intensa e sonhos vividos, sugerem que o número nove é um ser universalista, desprovido de egoísmo e cheio de amor para dar.'
       '\nORIENTAÇÃO: Deve desenvolver suas qualidades inatas, viajar pelo mundo, conhecer novas pessoas, novos ambientes e novos meios sociais. Caso negligencie a estes “conselhos”, pode-se tornar um ser solitário, introspectivo, tendendo para o isolamento, depressão e desânimo.',
    11: 'O 11 tem uma postura visionária, pois gostaria de se ver rodeado de pessoas que o admiram, fazendo prevalecer seus pontos de vista, pois pouco valor dá às ideias alheias.'
        ' Normalmente é atrativo, tanto na parte pessoal como em sociedade. É um ser idealista, e a inspiração e a inovação estão sempre presentes. De certa maneira contraditório, pois gosta que suas ideias prevaleçam, por vezes se torna “conciliador” e árbitro das mais diversas contendas, se saindo bem nessa função. Uma combinação de humildade e confiança o desafia a trabalhar para ter autodomínio material e espiritual. Através da experiência, pode aprender a lidar com os dois lados da sua natureza e desenvolver uma atitude menos extremista, confiando nos seus sentimentos.'
        ' É adepto da harmonia das formas, dos métodos persuasivos, da elevação moral das coisas e pessoas e de tudo que tem sentido superior. Apesar de não concordar, é muito vulnerável na sua sensibilidade, magoando-se profundamente com fatos que a outros nem sequer os preocupariam.'
        '\nORIENTAÇÃO: Como tem rara intuição e capacidade psíquica que deve desenvolver, é aconselhável seguir seus pressentimentos, sem se ater a conselhos alheios ou ideias preconcebidas. Como o medo sempre se relaciona à insegurança em relação ao dinheiro, precisa aprender a superar a tendência a ser arrogante ou calculista. Aprenda a expressar o seu talento único para liberar o seu verdadeiro potencial.',
    22: 'Apesar de ser reservado, muitas vezes é carinhoso e se preocupa com o bem-estar e a segurança de toda a humanidade, mas nunca perde de vista a sua posição pragmática ou realista...É prático, habilidoso, honesto, cordial, idealista, inspirado e um eficiente organizador com um grande potencial de realização. O 22, que é um número-mestre, quer dar a sua contribuição para o gênero humano.'
        ' Tem necessidade de afeto, carinho e amor. É trabalhador honesto, com capacidade de liderança inata, carismático e com profunda compreensão das pessoas e de suas motivações. Geralmente culto e mundano, tem muitos amigos e admiradores. Quando determinado, pode atingir o sucesso e a fortuna com a ajuda e o encorajamento dos companheiros de jornada.'
        ' A profundidade de seus sentimentos e necessidade de expressão indicam que é um ser dinâmico e motivado quando se sente inspirado. A sua capacidade de resistência e perseverança sugere que você mostre a sua verdadeira personalidade em momentos difíceis e estressantes. Embora seja generoso e entusiasmado, às vezes pode ser egoísta e arrogante.'
        '\nORIENTAÇÃO: Se tiver padrões muito elevados, pode sentir-se insatisfeito e se tornar crítico ou antipático. Por isso, deve cultivar a compreensão, a determinação e a persistência. Deve trabalhar sempre em benefício da humanidade, sem se esquecer de si próprio.'
}

Impressão = {

    0: 'É o número que descreve o que está oculto no ser humano e a imagem que uma pessoa tem de si mesma (geralmente sem perceber). Revela, ainda, a primeira impressão que os outros têm de nós, antes de nos conhecerem na realidade, ou seja, a condenação ou a absolvição antes do julgamento.',
    1: 'A impressão deste ser, é exatamente isto: um ser superior que se destaca dos demais, fazendo com que todos, invariavelmente, se virem para ele. É líder por natureza e vive patrocinando novas ideias, sejam elas suas ou de outrem. Nessa busca, sonha em ser corajoso, ousado, justo, leal e original, tanto nas qualidades profissionais, quanto no seu modo de se vestir e no falar. É persistente e dificilmente se deixa dominar pelo desânimo, por isso, para conseguir o que deseja, pode mostrar-se teimoso, egoísta, obstinado e por vezes opressor. Sabe mandar, dirigir e tem aparência sólida e autoconfiante.'
       ' Entre os amigos (que são muitos) é sempre escolhido para liderar, pois as suas ideias e posturas sólidas, lhe conferem as prerrogativas inerentes aos vencedores. Mostra grande potencial e habilidade executiva e de liderança, que pode se manifestar em trabalhos de especialização na sua área de atuação ou em postos de gerenciamento e administração, na carreira militar ou política.'
       '\nORIENTAÇÃO: Precisa aprender que o mundo não gira ao seu redor, e evitar a inclinação a ser autocentrado e ditatorial. Deve também policiar seus desejos gastronômicos, pois sendo o centro das atenções em festas e reuniões, pode-se deixar levar pela gula e sofrer de obesidade após os 50 anos.',
    2: 'O 2 quer amor, amizade, carinho, harmonia e paz. Tem feições que o destacam de todos os outros números: é ardoroso, idealista, calado, pacífico, diplomata, sendo muito hábil em conciliações de interesses seus e alheios. Na tentativa de agradar as pessoas de quem gosta, corre o risco de se tornar excessivamente dependente. Contudo, se desenvolver a autoconfiança, pode superar essa tendência e vir a ser uma pessoa bem-sucedida, sem grandes problemas.'
       ' É também responsável, justo, gosta de calma, de legitimidade, esclarecimento e também inspiração. Gosta também de se sentir compreendido e confortado. Quando só, parece sonhador, trazendo a cabeça nas “nuvens”. Pelos modos calmos, gentis e agradáveis, desperta em todos a segurança e a confiança. Normalmente age com naturalidade e simplicidade, não gostando de se exibir, seja em família ou numa roda social. A discrição é seu lema.'
       '\nORIENTAÇÃO: Infelizmente, em nossa sociedade, a simplicidade é vista como sinônimo de fraqueza. Por isso, deve aprender a ser mais dinâmico, arrojado, criativo e sempre procurando inovar, sem medo de errar ou ser diferente. Evite a estagnação ou a monotonia.',
    3: 'É do tipo social, amistoso, alegre e jovial. O 3 é popular, atraente, criativo, gostando de ser aplaudido, notado, admirado, ou seja, que o vejam como um verdadeiro artista. Normalmente cultiva a criatividade, os contatos sociais e a expressão de suas próprias ideias e sentimentos, especialmente através de alguma forma de arte.'
       ' Demonstra personalidade e um espírito de grande profundidade, justiça, esperança, filantropia, alegria e felicidade. Abençoado e com charme, sensibilidade e imaginação fértil, o três precisa equilibrar essas qualidades com alguma forma de fundação sólida na vida. Embora seja, em geral, artístico, charmoso e elegante, com excelentes relacionamentos, e tenha um bom-senso de humor, talvez precise desenvolver a autoestima e evitar tendências à preocupação ou insegurança emocional.'
       '\nORIENTAÇÃO: deve evitar constrangimentos e aborrecimentos, além da preocupação excessiva em relação aos projetos futuros, pois esses sentimentos perniciosos podem lhe prejudicar sobremaneira o fígado e o coração.',
    4: 'O 4 é o número da consumação e da manifestação da luz. É o número do entendimento e da ordem; é a chave que abrirá muitas portas mágicas fechadas ao homem comum. É o número dos seres rígidos, fortes, seguros, objetivos, trabalhadores, conservadores e dignos de confiança; o sustentáculo da família, da empresa e dos amigos.'
       ' É naturalmente preciso, simples, comedido, sendo apreciador de boa música, teatro e gastronomia. É elegante no vestir, preferindo o clássico ao moderno. A sólida estrutura e capacidade de organização mostra que o quatro precisa de estabilidade e ordem. Como a natureza o dotou de grande energia, habilidades práticas e forte determinação, com certeza conseguirá, com estes predicados, atingir seus objetivos. É detalhista, ordeiro, natural no andar, de gestos simples e graciosos. Bom ouvinte está sempre interessado nas conversas alheias quando estas lhe parecem sinceras e concretas.'
       '\nORIENTAÇÃO: Como é o esteio da família, deve ser empreendedor (colocar as ideias em ação), mesmo que a princípio pareçam “loucas”. Cuidado com a obesidade, excesso de autoridade e de controlar em demasia a vida dos outros.',
    5: 'É um número mágico e peculiar, que era usado pelos gregos e romanos como amuleto para proteger o portador dos espíritos malignos. O dono deste número é normalmente intuitivo, fazendo-o conhecer o âmago, o caráter e sentimentos alheios. É um ser cativante, interessante, agradável e por que não dizer, magnético. Sonha em viajar por todos os recantos do planeta levando uma vida de aventuras, sem vínculos ou até mesmo sem muita responsabilidade (não que não a tenha).'
       ' Sabe se defender dos inimigos é esperto e dificilmente consegue ser enganado. É também muito leal e sabe fazer amizades e também inimizades com muita facilidade. Está sempre à procura do que está além da superfície dos seres e das coisas, gostando de se mostrar antecipado, ora criando, descobrindo ou ditando modas. Uma das suas maiores virtudes é o entusiasmo com que encara qualquer inovação ou novidade, além de estar sempre bem informado sobre todas as coisas, sendo muito difícil surpreendê-lo.'
       ' Mesmo quando a idade chegar, terá sempre uma aparência alegre e juvenil. O desejo de explorar ou experimentar qualquer novidade e a abordagem entusiástica, sugerem que a vida tem muito a lhe oferecer. O charme é dos seus atributos naturais e, por isso, não tem problemas em atrair admiradores e amantes.'
       '\nORIENTAÇÃO: Como tem grande senso de humor e é atrativo fisicamente, precisa ser seletivo na escolha de amizades e parceiros para seus projetos.',
    6: 'O seis é visto como a perfeição dos números pelos cabalistas e o filósofo grego Nicomachus o chama de Vênus, deusa a quem era consagrado e razão pela qual é tido como o número do amor. O portador deste número é elegante, atraente, do tipo paternal e em certas ocasiões, encantador. Sonha com uma família maravilhosa, uma bela casa, com varandas, flores e sempre pronta para receber os amigos. É de certa maneira contraditório, pois enquanto tem o dom da harmonia, também atrai confusão; é o número da sedução, do vício e da virtude, das incertezas no casamento e também do amor puro e simples; gosta de segurança, tem senso de responsabilidade, ama o social (tem tendência à extravagância) e não suporta viver só.'
       ' O seis possui uma nobreza interior que se mostra especialmente quando ocupa posição de liderança ou de responsabilidade. Encara o trabalho com seriedade e dá o melhor de si quando tem liberdade de ação. Quando aprende a examinar todos os fatos de qualquer situação difícil ou a ceder em vez de se lançar em jogos de poder, tem resultados mais positivos.'
       '\nORIENTAÇÃO: Como é naturalmente generoso e por vezes ingênuo, dando valor a pessoas que não merecem, inclusive sendo explorado por eles, deve aprender a dizer “não”, a ouvir, meditar e só depois de muito pensar, tomar resoluções, ou seja, não agir precipitadamente. Como é digno de caráter, pensa que todos o são, mas não é verdade.',
    7: 'É um número classificado como próspero e geralmente é visto assim. É um número inteiramente religioso e como tal foi estimado pelos antigos; representa o triunfo do espírito sobre a matéria. Tem facilidade de adquirir numerosas simpatias no mundo social. Seu grande desejo é se ver cercado de livros numa atmosfera de paz, beleza e tranquilidade, e isso o leva a imaginar-se sábio e também o refúgio aonde todos venham consultar. É tremendamente místico, de natureza meditativa e também solitária, apesar de gostar da companhia dos amigos e também de elogios vivendo à procura do entendimento e também do esclarecimento de suas dúvidas e conflitos. Dispõe de natureza expressiva, única e imponente, pois a sua personalidade sete, não pode ser imitada pelos que não pertencem à mesma vibração. Possui personalidade exigente, sinceridade e honestidade com todos, exigindo, é claro, reciprocidade. Tem natureza distinta, delicada, sóbria, não gostando muito de aparecer.'
       ' Como é uma pessoa sensível, com sentimentos intensos e profundos, aprende a relaxar procurando ser alegre e evitando sobrecarregar sua mente quando enfrenta situações estressantes. A tendência a ser enigmático ou dissimulado, muitas vezes o leva a desenvolver a arte de fazer perguntas sutis sem deixar que ninguém saiba o que realmente está pensando.'
       '\nORIENTAÇÃO: Por ser detalhista, aparentemente descrente e de certa forma excêntrico, deve aprender a compreender os outros e aceitar como eles são, evitando a crítica, a calúnia ou dando conselhos a quem não lhe pediu.',
    8: 'Este número é peculiar e visto como de grande poder pelos antigos gregos e cabalistas, que diziam: “Todas as coisas são oito”. O oito aparenta imponência e influência, mesmo que esteja com problemas pessoais ou financeiros. Tem raciocínio vivo e perspicaz, e seus modos são calmos e equilibrados. Sonha em ser o comandante de grande empreendimento comercial, um poderoso executivo, sempre cercado de tudo e de todos e chefiando numerosos empregados.'
       ' É ordeiro, prático e está sempre procurando fazer do seu local de trabalho um lugar agradável e eficiente, onde existam todas as coisas necessárias ao bom desenvolvimento profissional. Quando é atingido em seu amor próprio, é passível de impulsos de mau humor e até certas repulsas a brincadeiras. Altamente criativo, o oito tem uma força vigorosa que o impele a sempre buscar novas formas de realização. Naturalmente dramático, mas com aversão a obedecer aos outros, não gosta de receber críticas. Tem uma incrível habilidade para lidar com as pessoas e fazer contatos, mas pode alterar entre ser caloroso e atencioso e ser frio e indiferente.'
       ' Gosta de ser notado e de dar impressão de estar bem financeiramente. Tem senso natural para os negócios e se beneficiará muito se desenvolver suas habilidades executivas e de organização.'
       '\nORIENTAÇÃO: Como a maioria das pessoas o vê como “superior”, tal superioridade não deve ser “arrogante” nem servir para menosprezar quem quer que seja. É imperioso que estude as Ciências Ocultas, a espiritualidade e desenvolva qualquer atividade humanitária.',
    9: 'É o ser que atrai sempre simpatias e antipatias na mesma proporção. Sonha em ser um humanitarista, ávido por servir, compreensivo em relação aos sofrimentos alheios e uma fonte de conforto e aconselhamento. Realiza-se pelo esforço que emprega em resolver os problemas alheios, dando-lhes amizade e carinho. É um visionário, sempre em busca de harmonia, esclarecimento e analisando tudo e todos, procurando desvendar os problemas mais obscuros e incompreendidos da Natureza.'
       ' Possui rara inteligência e também profundos conhecimentos sobre os mais diversos assuntos, mesmo que tenha tido uma educação insuficiente. Pela sua aparência (imponente), personalidade marcante, sabedoria, bondade, compreensão e ponderação, normalmente inspira confiança em todos que o conhecem, podendo produzir admiráveis resultados, pois normalmente atinge o fim a que se propõe. Quando vê as coisas correrem como quer, normalmente se sente leve e sabe julgar o caráter das pessoas. Isso lhe permite ser sociável e voltado para as coisas públicas, características que, quando combinadas, o levam a ajudar muitas pessoas.'
       '\nORIENTAÇÃO: Se tiver os seus desejos reprimidos ou limitados, tende a se comportar de forma rebelde ou a se tornar temperamental. Precisa ser mais flexível e menos orgulhoso. Adapte-se a situações novas e aprenda a esquecer o passado. Evite a cor preta, principalmente durante o dia.'

}
Expressão = {
    0: 'Descreve a maneira como um ser humano interage com outro. Ele diz quais são seus verdadeiros talentos e qual a melhor forma de expressá-los',
    1: 'Pontos positivos: liderança, criatividade, atitudes progressistas, otimismo, vigor físico, convicções fortes, competitividade, independência e sociabilidade.\n'
       '\nPontos negativos: arrogância, tendência ao ciúme, egocentrismo, antagonismo, falta de controle em ocasiões que deveria manter a paciência, instabilidade emocional e certa impaciência.\n'
       'O número um corresponde ao líder, às pessoas influentes, pioneiras, inventivas e planejadoras – embora muitas vezes essas pessoas realizem seus projetos sem levar em conta os envolvidos. Tendem a dominar, consciente ou inconscientemente, todos os seus conhecidos. Em vista da sua tremenda ousadia, vive a sonhar com grandes empreendimentos, pioneiros, sempre à procura de novidades, visando o presente e o futuro.'
       'O seu êxito profissional, afetivo ou pessoal, depende exclusivamente do seu modo de pensar, falar e proceder, pois como é individualista, dificilmente aceita conselhos ou sugestões de quem quer que seja. Em virtude da sua personalidade de aspecto arrogante e ditatorial, encontrará, ao longo da vida inúmeras dificuldades de relacionamento, sejam elas de cunho afetivo, pessoal ou mesmo profissional, levando-o a fazer amigos e inimigos com grande facilidade. Trabalhos subalternos, sem movimentação (como contador, caixa de banco, etc.), o deprimem, e diríamos até o deixam “estressado”, pouco à vontade e insatisfeito. Deve trabalhar por conta própria ou em cargos de chefia, pois como tem os ouvidos “sensíveis”, não gosta de receber ordens, principalmente quando essa ordem parte de alguém com estudo inferior. Possui emoções poderosas e habilidades para dar e receber amor. Logo, é importante encontrar uma forma de expressar os seus sentimentos em vez de buscar realização em interesses financeiros. Se ficar desanimado, pode cair nos dramas emocionais de outras pessoas, pois a sua sensibilidade precisa encontrar uma via de escape. A ordem e o método são, em geral, benefícios para você, embora um plano de vida definido seja um pré-requisito para fazer bom uso do seu tremendo potencial.'
       '\nORIENTAÇÃO: Em vista de ter tendência a começar muitas coisas e não terminá-las, deve adquirir conhecimentos, prudência, ponderação e sabedoria, para poder desenvolver o seu imenso potencial de realização e levar a bom termo os seus objetivos.',
    2: 'Pontos positivos: boas parcerias, gentileza, tato, receptividade, intuição, consideração, harmonia, natureza agradável, que demonstra sempre boa vontade em ajudar os outros.\n'
       '\nPontos negativos: desconfiança, subserviência, timidez, egoísmo, tendência a se magoar facilmente, ilusão e excesso de sensibilidade.\n'
       'O dois, apesar de não ser ambicioso, consegue tudo que deseja, mas sempre pela persuasão, nunca pela força. No lado negativo de seu caráter, destaca-se a hesitação constante e a falta de iniciativa, tendo tendência para adiar decisões importantes por qualquer motivo, levando-o a situações delicadas, principalmente no âmbito profissional e afetivo. Normalmente usa de destreza e sabedoria, trabalhando mais nos “bastidores”, ou seja, na retaguarda, para levar alegria e amor a todos os seus conhecidos. É cooperativo, de certa maneira tímido, vulnerável e até passivo, porém sempre atento a todos os detalhes de seu ambiente. Nasceu para trabalhar em conjunto com outras pessoas, dificilmente se adaptando ao comando, seja de que nível for. É detalhista e prefere fazer uso das suas habilidades inatas lidando com pessoas e satisfazendo seu desejo de eterna harmonia. Parte do seu sucesso provém do fato de aplicar suas habilidades e saber transformar a teoria em prática. Além   disso, o ser dois é muito trabalhador, sempre pensando no sucesso; é, também, de certa forma ambiciosa e com grande poder de liderança. Atento aos detalhes, é inventivo e sabe como ninguém solucionar problemas, seja de que nível for. Como é uma pessoa leal, grande parte da sua segurança é construída com amor e amizade; bons relacionamentos são importantes para a sua realização.'
       '\nORIENTAÇÃO: O dois era chamado pelos antigos de a “mãe dos números e do casamento”. Como a sua grande virtude é a compreensão dos processos do conhecimento oculto, deve sempre cultivar este dom, a fim de poder viver condignamente, sem esperar que alguém ou algo o faça movimentar-se.',
    3: 'Pontos positivos: humor, felicidade, amizade, lealdade produtividade, criatividade, franqueza, talento artístico, força de vontade, amor à amizade e talento com as palavras.\n'
       '\nPontos negativos: excesso de imaginação, tendência a se entediar com facilidade, tendência a ser pouco afetivo, extravagância, comodismo, preguiça, hipocrisia, vaidade e ostentação. O três é o número da mais alta sabedoria e valor, da harmonia e da ação, do amor perfeito, da ternura e da força d’alma.\n'
       'É tremendamente apaixonado pela família (mesmo que não aparente), gosta de receber e dar amor e sente-se muito feliz em poder contribuir para a felicidade alheia. É ambicioso e dificilmente desiste dos seus ideais, seus objetivos. É um ser criativo e normalmente tem muitos amigos divertidos, além de possuir uma elevada percepção de beleza, cor e forma. O aspecto leve e espirituoso da sua personalidade sugere que gosta de se divertir e de se comunicar, e esses predicados, quando num elemento culturalmente desenvolvido, leva-o à procura de respostas para as questões mais profundas da vida, o que, no final das contas, acaba levando-o a explorar áreas mais espirituais ou místicas. Detesta receber ordens e trabalhar em uma profissão ou posição inferior, pois como sabe ter poderes e qualidades superiores, também sabe como tirar proveito de qualquer situação que o possa fazer feliz.'
       '\nORIENTAÇÃO: A vontade de aprender e a aversão à rotina indicam que você pode passar por diversas experiências na sua busca por uma carreira que o desafie mentalmente.Com a maneira persuasiva com que usa as palavras, pode ser bem sucedido no mundo da comunicação, vendas, escrita, arte ou área editorial, sendo, é claro, o chefe ou trabalhando por conta própria, pois tem grande aversão a ser mandado.',
    4: 'Pontos positivos: capacidade de organização, autodisciplina, firmeza de caráter, trabalhador, sincero, paciente, confiável,  perseverante, obstinado e conservador.\n'
       '\nPontos negativos: comportamento destrutivo, pouco comunicativo, reprimido, rígido, pouca sensibilidade, procrastinação, muito econômico, autoritário, afeições ocultas e ressentimentos.\n'
       'O quatro é o número da perseverança e da imortalidade, da descoberta e da consumação, da firmeza e do propósito, da realização das esperanças, das regras, dos poderes e da vontade. O novo lhe assusta e sente-se mais à vontade lidando com atividades rotineiras e já consagradas pela experiência. Gosta de construir coisas, mesmo que aparentemente não tenham maior serventia, pois é dotado de grande habilidade manual. Não suporta pessoas hipócritas, desonestas, preguiçosas e com falta de visão do futuro, e estas deficiências o tornam deprimido e triste. Não gosta de discussões fúteis, sem motivo aparente. Jamais guarda rancor, ressentimentos ou ódio de qualquer espécie, esquecendo rapidamente qualquer ofensa ou constrangimento. Sendo um espírito superior, o seu amor ao próximo é, simultaneamente, glória e dever. Usualmente, pode cooperar e se sair bem em trabalhos de equipe; entretanto, como não gosta de receber ordens, é preferível trabalhar por conta própria. É muito importante ser original, criar as suas próprias ideias e tomar as próprias decisões. Precisa aprender a ser persuasivo e não vigoroso ou autoritário.'
       '\nORIENTAÇÃO: Como gosta da sua liberdade, evite fazer julgamentos ou criticar demais as pessoas, pois elas podem ficar ressentidas com as suas decisões e, assim, não lhe prestarem bons serviços ou não colaborarem com seus projetos.',
    5: 'Pontos positivos: versatilidade, adaptabilidade, instintos fortes, sedução, sorte, ousadia, amor pela liberdade, espiritualidade, curiosidade, sociabilidade e misticismo.\n'
       '\nPontos negativos: falta de confiabilidade, mutação, procrastinação, inconsistência, libidinosidade, confiança extrema e teimosia.\n'
       'O cinco é o número dos criativos, dos ousados, dos agitados e dos amantes da liberdade. A característica deste número o impele a viajar pelo mundo sempre em busca de saber, de novas experiências e também prazer. É um ser sociável, agradável e sempre bem-vindo em festas e reuniões. Gosta de ocupações diferentes, aquelas que o permitem estar em contato com pessoas, com o público, e que o deixem agir e exprimir-se livremente. Detesta receber ordens, principalmente de pessoas com estudo ou capacidade inferior às suas. Também não gosta de trabalhos pesados, enfadonhos, cansativos, preferindo os intelectuais, ou aqueles que o colocam em destaque, como chefe de vendas, de marketing, de criação, ou qualquer outro que implique no desejo de viver e pesquisar. Possui grande capacidade para lidar com quaisquer tipos de pessoas, sejam elas ricas, influentes ou intelectuais. O cinco, às vezes, tem surtos extraordinários de energia e coragem que lhe permitem reagir rapidamente a situações e agarrar as oportunidades. Possui uma mente excelente e precisa, e é capaz de ter pensamentos profundos, além de uma inclinação a pensar de forma técnica e analítica. O seu espírito independente o estimula a ver as coisas de uma maneira original e a buscar a liberdade pessoal. Porém, deve evitar ser voluntarioso ou obstinado e se tornar negativo e pouco comunicativo.'
       '\nORIENTAÇÃO: Como terá muitas experiências na vida, deve usar essas experiências e o seu temperamento corajoso para levar até o fim os seus projetos, pois é de certa forma dispersivo, não conseguindo terminar o que começa, tendo mais começo do que fins.',
    6: 'Pontos positivos: mundanismo, amabilidade, compaixão, confiabilidade, compreensão, solidariedade, idealismo, vida doméstica, humanitarismo, senso artístico, equilíbrio.\n'
       '\nPontos negativos: descontentamento, ansiedade, timidez, teimosia, franqueza excessiva, perfeccionismo, dominação, egoísmo, desconfiança, cinismo e egocentrismo.\n'
       'O seis é um número de certa maneira contraditório; é o número da confusão e junção, da união e sedução, do vício e  da virtude, da incerteza no casamento, do amor, da atração dos sexos e da beleza. Significa todos os tipos de problemas e discórdias, mas também é capaz de purificação pelo conhecimento e de uma vida elevada. A família é a base de sustentação de toda a sua vida. Amar e ser amado; eis o desejo maior do número seis. E quando privado desse amor, sente-se completamente perdido, sem saber o que fazer e como agir. Aparentemente calmo, pode explodir quando o contradizem ou criticam, principalmente quando nasce nos dias 29, 30, 31 ou primeiro de cada mês. O seis quer ver todos saudáveis, alegres e repletos de sucesso e felicidade. É pródigo em favorecer os outros, nada esperando em retribuição, passando inúmeras vezes por “otário”, pois as pessoas que favorece dificilmente retribuem os favores recebidos. Apesar de suas qualidades expressivas e brilhantes, pode ter dificuldades em fazer escolhas e tomar decisões. Como tem diversos interesses que o levam a diferentes direções, é muito importante ter um senso de propósito. Sem isso, pode ficar dividido entre os seus ideais e o desejo de satisfação material. Um lado da sua natureza pode ter uma forte atração por dinheiro, luxo e um estilo de vida indolente, mas o outro tem um desejo inspirador que pode levá-lo a trabalhar duro para realizar seus ideais. Qualquer que seja a sua escolha, terá muitas oportunidades e a habilidade de driblar as situações difíceis, graças ao seu charme e persuasão. Excelente marido, esposa ou amante, entrega-se por inteiro aos seus amores ou amizades, exigindo a mesma dedicação por parte dos outros, o que nem sempre acontece, frustrando-o e desestabilizando-o emocionalmente por longos períodos.'
       '\nORIENTAÇÃ: Para se elevar e conquistar fama e fortuna deve procurar profissões de interesse humanitário, sempre incentivado pelos que ama. Deve, também, concentrar sua energia e não se preocupar com o que outras pessoas dizem ou fazem. Aprenda a dizer “não” sem medo de ofender os outros.',
    7: 'Pontos positivos: educação, confiança, meticulosidade, idealismo, honestidade, poderes psíquicos, cientificidade, racionalidade, reflexão e pensador silencioso.\n'
       '\nPontos negativos: dissimulação, engano, pouco amável, fingimento, ceticismo, confusão e falta de sentimentos.\n'
       'É o número da realeza e do triunfo, da fama e da honra, da reputação e da vitória. O sete é perfeccionista, de certa maneira arredio, por vezes calado, demora a esquentar um relacionamento, interessado em filosofia, religiosidade, política e assuntos extra físicos. De grande intuição e discernimento, raramente segue as ideias alheias e só faz o que acredita. É vibrante em seus propósitos, honesto, leal, amigo, profundo, elevado moralmente, não gostando de futilidades e ilegalidades, e as atividades ou atos que as contenham, o deprimem e o entristecem. Como é um ator natural, muitas vezes passa uma imagem de autoconfiança. Em geral é perspicaz e inteligente, com um forte sexto sentido, pode avaliar rapidamente as pessoas e situações. Isso faz com que tenha um desempenho melhor quando ocupa uma posição de liderança em vez de se desgastar em trabalhos físicos. Mesmo que seja generoso, tanto quanto à sua disponibilidade de tempo quanto de dinheiro, pode se tornar expansivo e indulgente em excesso. Felizmente, é aberto às críticas, o que lhe permite usar a autoanálise como uma ferramenta valiosa, com poder de influenciar os demais. Por vezes é incompreendido, considerado estranho, um ser solitário, incomum, diferente.'
       '\nORIENTAÇÃO: Deve controlar a sua vida e estabelecer primeiro uma base sólida. Deve aproveitar o que a vida tem   de melhor e evitar o comodismo. Deve trabalhar em prol da humanidade, em serviços que beneficiem a coletividade e que de alguma maneira o satisfaçam intelectual e também espiritualmente. Precisa meditar e receber inspiração do eu interior.',
    8: 'Pontos positivos: liderança, meticulosidade, trabalho árduo, tradição, autoridade, proteção, poder de cura, bom juízo de valores.\n'
       '\nPontos negativos: impaciência, intolerância, mesquinhez, inquietude, excesso de trabalho, dominação, falta de planejamento, controlador.\n'
       'O oito é considerado o número da atração e da repulsão, da vida, de todos os tipos de lutas, da separação, ruptura, destruição, expectativas e ameaças. Também é conhecido como o “portão da eternidade”. Dificilmente consegue expressar seus sentimentos, parecendo frio e evasivo; a realidade é outra: é muito carinhoso, meigo, amoroso, carente de afeição e com muito calor humano. Em virtude destas características (ser uma coisa e parecer outra), tem tendência ao materialismo, ao acúmulo de bens materiais e também uma certa ganância de querer possuir a maior quantidade de dinheiro e bens possíveis. É organizadíssimo e muito dedicado à profissão, sendo excelente para lidar com situações concretas que requerem discernimento e constância. É austero, prático, direto, seguro de si, por vezes generoso, mas também se pode tornar violento quando as coisas não correm como deseja. Como é dotado de alto senso de justiça, jamais deve fazer qualquer transação comercial fraudulenta, pois essa “falta” poderá lhe ser muito prejudicial ao longo da existência. Normalmente não perde tempo lutando pela fortuna, pois sabe com absoluta certeza que conseguirá atingir seus objetivos. Com excelente capacidade de avaliação, lida com o dinheiro e as questões materiais de forma instintiva. Pode ter golpes de sorte ocasionais, combinando o seu sexto sentido com informações inesperadas. Como pode passar por períodos de instabilidade financeira, precisa equacionar o problema de como viver bem sem ser extravagante. A sua autoridade natural o leva a ocupar posições de responsabilidade, nas quais pode usar a sua capacidade de organização. E será melhor se isso envolver algum tipo de atividade criativa.'
       '\nORIENTAÇÃO: Como é progressista, dá grande importância ao dinheiro, ao status social e deseja constantemente expressar livremente os seus sentimentos, deve evitar situações monótonas ou excesso de rigidez. Terá mais chance de se realizar como alto executivo, comerciante ou trabalhando por conta própria.',
    9: 'Pontos positivos: idealismo, humanitarismo, criatividade, sensibilidade, generosidade, sedução, caridade, desapego e popularidade.\n'
       '\nPontos negativos: frustração, nervosismo, fragmentação, insegurança, egoísmo, pouca praticidade, complexo de inferioridade, medos e preocupação exagerada.\n'
       'O nove é misterioso, silencioso e liga-se ao plano astral. É o número da sabedoria, da virtude, da experiência, dos mistérios, da moralidade, do valor, da soberania, do amor humano, da obscuridade, da proteção e dos frutos do mérito. Enfim, é o número dos universalistas, dos generosos, dos seres humanos sensíveis aos sofrimentos alheios e vivem constantemente em busca da “verdade”, da compreensão, adaptando-se a trabalhos que visam o engrandecimento do ser humano, não medindo esforços e sacrifícios para fazer o que pensa ser certo, ou seja, amar, proteger e defender o próximo. O nove é um pensador nato e dispõe de grande capacidade de raciocínio, observação, criatividade e compreensão, levando-o a conhecer o âmago e as qualidades de coisas e fatos. Deve sempre ser impessoal, pois como é universalista, o materialismo pode-lhe causar sérios desapontamentos e aborrecimentos, principalmente com respeito a “amigos” e parentes que considera íntegros e incapazes de ações ilícitas, mas que na realidade são falsos. Tem excelente capacidade para ganhar dinheiro, e também para perdê-lo. Mas o mais extraordinário, é o fato de “nunca” lhe faltar dinheiro para as despesas do dia-a-dia. A fé interior é um importante fator para a sua autoconfiança e sem ela pode passar por fases de insegurança, reserva ou falta de autoestima. Perspicaz e de mente aguçada, entende as coisas rapidamente e as utiliza em seu benefício. Embora às vezes possa se preocupar muito com questões materiais, é através do uso da sabedoria interior e compreensão espiritual que pode superar   a perspectiva fria ou cética. Ousando ser espontâneo e justo ao mesmo tempo, mas sem perder a competitividade, pode desafiar a si mesmo e aos outros sem se sentir arrogante ou poderoso. Frequentemente atraído por indivíduos criativos, tem grande necessidade de amor e compreensão que se opõe à sua aparência forte e assertiva.'
       '\nORIENTAÇÃO: Como nasceu para ser livre, não deve se apegar a nada nem a ninguém, e sim amar e trabalhar em prol de toda a humanidade, sem nada esperar em troca. Evite ser destrutivo ou vingativo. Tente perceber a necessidade das outras pessoas e evite ser autocentrado ou crítico; evite dominar os outros com as suas opiniões. A compaixão e o amor trazem as maiores recompensas.',
    11: 'Pontos positivos: equilíbrio, concentração, objetividade, entusiasmo, espiritualidade, idealismo, intuitividade, inteligência, expansividade, inventividade, senso artístico, prestatividade, capacidade de cura, humanitarismo e poderes psíquicos.\n'
        '\nPontos negativos: complexo de superioridade, falta de objetivos, emotividade excessiva, magoa-se facilmente, falta de clareza, dominação e hipersensibilidade.\n'
        'O onze é o número do poder, da bravura, da energia, do sucesso em aventuras destemidas, da liberdade. Suas vibrações são lunares e seus possuidores são idealistas, intuitivos sensíveis, místicos, imaginativos e pensadores inspirados. O ser onze dificilmente é bem-sucedido como negociante, sentindo-se mais à vontade em profissões agitadas, como política, marketing, administração de grandes empresas ou grandes projetos esotéricos, sociais e ideológicos. Profundo conhecedor da mente humana, normalmente se destaca dos demais, pois sendo um “número mestre”, consegue vislumbrar em seus amigos e inimigos, defeitos e virtudes que em outros números seria impossível. Tem fé em seus ideais e dificilmente volta atrás em suas decisões. Procura sempre ser justo, leal, compreensivo e viver em harmonia com todos. A vibração especial deste número sugere que o idealismo, a inspiração e a inovação são altamente importantes para o seu possuidor. Uma combinação de humildade e confiança o desafia a trabalhar para alcançar a independência financeira e espiritual. Através da experiência, pode aprender a lidar com os dois lados da sua natureza e desenvolver uma atitude menos extremista, acreditando nos seus sentimentos. Em geral, o onze tem muita energia e gosta de vitalidade, mas deve evitar o excesso de ansiedade e a falta de sentido prático. Se não seguir os dogmas esotéricos e estudos profundos do “hermetismo”, pode-se tornar indolente, indeciso e sujeito à vontade alheia, sendo sugado e feito de empregado por seres muito mais inferiores do que ele. Como é um ser superior, normalmente atrai inveja e conflito das pessoas ao seu redor, fazendo amizades e inimizades com grande facilidade e, assim, dificultando a conclusão dos seus objetivos.'
        '\nORIENTAÇÃO: Como todos os grandes seres humanos, aqueles que se destacam dos demais, o onze também tem certas dificuldades de adaptação ao meio ambiente, ao convívio com seres de inferior categoria, que não o compreendem, não o apoiam e ainda o criticam. Deve ir em frente nos seus ideais sem se preocupar com o que os outros pensam de si.',
    22: 'Pontos positivos: universalidade, direção, elevada intuição, pragmatismo, praticidade, habilidade manual, construção, capacidade de organização, realismo, resolução de problemas e empreendedor.\n'
        '\nPontos negativos: esquemas de enriquecimento rápido, nervosismo, autoritarismo, materialismo, egoísmo, falta de visão, autopromoção, arrogância e preguiça.\n'
        'As pessoas cujos nomes correspondem a vinte e dois – o número supremo – possuem todas as qualidades boas dos outros números. Se desejar algo que em princípio pareça impossível, vá em frente, pois com certeza conseguirá atingir o seu objetivo, e até com certa facilidade, pois é habilidoso, idealista, organizado e tem grande potencial de realização. Tem, ainda, outras qualidades: integridade moral, honestidade nata, inspiração divina, disciplina técnica e social, constância nos objetivos e grande capacidade para levar até ao fim um projeto ou um objetivo. Não é muito fácil conviver com ele, pois sendo “superior”, normalmente não se adapta a situações preconcebidas e muito menos gosta de aceitar ordens de quem quer que seja, gostando de trabalhar sozinho ou então em cargos de chefia. A ambição pode levá-lo ao mundo dos negócios, no qual pode se sobressair graças à sua capacidade de organização e administração. Isto sugere também que pode vir a ser bem-sucedido nas áreas industrial, bancária, imobiliária e de merchandising. A sua visão e sua imaginação ativa podem levá-lo a seguir carreiras no teatro, fotografia, cinema ou decoração de interiores. A sua habilidade natural para lidar com as pessoas pode ser canalizada para profissões que envolvam o trato com o público, como comunicação, educação, medicina, trabalho social ou advocacia. O seu discernimento especial e compaixão podem levá-lo a trabalhar no mundo da cura, seja na medicina tradicional ou na alternativa.'
        '\nORIENTAÇÃO: Pode-se tornar famoso, tanto nacional como internacionalmente, pois como tem grande capacidade de persuasão e destreza com as palavras, consegue convencer a todos com seus argumentos e ideias, que outras não devem ser se não as que visem o bem da humanidade. Todo e qualquer vício lhe é altamente prejudicial, principalmente o fumo.'

}
Destino = {
    0: 'Este é um dos números mais importantes do Mapa Numerologico. Ele descreve as influências na personalidade, oportunidades e os obstáculos que uma pessoa irá encontrar ao longo da sua vida. Indica, ainda, as alternativas disponíveis e o provável resultado de cada uma delas.',
    1: 'PONTOS POSITIVOS: Pioneirismo, iniciativa, grande capacidade de liderança, inventividade, capacidade de comando e espírito executivo. Devido a estas características, o consulente com esta vibração, normalmente tem êxito e prosperidade em quase todas as atividades. Normalmente é um ser talentoso e espirituoso, em geral uma pessoa ativa e produtiva, dotada de uma inteligência aguçada e perspicaz se for estimulado por um projeto que considere valer a pena, tem o talento, a vitalidade e a capacidade para realizações extraordinárias. O 1 necessita de brilho, de reconhecimento e sucesso. Para conseguir esse estágio, se dedicam exaustivamente ao trabalho e normalmente são os melhores na sua área de atuação. Quem com ele convive, fica muito intrigado, pois demonstra sempre um semblante feliz, alegre, descontraído e, acima de tudo, de muita confiança em suas realizações, mesmo que tudo isso não seja verdade.\n'
       '\nPONTOS NEGATIVOS: Tendência à arrogância, egoísmo, prepotência, excesso de ânimo e dominação, por vezes achando que o sucesso do grupo dependeu exclusivamente dele. Esse tipo de postura pode tornar difícil o convívio com os amigos e colegas de trabalho, além de dificultar o relacionamento a dois, pois não é nada fácil conviver com esse tipo de pessoa. Um dos obstáculos à sua realização pessoal é a tendência a sobrecarregar o seu sistema nervoso com uma atitude crítica e preocupada. Como isso pode isolá-lo dos outros, precisa ter fé em si mesmo e nas suas habilidades e aproveitar ao máximo suas ideias originais e inventivas.\n'
       '\nORIENTAÇÃO:  Como não gosta de receber ordens, deve sempre procurar trabalhar por conta própria ou em cargo de chefia. Sendo individualista, íntegro e honesto em seus propósitos, deve aprender a ser pioneiro, a trabalhar em grupo e a ser comandante, não ditador; a ouvir conselhos, analisar situações e, após meditar sobre o assunto, tomar decisões. Ou seja, não agir precipitadamente.',
    2: 'PONTOS POSITIVOS: É o número das associações com outras pessoas; aquele capaz de unir, juntar ideias e fatos e contribuir para a realização de qualquer projeto. Possui grande capacidade para assimilar ideias alheias, valorizá-las e condensá-las, de forma a criar um clima de satisfação em todos os envolvidos. Tem habilidades diplomáticas naturais e sabe deixar os outros à vontade. Normalmente as pessoas se sentem atraídas por sua gentileza e maneiras refinadas e normalmente se ajusta bem a trabalhos em conjunto. É amigo leal, excelente pai ou mãe e faz grande questão de proteger os familiares, amigos ou simplesmente colegas de trabalho. Quando consegue ver as pessoas e situações a partir de uma perspectiva superior, pode dar muito de si, ao mesmo tempo em que se mantém distante, confiando que a vida irá naturalmente resolver as coisas no tempo devido.\n'
       '\nPONTOS NEGATIVOS: Tem tendência ao acomodamento e a esperar que os outros tomem a iniciativa, procurando manter-se na retaguarda, aparecendo pouco. Como não é pioneiro ou individualista quanto aos seus projetos de vida, é muito comum perder excelentes oportunidades por causa disso, não notando que o tempo vai passando e os projetos e ideais vão ficando cada vez mais distantes. No lado psicológico, quando está com disposição   negativa (o que não é raro), corre o risco de se sacrificar ou de se deixar levar pelo auto piedade ou comodismo.\n'
       '\nORIENTAÇÃO: Precisa cultivar a paciência, cooperação, tato, lealdade, aplicação e capacidade de seguir a orientação dos outros. Sendo o seu principal atributo a união, deve aprender a unir todos aqueles que desejam essa parceria. Porém, não deve interferir (o que normalmente faz) naqueles que não desejam estar unidos. Trabalhe sempre em parceria; se não existir, procure. Não fique parado esperando que as coisas caiam do céu ou que apareça alguém do nada para lhe ajudar a resolver seus problemas.',
    3: 'PONTOS POSITIVOS: É o Destino da sociabilidade. O 3 é alegre, versátil e talentoso. É criativo, social e possui o dom artístico, seja falando, representando, escrevendo, pintando ou desenvolvendo qualquer outro tipo de arte. Ao longo da sua existência deve sempre procurar cultivar esses dotes, ou seja, a criatividade, os contatos sociais e a expressão de suas ideias e sentimentos. Tudo o que faz cresce e se multiplica. Tem, na comunicação, como um todo, o seu principal trunfo para ser próspero e feliz. A emoção e a criatividade são as razões do seu sucesso, pois tem extraordinária facilidade de expressão e também grande capacidade executiva. Dinâmico e versátil, tem uma personalidade charmosa e habilidade para promover suas ideias. Quanto mais se dedicar ao trabalho, melhores serão os resultados e recompensas. Pode ser bem-sucedido como vendedor, professor, relações públicas, nas artes plásticas ou cênicas ou em contato com o público.\n'
       '\nPONTOS NEGATIVOS: Entedia-se facilmente, é vaidoso em demasia, impaciente, extravagante, exagerado, esbanjador, com certa irresponsabilidade para cumprir horários e desprezar pessoas com nível inferior ao seu. Quando essas características se manifestam, torna-se completamente insociável, disperso e com grande tendência a se entediar, a se isolar e ter crises nervosas.\n'
       '\nORIENTAÇÃO: Como nasceu com o dom da palavra, da comunicação, deve usar esse dom intuitivo para alcançar sucesso em trabalhos mentais. Embora o número três faça com que seja artístico e charmoso, dotado de um bom senso de humor, precisa desenvolver a autoestima e evitar as tendências a se preocupar ou se sentir emocionalmente inseguro. Deve, ainda, manter sob controle suas necessidades de movimentação e os excessos de atividades corriqueiras. Seja mais seleto nas amizades, ou seja, procure qualidade à quantidade.',
    4: 'PONTOS POSITIVOS: É o Destino da conquista, por meio do trabalho duro e incansável. Necessita cultivar a paciência, a confiança e a disposição para servir. Sua vida, em geral, não é muito fácil, em termos de retorno e exigências. Terá de trabalhar muito, com muita dedicação para manter a vida financeira, afetiva e social. Porém, tudo o que consegue   é sólido e eterno. Dedica-se de corpo e alma aos seus projetos    e conquistas e quase sempre é bem-sucedido profissionalmente. Tem poucos amigos, porém sinceros e não se interessa por quantidade, preferindo a qualidade em tudo e em todos. Sensível, cheio de ideias criativas e com senso de visão, tem potencial para expressar os seus conceitos originais e inspirados. Isso pode ajudá-lo a superar possíveis preocupações quanto à sua instabilidade financeira e a tomar as decisões corretas. Em si repousa a segurança e a estabilidade para    manter tudo e todos que dele dependem.\n'
       '\nPONTOS NEGATIVOS: Como é muito estruturado, conservador e inflexível, viver em grupo é extremamente difícil e quase sempre rejeita o novo, pois essa incerteza lhe transmite insegurança e possibilidades de perdas. É perfeccionista ao extremo e cobra demais dos outros, sejam familiares, amigos ou subordinados, tornando-se por vezes chato e incompreendido. Se levar uma vida negativa, poderá se inclinar para assuntos fraudulentos, conspirações, orgulho demasiado ou a planos astutos e maléficos para enriquecer rapidamente às custas do sofrimento alheio.\n'
       '\nORIENTAÇÃO: Deve viver ponderadamente, desenvolver o senso de responsabilidade moral, equilíbrio no poder e correta aplicação do saber em todos os seus projetos e realizações. Para viver sem muitos altos e baixos, precisa praticar a justiça, a ponderação e serviço a terceiros, sempre que a oportunidade se oferecer. Deve se relacionar positivamente com todos e respeitar o direito que eles têm de errarem. Deve, também, se relacionar com pessoas inteligentes e com as quais possa compartilhar os seus interesses.',
    5: 'PONTOS POSITIVOS: É o viajante, sempre desejoso de mudanças e de novidades. Esta vibração pertence às pessoas versáteis, no sentido de mudar, alterar e transformar. Está sempre à procura do que é novo, moderno, diferente. Esta procura também se estende pelo âmbito do saber, aprender e conhecer, muito mais do que por capricho ou prazer. É alegre, comunicativo, inteligente e receptivo, sempre bem-vindo em qualquer roda. Curioso, dotado de capacidade de discernimento, geralmente procura se fortalecer através do conhecimento. Com ele não existe “mau tempo”; está sempre de bem com a vida e, por isso, por vezes é confundido e até tachado de irresponsável, coisa que não é, em absoluto. Como é mentalmente ativo, gosta de um bom desafio intelectual. Contudo, a sua inclinação para se envolver em disputas implica que quando não consegue o que quer pode ficar irascível.\n'
       '\nPONTOS NEGATIVOS: Normalmente compra mais coisas do que necessita e pode pagar, levando-o a ter alguns problemas financeiros ao longo da vida. No lado sentimental, busca relações novas e intrigantes, que são as que o atraem, e às vezes se envolve em situações afetivas de que dificilmente se desvencilha. É um ser extrovertido com ampla visão de tudo e de todos, mas inconstante em seus objetivos, levando-o a ter muitos começos e poucos fins.\n'
       '\nORIENTAÇÃO: Para evitar os aspectos negativos prejudiciais ao desenvolvimento pessoal, deve desenvolver o poder de análise e seleção, aproveitando as experiências para aperfeiçoar e elevar seu caráter e sentir uma nobre afinidade humana isenta de interesses particulares, sexuais ou materiais. Como vive quase que exclusivamente no presente, é aconselhável trabalhar com um propósito definido, com um objetivo altruísta e levá-lo até o fim, aconteça o que acontecer. Lembre-se que a desordem não faz parte do Universo; logo, quando encontrar desordem, mude; mude a desordem em uma nova ordem. Será um bom psicólogo, historiador, professor, jurista, conferencista ou carreiras que tenham contato com o público. Como admira pessoas poderosas e independentes, precisa de um parceiro que o enfrente e que não se intimide com a sua personalidade vigorosa. Embora seja amigável e sociável, gosta de fazer as coisas sozinho e de enfrentar novos desafios. Como gosta de exercitar o seu poder cerebral, prefere a companhia de pessoas inteligentes.',
    6: 'PONTOS POSITIVOS: O 6 é normalmente dócil, educado, não suporta brigas e desentendimentos e quer sempre manter tudo calmo e tranquilo. É o número da harmonia, do lar e do amor. Compaixão, idealismo e natureza atenciosa são algumas das qualidades sugeridas por este dia admirável. É o número do perfeccionista ou do amigo universal e em muitos casos, humanitário, responsável e amoroso. Adora crianças, plantas e animais. A sua forte personalidade, combinada com sua capacidade amorosa e seu charme, pode ser muito atraente para os demais. É ambicioso, mas também compreensivo e encorajador, com disposição de fazer qualquer coisa pelas pessoas que ama. Detesta aglomerações e festividades excessivas, gosta da harmonia e de viver em paz. Geralmente se dá bem em relacionamentos afetivos, pois é amável e muito gentil.\n'
       '\nPONTOS NEGATIVOS: É por demais sentimental, vulnerável e crédulo, sempre confiando nas aparências sem analisar os fatos friamente. Tem, normalmente, dois problemas: o acomodamento e a falta de interesse. Por ser profundo, necessita de tranquilidade para, no mínimo, aparentar estar bem consigo mesmo. Quanto ao interesse, necessita ser constantemente incentivado, pois caso contrário tem tendência a deixar tudo a meio caminho. Na tentativa de manter tudo calmo e tranquilo, por vezes acaba interferindo de maneira não condizente com o momento, pois nem sempre a docilidade é o melhor remédio para certas situações, e isso o perturba sobremaneira, levando-o a se retrair e a acomodar-se, até que outro fato ou circunstância o incentivem.\n'
       '\nORIENTAÇÃO: Como é o Destino da responsabilidade e da liderança, necessita aprender a se ajustar a condições pouco harmoniosas, a assumir o fardo dos outros, a defender suas próprias ideias e a servir com alegria. Apesar de seus ideais elevados, precisa ser determinado, ter paciência e perseverança para atingir as suas metas. Deve trabalhar com o público, pois conquista amizades e simpatias com grande facilidade, sendo respeitado e amado por todos.',
    7: 'PONTOS POSITIVOS: Normalmente é uma pessoa que busca na vida o entendimento, o conhecimento mais profundo e qualitativo. No desenrolar de sua vida sempre procurará ser eficiente, intuitivo, perfeccionista, meticuloso, racional, reflexivo, concentrado, reservado, compreensivo e amoroso, apesar de não demonstrar qualquer desses sentimentos altruístas. Busca sempre o lado oculto das coisas e das pessoas. Nunca se satisfaz com as aparências, preferindo conhecer a essência. Com uma necessidade constante de ter uma maior autoconsciência, gosta de reunir informações e se interessa por ler, escrever por assuntos espirituais. A tendência a ser enigmático ou reservado sugere que pode ocasionalmente se sentir mal interpretado.\n'
       '\nPONTOS NEGATIVOS: Por vezes tem espírito dominador, crítico, exigente, perfeccionista, angustiado, impaciente, orgulhoso e aparentemente descrente. Embora seja pragmático por natureza, pode, às vezes, ter ideias pouco convencionais; tome cuidado, porém, para que isso não vire uma desculpa para ser do contra somente para se fazer de difícil. Deve, também, aprender a valorizar as pessoas como elas são, ou seja, respeitar o fato de os outros serem fruto daquilo que conseguiram absorver da vida e do convívio com os demais.\n'
       '\nORIENTAÇÃO: O 7 é o Destino da sabedoria e também o da solidão, por isso necessita desenvolver seus poderes mentais, estudar, meditar, buscar o significado último da vida. Enfim, tornar-se um especialista. Com sua inteligência arguta, sua habilidade de liderança e poder inerente a este dia, pode ter muitas oportunidades profissionais e ser bem-sucedido em qualquer uma delas. Pode ser escritor, advogado, pesquisador, religioso, pesquisador ou trabalhar simplesmente par desenvolver espiritualmente o planeta. Contudo, evite usar o poder de forma destrutiva e se rebelar contra os outros, pois o “tiro pode sair pela culatra”, ou seja, você ser o prejudicado no final.',
    8: 'PONTOS POSITIVOS: É aquele que tem capacidade para planejar, organizar e levar a efeito qualquer tipo de negócio. O seu Destino é o ganho financeiro e material. Individualista, inteligente e sensível, tem uma personalidade exigente. Embora seja ambicioso e vigoroso, e tenha coragem e lealdade aos seus ideais, é também muito gentil e generoso com os que ama. Geralmente não apresenta problemas financeiros, pois caso não seja herdeiro, é empreendedor. Objetividade é o seu lema. É justo, disciplinado, honesto, amigo, econômico, elegante e bom negociante. A sua necessidade de crescimento e habilidade para rever o todo sugere que gosta de pensar em grandes termos. Confiante e curioso, não gosta de interferências de outras pessoas e, com a inclinação à teimosia, pode tornar-se inquieto e impaciente.\n'
       '\nPONTOS NEGATIVOS: Como está sempre sendo impulsionado para ganhar dinheiro e para conquistas materiais, por vezes se esquece da família, do lado sentimental, social e, principalmente, de si próprio. Apresenta, também, acentuada dificuldade em demonstrar afeto e carinho, aparentando frieza nos relacionamentos em geral. Na realidade, é um ser altamente emotivo, necessitando de ajuda e compreensão por parte das pessoas mais chegadas.\n'
       '\nORIENTAÇÃO: Precisa cultivar a eficiência na arte de negociar e compreender as leis que comandam a acumulação, o poder e o uso do dinheiro. Para aproveitar ao máximo a sua inteligência rápida, sua necessidade de expressão e o lado sociável da sua natureza, precisa superar a inclinação à incerteza ou à indecisão nos seus relacionamentos emocionais. Necessita aprender que as pessoas precisam de atenção. Por isso, fique atento: quanto maior for   a ostentação, mais as pessoas concluirão que você não precisa delas. E, se em algum momento você realmente precisar, elas não estarão por perto. Como tem grande capacidade executiva, pode ser advogado, executivo de grandes empresas, vendedor (por conta própria), professor, político, missões religiosas ou conselheiro.',
    9: 'PONTOS POSITIVOS: É o Destino da universalidade. A pessoa com este Destino normalmente é espirituosa, charmosa, divertida e amorosa, mas o seu desejo de ser bem-sucedido deriva de uma interessante mistura de materialismo e idealismo. Benevolência, ponderação e sensibilidade emocional estão associadas a este dia. Tolerante e gentil, é na maior parte das vezes generoso e liberal. Sua intuição e seu poder psíquico indicam uma receptividade universal que, se positivamente canalizada, pode inspirá-lo a buscar o caminho espiritual. Necessita que as pessoas ao seu redor estejam bem, pois caso contrário tudo fará para que isso aconteça. Sua necessidade de auxílio é voltada para o todo, para a humanidade e não para uma pessoa ou um caso especial. Preocupa-se com tudo que esteja em desacordo com suas inspirações, é muito emotivo e por demais sentimental.\n'
       '\nPONTOS NEGATIVOS: Como vive para o mundo, por vezes acaba perdendo excelentes oportunidades, principalmente no âmbito profissional. Embora seja idealista e generoso, pode ser reservado e, em algumas ocasiões, ocultar sentimentos profundos que podem se transformar em ressentimentos. Frustra-se com certa facilidade, é pouco prático, de certa maneira egoísta, tem dias amargos, complexo de inferioridade, medos, preocupação exagerada e tendência ao isolamento.\n'
       '\nORIENTAÇÃO: Necessita servir e se divertir. Precisa aprender a amar o próximo, a livrar-se dos preconceitos, ser generoso, tolerante e compreensivo com seus semelhantes. Este número sugere a necessidade de superar desafios e a tendência a ser excessivamente sensível e passível de instabilidade emocional. Pode-se beneficiar de viagens internacionais e da interação com pessoas de diversos círculos sociais, mas deve evitar os sonhos irrealistas ou a inclinação ao escapismo.Com sua grande capacidade para adquirir conhecimentos, pode ser bem sucedido no meio acadêmico ou nas áreas de advocacia, psicologia, medicina, literatura, ou ainda nas artes plásticas de um modo geral e, principalmente, no meio esotérico.',
    11: 'PONTOS POSITIVOS: É o Destino da inspiração. É o número das pessoas altamente sensíveis e intuitivas. O 11 possui facilidade de ver, não somente a realidade atual, mas também os seus desdobramentos. Enxergam longe.   Normalmente são impacientes e dotados de certo nervosismo Sensível e idealista, é normalmente uma pessoa enérgica, com enorme potencial mental quando está disposto a se disciplinar através da concentração nos seus objetivos. Versátil e imaginativo, precisa expressar o seu poder criativo e emocional. A compostura e a perseverança são frequentemente a chave para o seu sucesso, e pode impressionar quando se especializa em uma área específica. O número-mestre 11 sugere que o idealismo, a inspiração e a inovação são altamente importante para ele. Uma combinação de humildade e confiança o desafia a trabalhar para alcançar a independência financeira e espiritual. É concentrado, objetivo, inspirador, muito espiritualizado, intuitivo, expansivo, prestativo, tem grande habilidade para curar, é humanitário e psíquico.\n'
        '\nPONTOS NEGATIVOS: Tendência a vícios, falta de objetivos, complexo de superioridade, magoa-se com facilidade, hipersensível, excessivamente emotivo, de certa forma egoísta e, quando as coisas não correm como quer, pode-se tornar desonesto no afã de conseguir o que quer. Como é altamente sensível, capta com maior intensidade as imperfeições do mundo e das pessoas. Com isso, vai acumulando insatisfações e sofrimentos. Essas insatisfações e sofrimentos, devem servir para melhorar sempre os seus conceitos sobre tudo e todos.\n'
        '\nORIENTAÇÃO: Precisa investigar o misticismo, confiar na própria intuição, ter fé, conservar-se humilde diante da notoriedade e inspirar os outros através do seu exemplo. Como a Natureza lhe dotou de certos atributos que não deu aos demais, deve usá-los para melhorar o planeta, os seus semelhantes e a própria vida. Não exija demais dos outros ou que eles compreendam ou façam o mesmo que você. Embora precise de liberdade para agir de forma independente, evite   se preocupar demais consigo mesmo e aprenda a trabalhar com as outras pessoas. Quando combina seu pensamento imaginativo com sua habilidade e praticidade, pode obter resultados extraordinários.',
    22: 'PONTOS POSITIVOS: É o Destino do mestre construtor que se dedica altruisticamente à humanidade. É, também, o número da Sabedoria. As pessoas com esta vibração são normalmente precoces em todas as atividades, sejam elas de cunho material, social, mental ou espiritual. Conhecem da vida o superficial, o profundo e o invisível, e possuem uma capacidade nata para o ensino. O Destino mestre 22 é orgulhoso, prático e altamente intuitivo. Normalmente honesto e trabalhador, com capacidade natural para a liderança, possui uma personalidade carismática e uma compreensão profunda das pessoas e de suas motivações. Apesar de ser reservado, muitas vezes tem uma preocupação atenciosa e protetora com o bem-estar das pessoas, mas nunca perde de vista a sua posição realista e pragmática. Em geral culto e mundano, tem muitos amigos e admiradores. Os que são mais competitivos, podem alcançar o sucesso e a fortuna com a ajuda e o encorajamento dos amigos e familiares. Charmoso, talentoso, comunicativo e diplomata natural, pode seguir carreiras nas áreas de vendas, promoções, agências ou de relações públicas. Da mesma forma, pode se dedicar ao mundo das comunicações, área editorial, política, educação, medicina e, principalmente, à área holística, ajudando a humanidade. Além de todos estes dons, ainda possui facilidade de entender e desenvolver processos de alta magia, entendida  como uma força extrafísica e não como uma forma racional.\n'
        '\nPONTOS  NEGATIVOS:  Tendência  à  arrogância, ao nervosismo, à preguiça, ao egoísmo, à autopromoção, à traição e a menosprezar o semelhante.\n'
        '\nORIENTAÇÃO: Precisa aprender que justiça, cooperação e serviço prestado também fazem parte do cotidiano. Precisa aprendera penetrar nos problemas, até às suas últimas motivações de seus atos e as verdadeiras razões de seu proceder. Apesar de emocional e materialmente generoso, pode, às vezes, radicalizar, deixando-se levar por uma vida de luxo e glamour. Se cair no escapismo e se recusar a amadurecer, precisará aprender a ser responsável e a ter uma perspectiva mais madura. Porém, quando faz algo de que realmente gosta, é capaz de trabalhar arduamente e, com dedicação e perseverança, pode se realizar e ser bem-sucedido. Está destinado a ser Líder, pela sua capacidade executiva que se estende por vasto campo, atingindo mesmo as fronteiras internacionais. No tocante à Magia, a mesma poderá ser construtiva ou destrutiva; o que semear, colherá!'
}
Liçõescarmicas = {
    0: 'Diferentemente das Dívidas Cármicas que são transgressões às Leis Naturais, as Lições Cármicas mostram obrigações que nos foram impostas na vida passada, mas que foram negligenciadas. Por isso, nesta existência, devemos “aprender” essas lições (eliminando o Carma), pois incidindo nos mesmos erros, com certeza teremos de passar pelos mesmos problemas nesta ou em outras encarnações.',
    1: 'LIÇÃO CÁRMICA 1 Esta Lição  mostra  claramente que em vidas passadas você foi um ser preguiçoso, sem qualquer iniciativa, ambição ou originalidade, e muito menos competência para cuidar de seus próprios negócios. Em outras palavras, você não estava “nem aí” para o que ocorria à sua volta. Provavelmente foi rico, gastador, esbanjador e sem qualquer criatividade ou objetivo concreto na vida, além de carecer de confiança em sua própria personalidade. Em virtude dessas cruciais faltas, nesta existência precisa aprender a ser pioneiro, criativo, a tomar decisões e não esperar por ninguém, além de aprender a ter amor próprio, a confiar na própria capacidade, a decidir e a concluir seus projetos e seus objetivos.',
    2: 'LIÇÃO CÁRMICA 2 Em vidas passadas você foi um ser totalmente indisciplinado, negligente e nada colaborador. O dinheiro tinha pouca importância, gastando-o a torto e a direito, sem se importar com o dia de amanhã. Nesta existência, em virtude dessas faltas no “passado”, poderá se tornar subserviente, introspectivo, com medo de tudo e de todos, evitando contatos com amigos e colegas. A Lição a aprender nesta existência é trabalhar em parcerias, grupos; deve, também, adquirir paciência e levar a bom termo os seus projetos e ambições. Deve ser econômico, poupador, pois caso contrário, terá graves problemas financeiros na velhice.',
    3: 'LIÇÃO CÁRMICA 3 Provavelmente “pobre” e sem muito estudo na vida passada, você foi um ser sem qualquer ambição, expressão, imaginação, criatividade, além da falta de vontade de progredir, limitando-se à “vidinha” insignificante, não tendo qualquer talento e confiança em si próprio, mostrando-se, isto sim, nervoso, irritadiço e até violento. A Lição a aprender nesta existência é eliminar a timidez, seja ela no falar, representar, escrever ou mesmo tratando com estranhos. Deve cultivar também a imaginação, agir com segurança, criatividade, paciência e sempre positivamente.',
    4: 'LIÇÃO CÁRMICA 4 Provavelmente “preguiçoso” na vida passada deixava tudo a meio caminho; não gostava de receber ordens, de trabalho metódico ou pesado. Tinha também certa dificuldade em terminar o que começava, ou seja, era dispersivo, imprudente e desajeitado. A Lição a aprender nesta existência é trabalhar metodicamente, ser constante (levar até o fim os seus objetivos), a ser concentrado, paciente e compenetrado no trabalho. Deve ainda aceitar qualquer tipo de trabalho, mesmo que em princípio não goste dele; ou seja, se não puder fazer o que gosta, deve gostar do que faz. Será bem sucedido na vida se seguir rigorosamente esta Lição.',
    5: 'LIÇÃO CÁRMICA 5 Na vida passada as mudanças e as novidades o assustavam, preferindo agir com cautela exagerada e em situações já consagradas pelo uso. Agora, em virtude dessas fraquezas, poderá se tornar incompreendido pelos que ama e convive, além de ter de enfrentar certas mudanças não desejadas, como mudança de casa, de trabalho, de parceiro, etc. A principal Lição a aprender nesta existência é enfrentar com naturalidade essas mudanças, a não ter medo de novidades e estar sempre por dentro da moda, da política, da moderna tecnologia, enfim, por dentro de tudo.',
    6: 'LIÇÃO CÁRMICA 6 Na vida passada você foi em ser completamente alheio e, digamos, até irresponsável para com obrigações familiares. Não estava nem aí para ajudar familiares doentes e necessitados. Nesta atual existência, até que aprenda esta Lição, que    é a de ser bom pai (ou mãe), irmão, parente, amigo ou cônjuge, poderá ser forçado a assumir responsabilidades domésticas (cuidar de doentes) ou mesmo prestar serviços a terceiros. Deverá ser compreensivo nessas ocasiões, pois os “parentes” e amigos pouco ou nada farão para ajudá-lo; muito pelo contrário, deverão sobrecarregá-lo com muitas exigências, até que aprenda definitivamente esta crucial Lição.',
    7: 'LIÇÃO CÁRMICA 7 Em outra existência você foi um ser indeciso; buscava o lado espiritual, mas queria ser rico a qualquer custo. Também pouco ou nada fez em matéria educacional, sendo um ser de pouca cultura, sem conhecimento de si próprio e das Leis que regem o Universo. Até que aprenda a Lição que outra não é senão a de ser culto, espiritualizado, analítico e principalmente decidido na tomada de atitudes, a vida o tornará indeciso, sempre ponderando e adiando indefinidamente os projetos. Deve, também, treinar a sua mente em vista do futuro que rapidamente se aproxima, ou seja, estudar modernas técnicas e agir sem medo.',
    8: 'LIÇÃO CÁRMICA 8 Provavelmente você foi “rico” na vida passada (herdeiro), sem qualquer capacidade ou vontade de ganhar dinheiro. Na verdade, era gastador, esbanjador, pouco se importando quem ganhava o dinheiro; só queria gastá-lo! Agora, nesta existência, até que aprenda esta Lição, a vida o levará a ter problemas financeiros, negócios fracassados e também longos períodos de paralisação. Deve dar valor ao dinheiro, gastá-lo com critério, sabedoria, sem esbanjá-lo, poupando-o, pois provavelmente terá de tratar de negócios e finanças próprias e também de outras pessoas, além de ter sérios problemas na velhice.',
    9: 'LIÇÃO CÁRMICA 9 Você foi na vida passada um ser frio, sem amor e compreensão para com seus semelhantes. Não aparentava qualquer sentimento de afeto; era desumano, indiferente, pouco se importando com os outros. Em virtude dessa “apatia” emocional, nesta existência poderá sofrer alguns reveses na vida, como perda de propriedades, longos períodos desempregado, doenças diversas e também desapontamentos com amigos, parentes, sócios, cônjuge ou simplesmente companheiros. A Lição a aprender nesta vida é a de ser bondoso, amoroso, generoso, compreensivo e interessado em tudo e com todos'

}
Tendênciasocultas = {
    0: 'Descrevem os desejos a que uma pessoa foi exposta em vidas passadas. Caso não estejamos atentos e não nos policiarmos constantemente, é comum tornarmos a ter estes mesmos desejos de conduta nesta existência.',
    1: 'TENDÊNCIA OCULTA 1 É o desejo da individualidade. Há uma tendência a ser autoritário, dominador, arrogante e um tanto egoísta.',
    2: 'TENDÊNCIA OCULTA 2 É o desejo de associações. Há uma tendência para depender demasiado dos outros, principalmente da família e dos amigos, tanto monetariamente como emocionalmente.',
    3: 'TENDÊNCIA OCULTA 3 É o desejo do auto expressão. Há uma tendência à vaidade, à impaciência, à presunção, a dispersar energias, viver sem objetivos concretos e buscando sempre diversões e festas, sem se preocupar muito com o dia de amanhã.',
    4: 'TENDÊNCIA OCULTA 4 É o desejo de trabalho. Há uma tendência para asfixiar sua personalidade com quantidade excessiva de detalhes (perfeccionista), de ser teimoso, intolerante e obstinado.',
    5: 'TENDÊNCIA OCULTA 5 É o desejo de mudança e também de liberdade pessoal. Há uma tendência para viver à custa dos outros, abusar do sexo, de algumas drogas (cigarro e álcool), para mudanças sem motivo aparente e para ser precipitado e impulsivo.',
    6: 'TENDÊNCIA OCULTA 6 É o desejo de realização e responsabilidade. Há uma tendência para se preocupar excessivamente com a família, principalmente com os filhos (se os tiver), com os deveres profissionais, para ser teimoso, perfeccionista e para apegar-se a modelos inflexíveis e obstinados.',
    7: 'TENDÊNCIA OCULTA 7 É o desejo de sabedoria e conhecimento. Há uma tendência para o fingimento, para a intriga e até para o alcoolismo, pois normalmente sente-se incompreendido e até rejeitado pelos parentes, amigos, cônjuge ou simplesmente conhecidos.',
    8: 'TENDÊNCIA OCULTA 8 É o desejo do materialismo. Há uma tendência para se preocupar excessivamente em como ganhar dinheiro, obter bens materiais e também poder.',
    9: 'TENDÊNCIA OCULTA 9 É o desejo de conhecimento e amor universal. Há uma tendência para se preocupar excessivamente com os problemas mundiais em detrimento de si próprio e da família.'
}
RespostaSubconsciente = {
    0: 'Este número nos diz como será a reação instintiva e automática de uma pessoa, quando em uma situação de emergência.',
    2: 'RESPOSTA SUBCONSCIENTE 2 É um ser arrogante, mentiroso, que não respeita regras, querendo que tudo e todos girem em torno de si (egocêntrico).',
    3: 'RESPOSTA SUBCONSCIENTE 3 É dispersivo e até indisciplinado; normalmente reage de forma explosiva e até de certa maneira destrutiva.',
    4: 'RESPOSTA SUBCONSCIENTE 4 Normalmente é um ser que vive perdido num labirinto de detalhes. As suas reações são fracas, tem tendência a vacilar e até atrapalhar os outros.',
    5: 'RESPOSTA SUBCONSCIENTE 5 É uma pessoa tensa e nervosa. Numa crise tem tendência a agir de forma confusa e impulsiva.',
    6: 'RESPOSTA SUBCONSCIENTE 6 É um ser sentimental. Sua primeira preocupação numa crise é com os entes queridos, seus objetos de estimação e animais (se os tiver).',
    7: 'RESPOSTA SUBCONSCIENTE 7 Normalmente é arredio e não gosta de se envolver com problemas alheios. Em uma emergência, considerará analiticamente a situação e, então, normalmente se retirará para dentro de si mesmo e rezará. Se existirem perdas materiais ou de vidas, poderá se entregar a vícios, como a bebida ou drogas.',
    8: 'RESPOSTA SUBCONSCIENTE 8 É eficiente e organizado. Numa crise ou em qualquer ocasião, pode-se contar com ele, pois é seguro e digno de confiança.',
    9: 'RESPOSTA SUBCONSCIENTE 9 É um ser entediado. A maioria das coisas tem pouca importância para ele. Numa crise é melhor não contar com ele, pois é impessoal, filósofo, introspectivo, resignado e indeciso.'
}
DívidasCármicas = {
    0: 'Quando o ser humano transgride certas Leis Naturais, está sujeito à penalidade que a Natureza lhe aplicará, seja nesta ou em outra existência. São transgressões, normalmente conscientes contra a vida de outras pessoas, ou contra a própria vida, que ferem o preceito da Lei da causa e do efeito.',
    13: 'DÍVIDA CÁRMICA 13 = Esta Dívida representa a morte em todas as suas concepções. Com absoluta certeza, em vidas passadas cometeu transgressões às Leis Naturais, tais como: (A) foi suicida, ou (B) foi assassino(a), ou (C) foi passivo(a) e negligente com a própria vida, ou (D) com a vida alheia, ou e por vícios apressou a própria morte. Uma coisa é certa: qualquer que tenha sido o caso ou os casos, os mesmos foram cometidos conscientemente, não sofrendo, em absoluto, qualquer influência externa. Nesta existência, em virtude dessas transgressões, pode se tornar preguiçoso, indiferente ao que acontece à sua volta    e tender para a crítica, seja ela de cunho pessoal, político ou existencial. Elimina-se esta crucial Dívida Cármica trabalhando duro em qualquer atividade profissional. Durante sua existência muito lhe será exigido e terá de lutar contra as adversidades da vida, administrar bens (seus e alheios), as más atitudes (suas e alheias) e contra o descontentamento e, desta maneira, além de “pagar a Dívida”, terá menos problemas financeiros ao longo da vida.',
    14: 'DÍVIDA CÁRMICA 14 = É a Dívida que trata dos bens materiais. O ser humano com esta Dívida Cármica terá ao longo da sua existência ou até que aprenda a “lição”, inúmeras atribulações: troca constante de emprego, de residência, perda de bens: móveis ou imóveis e negócios fracassados. Em vidas passadas, locupletou-se financeiramente em detrimento de outras pessoas, ou seja, para subir na vida, agiu fraudulentamente, prejudicando sócios, amigos e companheiros de jornada, arruinando-os e levando-os à miséria, à desgraça e à destruição moral. Nesta existência, terá de arcar com o peso de todos esses crimes hediondos, causados pela ambição desmedida e total inconsequência dos fatos. Tais crimes, até serem completamente “pagos”, lhe trarão sérios aborrecimentos nesta existência e certas perturbações, tais como, desapontamentos com amigos, sócios, familiares e também muitos obstáculos que se interporão aos seus objetivos. Tais desacertos na vida passada deve-se à sua total falta de compreensão e ambição desmedida. Ainda agora, ou seja, nesta vida atual, caso não seja espiritualmente desenvolvido, pode cometer os mesmos delitos e, assim, ter de passar pelos mesmos sacrifícios. A principal “lição” a aprender nesta vida é o desapego. Não se apegue a nada nem a ninguém e cultive a compreensão até às últimas consequências. Somente desta maneira, compreendendo o seu semelhante e desapegando-se de tudo e de todos, conseguirá desenvolver a humildade suficiente, o amor e a compaixão para resgatar essa tão triste e cruel Dívida Cármica.',
    16: 'DÍVIDA CÁRMICA 16 = Esta Dívida Cármica é a manifestação do orgulho, do autoritarismo, da traição e da vaidade. Com absoluta certeza em vidas passadas você viveu exclusivamente para seu deleite (bem-estar), pouco se importando com o direito ou sentimento do próximo, subjugando todos e agindo com desmedido autoritarismo, frieza e até violência. Nesta existência, em virtude dessas “faltas”, deverá sofrer vários contratempos ou até algumas calamidades, tais como: acidentes os mais diversos, decepções com amigos, colegas de trabalho, parentes, cônjuge (namorado ou companheiro), fracasso nos negócios e também problemas com a justiça. Este Carma ainda carrega o peso de traições conjugais, crimes passionais e uniões ilegítimas. Além dos contratempos acima citados, poderá passar por momentos muito angustiantes nesta existência, como trabalhar e não conseguir atingir os objetivos, perder a fortuna ou o poder, além de ser traído pelos que ama e confia. Para eliminar este Carma, deve ser compreensivo, justo e solidário com seu semelhante. Deve andar na mais perfeita retidão e não se desviar deste caminho, sob pena de continuar indefinidamente sofrendo dos mesmos males.',
    19: 'DÍVIDA CÁRMICA 19 = É o número do “equilíbrio”. Como sabemos, a lei da Causa e Efeito é perfeita e aqueles que a transgridem com absoluta certeza pagarão caro por essa transgressão. Em vidas passadas você agiu contrário a essa Lei e, agora, deve pagar o preço da iniquidade ou então repor aquilo que tirou aos outros. Mais claramente, você se locupletou financeiramente e socialmente, usando de artifícios fraudulentos, prejudicando pessoas e até fazendo que sofressem por sua causa, ou seja, levando culpas que deveriam ser imputadas a você, como serem aprisionados, terem seus bens confiscados e se verem desonrados por sua causa. Mais ainda, ganhou notoriedade e premiações que não lhe eram absoluto de direito, e sim de outrem. Nesta existência, poderá ter segredos desvendados, ser traído por sócios e companheiros e, principalmente, ter de conviver com o sucesso e fracasso contínuos, até que aprenda a lição que é a que só colhemos aquilo que semearmos. Basicamente, é o aviso para que não seja dominador, orgulhoso, egoísta, autoritário e exerça a tolerância em tudo e com todos. Elimina-se esta Dívida Cármica levando uma vida reta e elevada, altruísta, carregada de amor, afeto, bondade e voltada para o bem-estar de toda a humanidade. Agindo desta maneira, ou seja, plantando bem e cuidando da horta, a colheita será farta de sucesso e prosperidade.',
}
Missão = {
    0: 'A Numerologia Cabalística dá grande importância a este número, sendo considerado mesmo de “alta importância”, pois reflete, na essência, o que a pessoa veio fazer neste planeta, nesta existência. É fundamental e muito importante esclarecer que toda e qualquer pessoa tem “livre arbítrio” e pode fazer o que bem entender com a sua vida, porém, também é importante se saber que os números obedecem a uma ordem rigorosa de harmonia, compatibilidade, neutralidade e incompatibilidade que, se não respeitada, pode causar ao seu portador inúmeros aborrecimentos ou mesmo derrocadas na vida. Lembre-se de que a Numerologia Cabalística existe para facilitar a vida das pessoas e não para complicá-la. Todos nós temos uma “Missão” sobre a Terra, nesta e em outras existências (A Natureza jamais desejou que qualquer ser humano viesse a sofrer; e se tal fato acontece, a culpa é exclusiva do ser humano); logo, este item é de grande utilidade, pois a sua essência mostra como podemos tirar o melhor proveito da vida, sem que com isso prejudiquemos qualquer outra pessoa.',
    1: 'MISSÃO – 1 (LIDERANÇA)\nPontos positivos: liderança, criatividade, caráter progressista, vigor, otimismo, convicções fortes, competitividade, independência e espírito de união. \nPontos negativos: arrogância, ciúme, excesso de orgulho, antagonismo, falta de controle, egoísmo, instabilidade e impaciência. As atribulações, contratempos, obstáculos e vitórias que obteve ou obterá na primeira metade da vida (até os 45 anos), servem ou serviram para desenvolver a força moral necessária na segunda metade, quando necessitará se estabilizar materialmente e cumprir seus deveres sociais e espirituais. Com desejo de ser o primeiro e independente, o possuidor desta Missão tem inclinação a ser individualista, inovador e corajoso, com muita energia, predicados que deve cultivar ao longo da sua vida. O espírito pioneiro o encoraja a se aventurar sozinho nos seus empreendimentos. Esta força também pode estimulá-lo a desenvolver habilidades executivas, de organização, criatividade e originalidade. A Missão 1, é a do ser inteligente, mas também o é do ambicioso e do agressivo. É confiante em seus propósitos, independente e, se quiser, poderá aspirar aos maiores cargos na sua profissão ou atividade. Tem tendência a ser autoritário, dominador e inflexível em seus propósitos, o que lhe causará alguns problemas, principalmente com respeito aos subordinados e pessoas mais chegadas, como parentes e amigos. Será sempre invejado e deve se proteger espiritualmente contra essa inveja e contra inúmeras maledicências que dirão sobre sua pessoa. Será também admirado, respeitado e terá a consideração dos outros. Se agir com prudência e sabedoria, poderá ser muito bem-sucedido na política, na administração de grandes empresas, como líder de povos ou excelente místico. Idealista e sociável busca estabilidade emocional e segurança. Tem uma personalidade carismática e atrai sempre amigos e admiradores respeitáveis.',
    2: 'MISSÃO – 2 (HARMONIA)\nPontos positivos: gentileza, tato, receptividade, intuição, consideração, harmonia, solidez, inteligência, conservadorismo, economia e honestidade. \nPontos negativos: falta de confiança, subserviência, excesso de sensibilidade, mau humor e se ofende com facilidade. Paz e disposição ordenada como um todo, é seu principal lema. A Missão 2 representa a solidificação do ser humano, a capacidade de trabalhar em grupo, a disposição em servir e também em querer que tudo esteja em perfeita ordem e em perfeito estado de funcionamento. Harmonia entre pesquisar e executar, coerência entre os propósitos e os objetivos são, com certeza, suas maiores virtudes. Aparenta ser frio, calculista, obstinado, incrédulo, mas na realidade é muito afetuoso, bom amante e sempre pronto a ajudar aqueles que o procuram e também aqueles que necessitam sem o procurarem, ou seja, gosta de ajudar a todos   indistintamente. Em virtude desse seu procedimento, por vezes é enganado por aqueles em quem confia, pois como é íntegro e honesto em tudo e com todos, acha que os outros também o são. Tal procedimento o deprime e o chateia, mas como não sabe guardar rancor nem ódio, rapidamente se recupera desses inconvenientes, perdoando os elementos que lhe fizeram mal. É o melhor mediador que existe e, em sua longa existência, será excelente juiz, advogado, médico, professor ou trabalhando como elemento de retaguarda em uma grande organização, além, é claro, de ser o melhor diplomata que se conhece. Bebida alcoólica, cigarro e drogas de qualquer espécie, são venenos para o seu organismo.',
    3: 'MISSÃO – 3 (CRIATIVIDADE)\nPontos positivos: amigável, criativo, artístico, amante da liberdade, engraçado, com excelente senso de humor, talento com as palavras e com grande poder de desejar. \nPontos negativos: entendia-se com facilidade, vaidoso, excessivamente imaginativo, exagerado, convencido, extravagante, comodista e preguiçoso. É sociável, popular, criativo, artístico e imaginativo. A palavra que exprime seu número é a auto expressão, seja na oratória, escrevendo ou mesmo representando. De excelente senso de humor, o 3 nasceu para brilhar, pois é dotado de grande imaginação e criatividade. Mesmo quando a idade chegar, parecerá sempre jovial e sonhador, procurando novas aventuras e criando novos projetos. Se interpretar estas palavras corretamente, poderá ser ou conseguir tudo o que desejar. Durante sua longa vida, muitas oportunidades de sucesso surgirão; cabe a você interpretá-las e segui a sua própria intuição, não se deixando seduzir ou levar pela conversa dos outros. Como não se deixa jamais dominar pela depressão ou desânimo, quase sempre é bem-sucedido no que pretende realizar. Porém, de acordo com a sua personalidade extrovertida, que por vezes é interpretada como sendo irresponsável, estará exposto a críticas, do tipo: “Você é irresponsável, só pensa em farra...” Deve evitar os trabalhos rotineiros e enfadonhos, como escriturário, caixa de banco, ou qualquer outro que não seja brilhar. Deve estimular a arte na sua plenitude, a oratória, a comunicação, a música, a literatura, o paisagismo, etc. Mesmo quando muito tenso ou preocupado, jamais demonstra seus sentimentos, mostrando aos outros somente o que é belo e agradável em seu ser, o que na realidade é muito bom, dando vida à antiga máxima: “Demonstre sempre alegria; desse modo, deixarás os amigos sempre alegres e os inimigos perplexos e intrigados”. O convívio familiar, os compromissos sociais e o trabalho, normalmente serão tratados com espírito elevado e qualquer problema que surja, será resolvido com sabedoria e criatividade.',
    4: 'MISSÃO – 4 (VONTADE)\nPontos positivos: organizado, disciplinado, estável, trabalhador, habilidoso, perseverante, sincero, íntegro, paciente, conservador, esperançoso, confiável, pragmático e autodisciplinado. \nPontos negativos: pouco comunicativo, autoritário, muito econômico, rígido e guarda ressentimentos. Sua Missão não é a de um visionário e sim de alguém com os pés no chão, alguém que persegue e consegue seus objetivos por métodos já consagrados e não por iniciativa própria, usando projetos inéditos. É um lutador; coragem e determinação não lhe faltam e enfrenta qualquer obstáculo com valentia e determinação. A sua Missão é construir coisas; porém, a sua realização pessoal será de âmbito mais coletivo e não em prazer próprio. O amor, a tolerância, a paciência, a harmonia, fazem parte desta Missão, pois todos estes predicados lhe serão exigidos durante toda a vida. Na maturidade, ou seja, após os 45 anos, tais virtudes deverão ficar muito mais latentes, pois será comparado a um pai ou a uma mãe, aonde todos virão se aconselhar. Como é empreendedor, ambicioso e trabalhador, tem força de vontade e habilidade para perceber as oportunidades que podem ajudá-lo na sua escalada para o sucesso. A construção civil, a química, a metalurgia, a odontologia, a segurança pública, política, e as finanças de um modo geral, devem ou deveriam ser os seus principais objetivos. As parcerias serão sempre preferíveis à iniciativa individual. Se compreender bem estas palavras e possuir tempo e discernimento compatíveis, conseguirá reconhecimento tanto   no âmbito nacional como internacional, pois é carismático, amigo e interessado no bem-estar do seu semelhante, e o planeta necessita muito desse tipo de indivíduo.',
    5: 'MISSÃO - 5 (VERSATILIDADE)\nPontos positivos: versátil, adaptável, progressista, magnético, ousado, rápido, curioso, místico e sociável. \nPontos negativos: instável, procrastinador, inconsistente e excessivamente confiante. Além dos predicados acima expostos, esta Missão mostra um ser ousado, enérgico e amante da liberdade. Normalmente é atraente ao sexo oposto e está fadado a trocar várias vezes de parceiro durante a vida, pois também é bastante volúvel. Amante do novo gosta de viajar pelo mundo sempre em busca de novidades, estudando e tentando compreender seus semelhantes. Tem mente investigadora, grande versatilidade mental, excelente memória e extraordinária capacidade para fazer várias coisas ou mesmo ler vários livros ao mesmo tempo. Tem, também, a tendência para começar várias coisas e não as terminar, ou seja, tem muitos começos e poucos fins. Tal tendência deve-se principalmente ao desejo quase obstinado de conhecer, o que o torna por vezes dispersivo, instável e inquieto, mudando de atividade tão logo encontre outra que lhe proporcione mais prazer e interesse. Como é adaptável e ousado, deve assumir responsabilidades, pois pode naturalmente ser levado a posições de autoridade. Sua inteligência aguçada e suas ideias originais podem levá-lo a seguir carreiras nas áreas de educação, filosofia ou pesquisa científica. Magnânimo e gentil, é excelente para lidar com as pessoas e normalmente sabe perceber as oportunidades. Como não gosta de obedecer a ordens de outros, prefere ocupar posições gerenciais ou trabalhar por conta própria, pois quer ter liberdade para agir do seu jeito. Seu humanitarismo e suas   aspirações espirituais podem levá-lo a ter atividades ligadas a reformas sociais ou à religião. Mas o seu senso dramático natural pode fazer com que seja atraído pelo mundo do entretenimento. Por ser muito popular, por vezes se deixa levar por palavras elogiosas e comete erros infantis de julgamento de seus semelhantes, sendo muitas vezes vítima de intrigas e ciúmes. É inteligente e deve aprender a desligar-se das opiniões alheias e a seguir sua própria intuição, mesmo que a princípio lhe pareça tarefa das mais difíceis. Novos ambientes, novos amigos, novas paisagens, novos horizontes, farão parte ativa de sua vida.',
    6: 'MISSÃO – 6 (AMOR)\nPontos positivos: fraternidade universal, compaixão, confiabilidade, solidariedade, idealismo, vida doméstica, humanitarismo, firmeza e temperamento artístico. \nPontos negativos: ansiedade, timidez, teimosia, perfeccionismo, egocentrismo e desconfiança. A família é a sua principal fonte de preocupação. É extremamente prático, equilibrado, sentimental, leal, compreensivo e sempre pronto a ajudar aqueles que solicitam sua ajuda. É honesto, digno de confiança, bondoso, alegre e muito amoroso. Sente-se muito feliz em cargos que requeiram alta responsabilidade, na qual possa contribuir para regularizar, ajustar e harmonizar as coisas — ocupações relacionadas com instituições de caridade, ensino de modo geral, treinamento, ocupações esotéricas, tanto em consultas como didáticas, ou trabalhos em comunidades que requeiram uma liderança justa e íntegra. O Amor é a sua palavra mágica; amor pelas crianças, velhos, familiares, principalmente os filhos (se os tiver), pelos amigos e, em alguns casos, pelos inimigos. Às vezes pode ser explosivo, principalmente quando duvidam de sua capacidade profissional ou põem à prova    sua honestidade. Porém, logo esquece as ofensas e os desafetos    e perdoa todos, indistintamente. A sua Missão é, invariavelmente, o Amor Universal. Caso se desvie dela, sofrerá na maturidade o peso dessa negligência, culpando tudo e a todos pelos fracassos e decepções que por certo aparecerão, mesmo que, materialmente seja bem-sucedido. Se atentar para estas palavras, terá êxito social e será muito popular, principalmente se se dedicar a algum tipo de Ciência Oculta ou à religião, pois é muito convincente e consegue facilmente fazer com que as pessoas o sigam. Cuidado com o fanatismo, pois ele poderá levá-lo a sofrer do mal que anunciar. Amigável e perceptivo, prefere a companhia de pessoas inteligentes e enérgicas, com quem possa dialogar e ter estímulo mental. A sua natureza amigável e solidária sugere que os outros o procuram para pedir conselhos e apoio. Nos seus relacionamentos aparenta ser forte e enérgico, mas o lado sensível da sua natureza implica que pode, ocasionalmente, ficar inquieto e entediado. Embora tenha interesse por outros países e   pessoas estrangeiras, prefere ter uma vida doméstica estável.',
    7: 'MISSÃO – 7 (SABEDORIA)\nA Sabedoria é a sua palavra de ordem. É um número por vezes enigmático, pois tanto pode ser profundo, calado e introspectivo, como estranho, sociável, falante e cativante. À primeira impressão, as pessoas sentem até um certo receio de se aproximar dele, pois seu ar superior e até assustador, leva-as a se afastar em princípio. Porém após algumas horas ou somente minutos de convivência, todos lhe são simpáticos e anseiam por permanecer em sua companhia, e se alguns não o fazem, é por pura inveja ou despeito. O 7 a vibração da perfeição e da sabedoria, tanto na vida material como no mundo espiritual. É o número das qualidades psíquicas e também da sensibilidade. Como não gosta de ser mandado, deve evitar qualquer trabalho subalterno, saindo-se melhor como educador, guia espiritual, orador, pregador, pesquisador, explorador, historiador, ou em atividades ligadas ao saber, na mais pura expressão da verdade. É autoritário e gosta por demais que os outros o venham consultar. Poderá também ser excelente escritor metafísico, de esoterismo, religião ou de assuntos relacionados à autoajuda. Existe tendência a tornar-se introvertido e meditativo, na maturidade. É fundamental que se cerque de pessoas inteligentes e, de preferência, que comunguem de seus ideais. O sucesso mundano, ou seja, material, é importante para a sua realização espiritual, por isso, não deve ser negligenciado qualquer projeto que possa trazer benefícios financeiros, pois sua Missão depende, em muito, desse poder material para se desenvolver e obter o êxito que almeja. Carismático e ambicioso, tem charme e calor humano. Como é sociável e generoso, se sai especialmente bem em atividades que envolvam o trato com as outras pessoas. Independente e voltado para o sucesso, gosta de se manter em atividade e pensa em termos grandiosos. Com as suas poderosas convicções e seu desejo de servir, pode se dedicar a um ideal ou liderar um movimento que seja de utilidade prática para as outras pessoas. Ainda que consiga demonstrar uma profunda compaixão, as suas emoções fortes podem indicar que precisa evitar extremos ou agir muito impulsivamente. ',
    8: 'MISSÃO – 8 (JUSTIÇA)\nPontos positivos: liderança, meticulosidade, autoritário, protetor, bons princípios, trabalhador, com grande poder de cura e justiça. \nPontos negativos: impaciente, dominador e intolerante. É organizado, muito dedicado aos seus negócios, excelente para lidar com situações concretas, dominador, prático, direto, seguro de sie generoso. A sua meta principal é o progresso material; progresso sólido, feito com justiça e elevado senso moral. É persistente, tem grande força de vontade e está sempre à espreita de oportunidades, dificilmente deixando que se lhe escapem. Tais características ficam, com o passar dos anos, cada vez mais acentuadas. O grande problema que deve encontrar na primeira metade da vida (até aos 45 anos), é a disciplina. Caso não seja disciplinado nessa fase, sofrerá muito na segunda, onde as oportunidades já não serão tantas, e a vitalidade já não será tão grande. Se disciplinado, sempre encontrará meios de se locupletar, seja trabalhando com afinco e esperteza, seja fazendo os outros trabalharem em seu proveito. Não é dado a discussões e a teimosias, preferindo contornar qualquer obstáculo que se entreponha entre si e o objetivo desejado. Quanto aos amigos, poderá ter tantos quanto os que lhe querem bem, pois sabe fazer amizades e também sabe desfazê-las, tão logo perceba que são interesseiras. Como dá grande importância à riqueza material, terá mais sucesso como alto executivo de uma grande empresa, proprietário de um complexo industrial, comércio em geral, advogado, político ou dedicar-se às Leis. Para que isso se concretize, deverá ter vontade firme, muita energia e autoridade. Através da sua inspirada conscientização e autodisciplina, poderá ser muito feliz trabalhando em prol da humanidade. É necessário aprender a se desapegar e um nível profundo e pode ser que só compreenda isso numa idade avançada. Quando reconhece esta necessidade adquire liberdade interior e passa a ter ideais mais profundos e humanitários. Altamente intuitivo, tem necessidade de transcender o lado mundano da vida que pode estimulá-lo a produzir um trabalho original ou fazê-lo entrar em contato com outros países.',
    9: 'MISSÃO – 9 (CONHECIMENTO)\nPontos positivos: idealismo, criatividade, sensibilidade, generosidade, magnetismo, caridade, desapego e popularidade. \nPontos negativos: nervosismo, falta de praticidade, preocupação exagerada, algum egoísmo e frustração. É aquele que ao longo da vida deve obter o máximo de conhecimentos possíveis, a fim de adquirir a Sabedoria. É universalista, generoso, sente compaixão por todos, quer melhorar a espécie humana e busca sempre a verdade. Por ser universalista, tanto o passado, o presente e o futuro fazem parte do seu dia-a-dia, vivendo nessa eterna busca, sem nada encontrar, pois já sabe de antemão todas as respostas. Como tem ampla visão do mundo e das coisas que o   cercam, principalmente os seus semelhantes, sofre por demais quando percebe injustiças ou mesmo quando alguém é maltratado, mesmo que todos digam que merece, sejam eles quem forem. Normalmente é impessoal, porque os interesses pessoais geralmente levam-no a desapontamentos. Por seu caráter, não deve ser egoísta e jamais esperar qualquer coisa dos outros, principalmente agradecimentos. Como é incompreendido, por ser universal, tem alguma dificuldade de relacionamento afetivo, dificilmente se dando bem no amor. Por vezes será chamado “irmão(ã) mais velho(a) da humanidade”, pois sua responsabilidade é quase sempre com os outros e muito pouco faz em prol de si mesmo. Pelo seu caráter justo e humanitário, encontrará muitas pessoas boas e prosperará e terá sucesso e felicidade. O amor que dedicará ao seu semelhante voltará duplicado e encontrará paz e harmonia. Na maturidade, caso tenha destinado a primeira parte da vida ao altruísmo, deve se dedicar ao ensino religioso, esotérico, a escrever ou a falar em público. Nessa fase, muitas pessoas dependerão de sua sabedoria e o (a) procurarão, esperando encontrar uma palavra de apoio, solidariedade e sabedoria. Deve, enfim, viver tão altruisticamente quanto possível e sempre buscar a harmonia entre seus irmãos. ',
    11: 'MISSÃO – 11 (FÉ)\nPontos positivos: equilibrado, concentrado, objetivo, entusiástico, espiritualizado, idealista, intuitivo, habilidade para cura, humanitário e grande capacidade psíquica. \nPontos negativos: complexo de superioridade, excessivamente emotiva, certo egoísmo, dominador e magoa-se facilmente. A grande virtude do 11 é a Fé. Esta Fé, em si mesmo, não representa propriamente religiosidade, mas também a Fé nos seus ideais, propósitos, pressentimentos, “sorte” ou em projetos que elabora. Consciente ou inconscientemente, a Fé em alguma coisa é a sua “mola mestra” que lhe empurra para o destino. Agindo corretamente, tem o potencial para inspirar as pessoas com seus ideais e imaginação; assim, evite investir suas energias mentais em coisas de pouca importância para não se desviar das suas metas positivas. Quando se concentra em suas energias positivas, é capaz de obter resultados produtivos. Cuide da sua dieta e da sua saúde e aprenda a relaxar de vez em quando. Tem a habilidade para encantar as pessoas, mas elas talvez não vejam que atrás da sua aparência ousada, sofre de confusão interior. Em idade avançada, terá uma sabedoria acumulada através do desenvolvimento espiritual, da compreensão universal e do potencial místico. Evite ser impulsivo, agir de forma extravagante ou se deixar levar por esquemas de enriquecimento rápido, investindo em planos de longo prazo. É um diplomata por excelência e possui a rara capacidade da harmonia e compreensão. É justo e está sempre pronto a “arbitrar” qualquer divergência entre colegas de profissão, amigos ou discórdias familiares. É incapaz de ferir propositadamente um seu semelhante e, quando isso acontece, prontamente pede desculpas e não repete mais o erro. Com o passar dos anos, as suas capacidades de liderança ficarão mais latentes e sólidas e as oportunidades de sucesso, prosperidade e fama aparecerão e com certeza as aproveitará. Estas características poderão levá-lo a ser um famoso orador, um renomado político ou mesmo um diplomata, incumbido de uma grande missão, principalmente de Paz, que lhe será imposta e que terá de usar de toda a sua perspicácia e habilidade para saber concretizá-la a bom termo. A parte final da vida (após os 60 anos) será repleta de paz, harmonia e felicidade.',
    22: 'MISSÃO – 22 (ESPERANÇA)\nPontos positivos: universalidade, intuição elevada, pragmatismo, praticidade, capacidade de organização e capacidade de resolução de problemas. \nPontos negativos: nervosismo, complexo de inferioridade, autoritarismo, preguiça e egoísmo. A Missão 22 é muito parecida com o “Destino 22”, pois é um número altruísta e voltado quase exclusivamente para a humanidade, para o todo, para a sabedoria. Vê tudo em larga escala e é altamente capaz de levar a bom termo qualquer projeto que vise o bem do Planeta. Mentalmente arguto e intuitivo, gosta de tomar as próprias decisões. É uma pessoa humanitária, que tem uma visão realista da vida e, com sua força, pode dar apoio a quem precisa. Com grande habilidade social, normalmente tem muitos amigos e admiradores. Para ter paz e harmonia, muitas vezes está disposto a fazer concessões ou sacrifícios. Considera a amizade e gosta de se dar com as pessoas que o inspirem a ser aventureiro e o divirtam. Se sente bem quando está na companhia dos que ama. A sua existência se pautará por testes: financeiros, amorosos ou existenciais. Porém, como é um ser superior, uma vez ciente deste fato, conseguirá se sair bem em qualquer atividade ou profissão e terá grandes satisfações no decurso de sua longa existência. Pode-se dedicar à política, ciências, filosofia, relações públicas, filantropia, esoterismo ou religiosidade. Deverá aprender a ser prático e idealista, a ter os “pés no chão”, a estudar os princípios fundamentais do ser humano e a trabalhar para construir um mundo melhor. Quaisquer tipos de vício (bebida, cigarro, drogas) são verdadeiros venenos para o seu organismo.',

}
DiadoNascimento = {
    1: 'DIA NATALÍCIO 1 = DIA DA LIDERANÇA Pontos positivos: liderança, criatividade, caráter progressista, vigor, otimismo, fortes convicções, competitividade, independência, sociabilidade. \nPontos negativos: arrogância, egoísmo, ciúme, antagonismo, egoísmo, excesso de orgulho, hesitação, impaciência. É um líder nato; gosta de mandar em vez de fazer. É criativo e original, tem raciocínio lógico e rápido, e é capaz de discutir sobre os mais variados assuntos, até mesmo aqueles que conhece superficialmente. Tem tendência a ser autoritário, de certa maneira possessivo e um tanto egocêntrico. Ainda que o lado aventureiro da sua natureza queira se expressar, a sua visão altamente pragmática e a preocupação com a segurança sugerem um certo materialismo. Isso indica que mesmo que o idealismo e a visão tenham um papel importante na sua vida, a preocupação com o dinheiro, ou com a sua falta, mantém os seus pés firmes no chão. Todavia, você é uma pessoa amável, que tem emoções fortes, inclinações humanitárias e grande capacidade de liderança. Embora não o sendo, o número 1 é considerado “frio”, calculista e pouco comunicativo. Raramente demonstra afeto e compaixão. Contudo, gosta de ser elogiado e admirado. Quando seus projetos ou desejos não se realizam ou são frustrados, o seu sistema metabólico sofre sobremaneira e os nervos e a pressão arterial são sobrecarregados, podendo causar graves danos ao organismo. Como líder, sente-se terrivelmente frustrado em posição subalterna e, por vezes, torna-se irascível, violento e inconsequente, sendo muito difícil trabalhar e conviver em sua companhia. Precisa também aprender que o mundo não gira ao seu redor, evitando a inclinação a ser autocentrado ou ditatorial. O nascido no dia um necessita saber para poder e querer. Estudar, projetar, manter a consistência no objetivo deve ser sua principal característica, pois tem tendência a deixar tudo pela metade ou a glória aos outros. Individualista e dinâmico, gosta de conhecer gente de diferentes círculos sociais. Como é sociável, aprecia a companhia de outras pessoas, especialmente as que são criativas e que o estimulam a se expressar. Ainda que seja leal e disposto a apoiar entes queridos, tem dúvidas e se sente indeciso nos seus relacionamentos. Para evitar desapontamentos, precisa encarar a sua vida emocional com leveza e lembrar que deseja ser feliz.',
    2: 'DIA NATALÍCIO 2 = DIA DA DIPLOMACIA Pontos positivos: gentileza, tato, boas parcerias, receptividade, intuição, consideração, harmonia, em princípio agrada a todos. \nPontos negativos: desconfiança, falta de objetividade, subserviência, excesso de sensibilidade, emotividade, egoísmo e tendência a ser desonesto quando não consegue atingir os objetivos pela maneira correta, ou seja, honestamente. É um ser diplomata por excelência. É aquele que harmoniza o grupo e a família; o que possui o dom da reconciliação. É cooperativo, aparentemente tímido e vulnerável, de certa maneira passivo, mas sempre atento aos detalhes de seu ambiente. Amante das diversões e sociável pode ser um grande amigo e uma boa companhia. Mesmo sendo amoroso e afetuoso, deve tomar cuidado na escolha dos seus relacionamentos para que eles sejam duradouros. Pode ser romântico com as pessoas de que gosta e mostrar honestamente os seus sentimentos, mas, para ser feliz, deve procurar segurança financeira. Com todo o seu charme, é muito atraente ao sexo oposto, porém não gosta de ser tolhido em seus movimentos. No trabalho sente-se melhor desenvolvendo atividades ligadas a grupos, pois com sua personalidade de certa forma amável e passiva, é amado por todos e é sempre excelente profissional. Não suporta ficar parado procurando sempre algo para fazer. É compreensivo com os sentimentos alheios e incapaz de ferir quem quer que seja. Um dos seus grandes defeitos é a inclinação para subestimar seus dotes e capacidades, tanto intelectuais como profissionais, sendo, muitas vezes, subordinado de pessoas com capacidades inferiores às suas. Evite deixar-se levar pela vontade de ter poder e de controlar, superando a tendência a ser crítico e a exigir a perfeição. Precisa encontrar o equilíbrio entre ajudar demais e sair de cena subitamente. No caso de se sentir inseguro, pode ser difícil levar adiante os seus planos pessoais, mas se encontrar alguma forma de expressão artística ou criativa, pode transformar seus sonhos em realidade. Com uma excelente habilidade social, pode prosperar como escritor, em atividades educacionais, publicidade, relações públicas, advocacia, diplomacia e atividades artísticas as mais variadas. Caso não seja culturalmente desenvolvido, pode-se tornar cruel, inescrupuloso e até violento, no propósito de atingir seus objetivos.',
    3: 'DIA NATALÍCIO 3 = DIA DA POPULARIDADE Pontos positivos: bem-humorado, feliz, amigável, produtivo, criativo, artístico, amor à liberdade, talento com as palavras, poder para desejar. \nPontos negativos: entendia-se facilmente, vaidoso, excesso de imaginação, orgulhoso, extravagante, comodista, preguiçoso e esbanjador. O nativo deste dia é um ser de rara animação, criatividade, expressão e popularidade. Pode parecer irresponsável para alguns que não o conhecem bem, mas na realidade é um ser altamente responsável, e prestativo com tudo e com todos. O três é intuitivo, original, honesto, dedicado à família e aos amigos (aos últimos, às vezes até demais). Sua ambição e sua personalidade atraente podem levá-lo ao topo de qualquer carreira. Seja na área bancária ou no mundo artístico, a sua aversão a ordens alheias pode levá-lo a posições superiores, como gerente ou executivo. Pode se sair particularmente bem na área teatral, como ator, diretor ou dramaturgo. Mas pode também usar a sua habilidade para lidar com as pessoas na área dos negócios, em que pode dar início a novos projetos nos quais ocupe uma posição de liderança. Como sabe delegar tarefas a seus subordinados, pode ser um excelente administrador ou trabalhar por conta própria. Pelo seu lado extrovertido e amistoso, não suporta ser criticado, apelando para o sentimentalismo daqueles que o criticam ou simplesmente lhe dão “conselhos”. É do tipo que trabalha em inúmeras atividades ao mesmo tempo e que quase sempre as deixa a meio caminho, ou seja, tem muitos começos e poucos fins. No decorrer de sua longa vida, terá inúmeras frustrações, e estas podem levá-lo a ter certos problemas físicos, principalmente o sistema nervoso abalado e também a contrair certas doenças de pele, que não se sabe como as supera com grande senso de humor. Na realidade, é um ser lutador, criativo e liberal, que usa de seus predicados, principalmente a oratória, como sustentáculo de sua vida.',
    4: 'DIA NATALÍCIO 4 = DIA DA PERSISTÊNCIA Pontos positivos: organização, autodisciplina, firmeza, trabalhador, habilidade, talento com as mãos, pragmático, confiável, preciso. \nPontos negativos: Falta de comunicação, rigidez, falta de sentimentos, procrastinação, autoritarismo, afeições ocultas, ressentimento, severidade. Os nascidos neste dia são muito disciplinados, constantes, regulares e ordeiros. Perseverantes em seus propósitos, incansáveis trabalhadores, dificilmente deixam de atingir seus objetivos. Normalmente honesto, sincero e conservador, o quatro adapta-se a trabalhos rotineiros, metódicos e que requerem esforço concentrado. Em virtude dos predicados acima descritos, o nascido neste dia sente-se mais à vontade lidando com situações rotineiras, já consagradas pelo uso, evitando o novo ou o incerto. A sólida estrutura e o poder de organização sugeridos por este dia, indicam que o seu possuidor precisa de estabilidade e ordem. Dotado de energia, habilidades práticas e muita determinação, pode ser bem-sucedido se estiver disposto a trabalhar duro. Preocupado com a segurança, quer ter uma base sólida para si mesmo e para os familiares. A maneira pragmática como encara a vida lhe confere um bom senso para os negócios e habilidade para alcançar o sucesso material. É muito econômico, amigos dos amigos, companheiro leal e dedicado, de grande responsabilidade e respeitador do sexo oposto, gostando também de construir coisas e trabalhar com as mãos. Tem tendência à obstinação, ao apego e à sua maneira de ser. Quando o contrariam ou fazem coisas que não gosta, pode se tornar rude, ofensivo e dominador, pois se ofende com certa facilidade. Porém, como é uma alma superior, não sabe guardar raiva, ressentimentos ou qualquer sentimento de revolta. Quando não consegue atingir seus objetivos ou os mesmos são adiados por circunstâncias adversas à sua vontade, fica tremendamente frustrado causando-lhe dor de cabeça, melancolia, distúrbios cardíacos e problemas renais. Em matéria de amor (relacionamentos) a mulher nascida neste dia é mais acessível e menos dramática do que o homem.',
    5: 'DIA NATALÍCIO 5 = DIA DA VERSATILIDADE Pontos positivos: versátil, adaptável, progressista, fortes instintos, magnético, sortudo, ousado, amante da liberdade, perspicaz e rápido, curioso, sociável. \nPontos negativos: pouco confiável, instável, procrastinador, inconsistente, excessivamente confiante, cabeça-dura. O nascido neste dia é normalmente divertido, alegre, ousado, dotado de poderes psíquicos, imaginação fértil, versatilidade e também amante da liberdade. Em virtude de ter os ouvidos muito sensíveis (não gosta de receber ordens), vive constantemente em busca de dinheiro, por vezes de maneiras totalmente inusitadas, sem qualquer medo de correr riscos. Gosta de estudar e de saber, para poder conseguir atingir seus objetivos com mais facilidade. Viajar por diversão, estudo ou satisfação do ego, também fazem parte da sua personalidade. É obstinado em seus propósitos (impaciente e impulsivo) e não descansa enquanto não consegue atingir seus objetivos, mesmo que tenha de usar de artifícios pouco convencionais ou prejudicar alguém. Gosta de estar em contato com o público, de preferência sendo o centro das atenções, e no trabalho sente-se melhor em ocupações que o coloquem em contato com pessoas, mas que estas lhe permitam agir e exprimir-se livremente. As frustrações, principalmente no âmbito profissional, que atrapalham seus planos, causam-lhe insônias, distúrbios psicológicos, falta de controle emocional que podem se transformar em violência. Quando quer ou é incentivado, consegue resultados fantásticos no terreno profissional, pois tem grande capacidade de discernimento, amplos conhecimentos e satisfação naquilo que faz. Porém, o seu lado obscuro, o lado “libertino”, leva-o a ter muitos começos e poucos fins. Quase nunca é bem-sucedido no amor (existe, é claro, as exceções), levando-o a trocar várias vezes de parceiros ao longo da sua duradoura vida. Tem, ainda, grande habilidade para encantar as pessoas com o seu otimismo e generosidade, mas deve tomar cuidado para não se tornar muito autoritário com quem ama.',
    6: 'DIA NATALÍCIO 6 = DIA DO AMOR Pontos positivos: universal, fraterno, compassivo, confiável, compreensivo, solidário, idealista, com inclinação doméstica, humanitarista, equilibrado, artístico. \nPontos negativos: insatisfeito, ansiedade, timidez, irracionalidade, teimosia, falta de harmonia, dominação, egoísmo, desconfiança, cinismo, egocentrismo. O nascido no dia seis é normalmente sentimental, muito equilibrado, compreensivo, adora a família, a casa, os amigos, os filhos (se os tiver) e é também excelente amante. Tem personalidade magnética e atrai sempre as atenções: em festas, reuniões, cursos, etc. Profissionalmente sente-se realizado numa posição superior, onde pode contribuir para o desenvolvimento da empresa, das coisas e principalmente das pessoas envolvidas. É perseverante e luta até o fim para atingir seus objetivos. Como é altamente sensível, quando contrariado, ou quando as coisas não correm como quer, pode-se tornar ciumento, nervoso e demonstrar possessividade, levando-o a ter atitudes enérgicas para defender seus princípios ideológicos. No tocante às frustrações amorosas, estas lhe causam quase sempre complicações nervosas e problemas ósseos. Gosta de se divertir, é atencioso, confiável e gentil. Tem uma alegria infantil e sempre será jovial. Naturalmente charmoso e sociável, tem muitos amigos e admiradores. Nos relacionamentos sérios, é uma pessoa romântica, idealista e leal, mas deve evitar se martirizar ou ser muito possessivo. Pode ser um parceiro dedicado, caloroso, atencioso e amoroso. Prático e econômico, tem tendência a ser idealista e despreocupado. Confiando na sua intuição e aprendendo a se conhecer, pode ter um bom relacionamento com as pessoas. Será excelente profissional no setor social (trabalhar com idosos, crianças, deficientes físicos ou mentais), em áreas esotéricas e religiosas, como professor, decorador, cozinheiro, ou tratando de embelezar o planeta.',
    7: 'DIA NATALÍCIO 7 = DIA DA INSPIRAÇÃO Pontos positivos: culto, confiável, meticuloso, idealista, honesto, com poderes psíquicos, científico, racional, reflexivo. \nPontos negativos: dissimulado, pouco amigável, fingido, cético, confuso quando lida com detalhes, inoportuno, indiferente, pouco sentimental, sensível às críticas. Além da integridade inerente do número sete, ele também possui em larga escala a independência de pensamento, a iniciativa e a ponderação. É também um perfeccionista e um tanto arredio a coisas e a novas amizades. Analítico e ponderado, o nascido neste dia é muitas vezes crítico e concentrado nos seus próprios interesses. Com uma necessidade constante de ter uma maior autoconsciência, gosta de reunir informações e se interessa por ler, escrever e por assuntos espirituais. A tendência a ser enigmático ou reservado sugere que pode ocasionalmente se sentir incompreendido. Embora não goste de contrair dívidas, precisa aprender a administrar a sua vida financeira. O desejo de desfrutar de uma boa vida sugere que precisa de algo que o motive e o inspire, pois há o perigo de cair na rotina e perder o rumo. Se o querem ver feliz, falem de religião, de filosofia (pura e simples), de ensino ou qualquer atividade ligada à espiritualidade. Companheiro (ou companheira) dedicado, quando se separa é por demais difícil se adaptar a uma nova relação. A Numerologia aconselha que se case tarde (após os 30 anos), depois de namorar muito e conhecer plenamente as características do companheiro. Durante a vida procura mais as coisas filosóficas e abstratas, que o atraem. Gosta sobremaneira que todos o venham consultar; tem tendência a se sufocar pelo trabalho e por isso pode ter ataques de nervosismo e alguma depressão. Normalmente sem vícios (álcool, cigarro e drogas); quando os tem lhe são altamente prejudiciais ao fígado, pâncreas e estômago. Também deve evitar qualquer tipo de jogo (dificilmente ganhará em loterias e seus afins).',
    8: 'DIA NATALÍCIO 8 = DIA DO ÊXITO MATERIAL Pontos positivos: liderança, meticulosidade, trabalhador, autoridade, proteção, poder de cura, bom juízo de valores. \nPontos negativos: impaciência, desperdício, intolerância, excesso de trabalho, dominação, desencoraja-se facilmente, falta de planejamento. O nascido neste dia é normalmente organizado, muito dedicado aos negócios, criativo e com enorme potencial para ganhar dinheiro. É justo, leal, prático, generoso (quando quer) e tem grande capacidade executiva e grande senso de justiça. Quando as coisas não saem como deseja, pode-se tornar direto (às vezes até demais), agressivo, com acessos de mau humor e com grande tendência a dominar a todos, indiscriminadamente, sejam eles parentem, amigos ou empregados. Em vista do seu grande potencial para ganhar dinheiro (ou outros bens materiais), será mais bem-sucedido como alto executivo, trabalhando por conta própria ou em alguma atividade em que o dinheiro esteja presente e em grande quantidade. Pertence ao mundo dos negócios e, por isso, deve sempre desenvolver a sua capacidade criativa para levar avante os seus projetos. Com a sua percepção aguçada e rápida, sabe avaliar as pessoas e as situações. Muito eficiente no trabalho, tem disposição para trabalhar duro e assumir responsabilidades. No entanto, precisa aprender a administrar ou a delegar a sua autoridade de maneira justa e imparcial. Manifesta, normalmente, uma aparência austera, fria e calculista; na realidade, é tremendamente carente, sentimental, justo e sempre pronto a ajudar quem necessita dele. Do outro lado, ou seja, do material, como “adora” dinheiro, vive desconfiado, descrente de quase tudo e até certo ponto pessimista, o que o leva a ter repentes de solidão, mau humor  e até um tanto ranzinza. Apesar disso, dificilmente é derrotado, superando todos os obstáculos que por ventura lhe apareçam pela frente. Qualquer vício lhe é prejudicial à saúde, principalmente o álcool, que pode lhe provocar graves distúrbios ao fígado e estômago.',
    9: 'DIA NATALÍCIO 9 = DIA DO HUMANISMO Pontos positivos: idealista, humanitário, criativo, sensível, generoso, magnético, poético, caridoso, desapegado, sortudo, popular. \nPontos negativos: frustrado, fragmentado, inseguro, egoísta, pouco prático, preocupado. O nativo deste dia é normalmente universalista: sente compaixão por todos e quer melhorar o gênero humano. Benevolência, ponderação e sensibilidade emocional estão associadas a este dia de nascimento. Tolerante e gentil, é na maior parte das vezes generoso e liberal. Sua intuição e seu poder psíquico indicam uma receptividade universal que, se for positivamente canalizada, pode inspirá-lo a buscar o caminho espiritual. Amante da verdade, normalmente generoso, independente, liberal, audacioso, corajoso, combativo, e não tem medo da derrota eventual, pois sabe que conseguirá o que deseja. Dificilmente tem paz de espírito e tranquilidade, pois tem facilidade em atrair discórdia e desentendimentos, afastando os amigos e as pessoas que o amam. É um ser muito contraditório, pois sendo humanista e bondoso não deveria ser arrogante e revoltoso, mas o é, e dessa maneira, destrói em minutos o que levou anos a construir. Em virtude da sua autoconfiança, normalmente    protela tudo, e às vezes acaba ficando em dificuldades financeiras; mas no final acaba se saindo bem. Não gosta de receber ordens e será mais bem-sucedido em assuntos relacionados com a religião, filantropia ou associações beneficentes, nas quais a inspiração, a bondade e a compreensão sejam necessárias. Quando transmite os seus poderes sentimentais, expressa o seu amor e afeição. Normalmente não se apega a nada nem a ninguém, sejam bens materiais, amigos, companheiros de jornada ou mesmo a pessoa amada, tendo ao longo da sua vida muitos desapontamentos amorosos e também algumas perdas de amizade. Adora viajar e conhecer novos lugares, novos países, novas pessoas. Qualquer vício lhe é tremendamente nocivo ao organismo, seja o hábito de beber, fumar ou qualquer outro, pois possui um organismo muito sensível e os vícios lhe prejudicam terrivelmente o sistema nervoso e o respiratório.',
    10: 'DIA NATALÍCIO 10 = DIA DA AUTOCONFIANÇA Pontos positivos: liderança, criatividade, caráter progressista, vigor, otimismo, fortes convicções, competitividade, independência. \nPontos negativos: arrogância, ciúme, egoísmo, orgulho, antagonismo, falta de controle, hesitação, impaciência. O nativo deste dia é audacioso, progressista, independente, prestativo, amigo, atraente fisicamente, cativante e sempre pronto a ajudar àqueles que lhe pedem auxílio. Pelo seu lado atraente e de certa forma arrogante, normalmente desperta inveja e antipatias. Tem grandes ambições e se esforça para realizá-las, mas talvez precise, antes, superar alguns obstáculos. Vigoroso e original, defende aquilo em que acredita, mesmo quando os outros não concordam. A sua habilidade para dar início às coisas com um espírito pioneiro o encoraja a viajar por lugares distantes e a se arriscar sozinho. Para ter sucesso profissionalmente, deve desenvolver a espiritualidade, pois caso contrário pode ser envolvido por pessoas inescrupulosas que tudo farão para o arruinar, e caso não possua esta característica, dificilmente terá competência para solucionar seus problemas. Sendo líder por natureza, ou trabalha só, ou em cargos de chefia, de preferência no ramo da engenharia, metalurgia, comércio, vendas ou diretamente com o público, pois é muito convincente. Tem admiração por pessoas que se realizam na vida. Pode ser que sonhe com um amor tão elevado e ideal que acabe sendo difícil encontrar alguém que corresponda às suas expectativas. Como pode alterar entre ser amoroso, espontâneo e afetuoso ou frio e distante, é vital ter um espaço para si mesmo a fim de equilibrar a sua sensibilidade. A sua personalidade naturalmente amigável é a garantia de muitas amizades e a sua hospitalidade sugere que é um bom anfitrião. Deve evitar todo e qualquer tipo de vício, principalmente o fumo, pois as suas vias respiratórias são frágeis e sofrerão terrivelmente com este vício.',
    11: 'DIA NATALÍCIO 11 = DIA DA HARMONIA Pontos positivos: equilíbrio, concentração, objetividade, entusiasmo, inspiração, espiritualidade, idealismo, intuição, capacidade de curar, humanitarismo, psiquismo. \nPontos negativos: complexo de superioridade, falta de objetivos, excesso de emotividade, magoa-se facilmente, egoísta, falta de clareza, dominador, mesquinho, hipersensível. Apesar de ter como lema a harmonia, a inspiração está sempre presente em sua vida. É um diplomata por   excelência; delicado nos termos, ações, possuindo tato e discernimento para qualquer problema ou ocasião. É um número mestre, e os seus possuidores quase sempre são carinhosos, sentimentais e necessitam tremendamente de um lar para se sentirem seguros, protegidos, pois não gostam de viver sozinhos. Inteligente e observador, é muitas vezes um parceiro ou amigo leal. Sincero e atencioso, precisa de um relacionamento estável e honesto. Embora possa refinar e melhorar as situações, precisa aprender a diferenciar entre ser autoritário e crítico e ajudar os que estão à sua volta. Parece uma contradição, e é, pois, sendo o 11 compreensivo por natureza, não deveria se importar com o pensamento alheio, mas se importa. Como ama a liberdade, necessita estar sempre ocupado para se sentir útil e feliz. É eficiente profissionalmente e poucos o acompanham em qualquer atividade, apesar de ser mais sonhador do que realizador e, em virtude disso, deve sempre procurar orientação técnica profissional para ser bem-sucedido, ter sucesso e ser feliz. Deve, ainda, tomar muito cuidado para que o seu intelecto não sufoque sua intuição, pois sendo psíquico, não pode vacilar ante os problemas. Para atingir seus objetivos, por vezes pode parecer submisso, mas, na realidade, consegue tudo o que deseja, pois é convincente, inteligente e perspicaz. Tem tendência à arrogância e qualquer vício lhe é prejudicial ao organismo.',
    12: 'DIA NATALÍCIO 12 = DIA DA AUTO-EXPRESSÃO Pontos positivos: criativo, atraente, capacidade de iniciativa, disciplinador, assertivo, confiante. \nPontos negativos: reclusivo, excêntrico, pouco cooperativo, excessivamente sensível, falta de autoestima. É comunicador nato; pela sua criatividade, expressão, e argumentação, consegue convencer todas as pessoas. Tem gosto artístico, habilidade manual, é idealista, quase sempre está de bom humor e dificilmente desiste dos seus ideais. A sua praticidade e agilidade em fazer as coisas, leva-o a assumir mais compromissos do que pode cumprir e, dessa maneira, está sempre atarefado, cheio de trabalho, necessitando aprender a dosar suas energias, pois com certeza vai precisar delas em casos especiais. O nativo deste dia é íntegro em seus propósitos, justo, de caráter leal, franco, liberal, de natureza ativa, qualidades comerciais, diplomacia, aptidão para o comando e gosta de ajudar o próximo. É muito respeitador da fé e das crenças, suas e alheias. Gosta de ser popular, tendo mesmo tendência para se tornar político ou trabalhar em atividades de interesses sociais. É bondoso por natureza, e por vezes em vista dessa característica, tende a ser explorado pelas pessoas mais chegadas a ele. O nativo deste dia, quando quer captar simpatias ou simplesmente fazer novas amizades, é capaz de assumir papel de vítima, mantendo as suas “presas” em constante ansiedade e insegurança, ora pelo apaixonado, ora indiferente e distante. Tem natural impaciência e quase sempre se deixa dominar nervosismo, estando sujeito a crise de depressão, insatisfação e se entendia com facilidade. Para evitar este lado negativo da sua personalidade, deve-se manter sempre interessado em algo construtivo, de preferência ligado a atividades intelectuais. Os eventuais fracassos, as decepções e frustrações (principalmente com os “amigos”), podem lhe causar problemas no sistema nervoso, hipertensão e, em certos casos, problemas renais.',
    13: 'DIA NATALÍCIO 13 = DIA DA PERÍCIA Pontos positivos: ambicioso, criativo, tem amor pela liberdade, é expressivo e com grande iniciativa. \nPontos negativos: impulsivo, indeciso, autoritário, pouco emotivo e rebelde. O nativo deste dia é meticuloso, autoritário, sistemático, prático, econômico, trabalhador incansável, sempre lutando em prol dos seus objetivos, não poupando esforços para atingi-los. Como o nome do dia sugere (dia da Perícia), é tremendamente habilidoso em reformas, em transformações e mudanças, quando estas são de seu interesse. Ousado e dinâmico, prefere fazer acontecer a esperar. Está associado ao trabalho duro e pode realizar muito através da determinação e do talento. Como é influenciado pelo ambiente, precisa ser decidido e controlar a direção da sua vida. É alegre e talentoso, podendo se sair bem em profissões artísticas ou de entretenimento. É justo, bondoso e fica muito chateado quando pressente ou constata que alguém foi injustiçado e também se revolta com as competições desleais ou quando alguém é enganado. É muito amoroso, mas encontra certa dificuldade em expressar seus sentimentos. É prestativo, dedicado, bom amigo, mas quase nunca expressa essas emoções, parecendo isto sim, indiferente, frio, materialista e calculista, reprimindo seus sentimentos, sejam eles de dor, decepções ou mesmo de alegria. Não sabe viver sem amor, carinho e afeto. Porém, como não expressa esses sentimentos, poucos o compreendem e o conhecem realmente. A sua vida é pautada pelos negócios, pela dedicação à profissão (de preferência em indústrias de grande porte, construção civil ou administração pública), pois sendo íntegro, organizado, honesto e eficiente, rapidamente consegue um lugar de destaque nesses campos. Contrariedades e decepções, principalmente com parentes e amigos, podem lhe causar dores de cabeça, problemas na fala e no sistema respiratório; porém, consegue superar esses inconvenientes de modo admirável.',
    14: 'DIA NATALÍCIO 14 = DIA DA COMPREENSÃO Pontos positivos: ações decididas. Trabalhador, sortudo, criativo, pragmático, imaginativo e habilidoso. \nPontos negativos: excessivamente cauteloso ou impulsivo, instável, sem consideração e teimoso. Potencial intelectual, perspectiva pragmática e forte determinação são algumas das qualidades associadas a esta data de nascimento. Tem grande desejo de estabelecer uma base sólida e ter sucesso através do trabalho duro, podendo alcançar o topo da sua profissão. Quase sempre vive no presente, pouco se preocupando com o futuro, sendo normalmente líder em qualquer situação ou grupo. É negociante nato; não gosta de ficar parado e está sempre procurando alguma coisa para fazer, para ganhar dinheiro, pouco se importando se terá lucro ou prejuízo. Tem uma boa capacidade de gerenciamento e talento para solucionar problemas além de uma aptidão natural para escrever – que sempre pode usar de forma criativa ou na vida profissional. Um dos seus pontos mais fortes é trabalhar com muita movimentação de capitais, seja como especulador financeiro, agente imobiliário ou negociador. Admira as pessoas que têm uma maneira pouco convencional ou original de abordar a vida e deseja ter um parceiro que seja disciplinado e trabalhador. Com seu charme e seu comportamento naturalmente assertivo, sempre atrai pessoas que acreditam em você. Para ser feliz no casamento, o cônjuge deve gostar de vida agitada, deve gostar de viajar, de conhecer novos lugares e jamais se prender a um único local, pois caso contrário a união não terá final feliz. Por ser naturalmente bondoso e emotivo, normalmente é presa fácil dos inescrupulosos, principalmente quando querem seus favores. Por vezes (quando desenvolvido espiritualmente) é profético, com tendências construtivas e destrutivas. Quando o contrariam ou frustram seus ideais, pode ter problemas respiratórios, algumas alergias e até desenvolver doenças imaginárias. Caso consiga controlar a impulsividade e adquira prudência, pode-se tornar muito bem-sucedido financeiramente e também socialmente.',
    15: 'DIA NATALÍCIO 15 = DIA DO MAGNETISMO PESSOAL Pontos positivos: disposição, generosidade, responsabilidade, amabilidade, cooperação, apreço e ideias criativas. \nPontos negativos: inquieto, autocentrado, medo de mudar, perda da fé, preocupação e certa indecisão. Como o dia sugere: “magnetismo pessoal”, esse magnetismo é levado às últimas consequências, pois tanto homens como mulheres lhe acham simpático, agradável, afetuoso e interessante. No comércio ou num escritório, normalmente é o alvo das atenções, mais, é claro, do sexo oposto, sendo dessa maneira, invejado por todos. Não tem grandes problemas financeiros (raramente fica pobre), pois é esperto, inteligente, perseverante e sempre encontra uma maneira de ganhar dinheiro. Não tem veia de avarento, mas é de certa forma apegado ao dinheiro, preferindo gastá-lo com o companheiro (a), na casa ou em algo proveitoso para si. Tem grande calor humano, adora reuniões sociais, ama   a vida e as pessoas e não se importa em gastar, principalmente quando está em companhia da pessoa amada. Deve-se casar (ou unir-se) com pessoa que tenha    afinidade com seus propósitos e gostos; caso contrário, a união não dará certo e qualquer separação o faz sofrer em demasia, levando-o inclusive ao isolamento. Tem a capacidade de atrair oportunidades e condições harmônicas. Entre outros, pode assumir cargos comunitários, direção de espaços esotéricos, ou outros que exijam compreensão e mão firme para manter a ordem e a disciplina em grupos, isto pela sua capacidade, responsabilidade e habilidade natural de compreender, unir e harmonizar. Mesmo em idade avançada, parecerá sempre jovem. É honesto, digno de confiança, amável e bondoso. Tem inclinações artísticas, podendo se destacar na oratória, nas artes plásticas, música ou representação. Os seus pontos fracos organicamente são a garganta, os pulmões e o fígado; por isso, qualquer vício lhe é prejudicial, principalmente o álcool e o cigarro.',
    16: 'DIA NATALÍCIO 16 = DIA DO TRIUNFO Pontos positivos: responsabilidade, integridade, intuição, sociabilidade, cooperação e discernimento. \nPontos negativos: preocupação, insatisfação, irritabilidade, egoísmo, ceticismo e falta de solidariedade. É um extremista! Quem nasce neste dia pode ser o mais miserável dos seres, ou o maior dos ricos. Dependendo da vida que levar, pode transformar seu possuidor numa pessoa poderosa, rica, um ser de pleno sucesso e felicidade; no outro extremo, pode arruinar, levar ao desmando, transformar o ser em um elemento arrogante, prepotente, orgulhoso e dominador. Aconselha-se que os nascidos neste dia vivam tão altruisticamente quanto possível, que tenham pensamentos positivos, sentimentos elevados e, desta maneira, com absoluta certeza atingirão o sucesso e serão muito felizes. É, também, o número do equilíbrio entre o material e o espiritual. Se teimar em viver fraudulentamente, querendo levar vantagem em tudo e com todos, poderá ver seus planos frustrados, ser traído por amigos e ainda contrair doenças inesperadas. É analítico, cético (só acredita no que vê ou é comprovado), gosta de conhecer a essência e o âmago das coisas e pessoas e também apresenta acentuado caráter perfeccionista. Pelo seu senso de perspicácia, gosta e consegue desvendar coisas misteriosas e também de acumular conhecimentos. É um ser de grande sensibilidade, intuição e inspiração, tendo mesmo qualidades psíquicas sem qualquer estudo do assunto. Um dos seus grandes defeitos é gostar que as pessoas que o rodeiam vivam conforme seus moldes e, quando isso não ocorre, torna-se mal-humorado e até colérico. Por esse seu temperamento de presunção, geralmente vive isolado, porém, na realidade tem grande desejo de afeto e principalmente compreensão. Apesar de tudo isso, não gosta que interfiram na sua vida, nos seus projetos, mesmo quando estes não dão ou não estão dando certo. Em vista da sua grande sensibilidade, que é atrativa em vários segmentos, deve tomar muito cuidado com falsos amigos, descontentamentos, com a ansiedade e principalmente com alguns perigos físicos, como o excesso de velocidade em automóveis. Deveria, portanto, fugir da agitação das grandes cidades, dando preferência a viver no campo ou então perto da água (rios, lagos e oceano). Pela sua habilidade analítica aguçada, pode trabalhar na área das ciências, da matemática ou da informática. Por outro lado, como tem habilidade para se comunicar e lidar com as pessoas, pode ser atraído por uma carreira na área de educação, advocacia ou atividades que envolvam o trato com os outros. As frustrações ao longo da vida (que não são poucas), podem levá-lo à obesidade, a ter problemas estomacais, de pele e até algumas doenças imaginárias (hipocondrianismo).',
    17: 'DIA NATALÍCIO 17 = DIA DA PERSPICÁCIA Pontos positivos: ponderação, especialização, capacidade de planejamento, senso para os negócios, trabalhador, científico e atrai dinheiro. \nPontos negativos: distante, teimoso, descuidado, mal-humorado, sensível, crítico, preocupado e desconfiado. O nativo deste dia é naturalmente um líder inteligente e arguto. Está quase sempre de bom humor e consegue ser simpático até com os opositores. Diferente do líder (1), se quiser ser bem-sucedido profissionalmente, deve trabalhar em grupo, com outras pessoas, e somente em casos esporádicos consegue se sair bem trabalhando só. Nasceu para ser bem-sucedido no plano material, através de muito trabalho, persistência e determinação. Por qualquer destes caminhos que se aventurar, será um vencedor: negócios imobiliários, comércio em geral, como executivo trabalhando com muitos subordinados ou comércio exterior, pois é eficiente profissionalmente, econômico (não avarento) e grande articulador. É popular, está sempre bem-humorado (mesmo que seja para disfarçar alguma decepção ou tristeza), e só depende dele ser ou não bem-sucedido financeiramente. Normalmente é bem relacionado com o sexo oposto e está sempre cercado de admiradores, apesar de não gostar de revelar os seus pensamentos ou sentimentos. Busca a companhia de pessoas ambiciosas, determinadas e trabalhadoras. No plano afetivo, as pessoas que mais o atraem são aquelas que demonstram gerar harmonia e paz.',
    18: 'DIA NATALÍCIO 18 = DIA DO PODER MENTAL Pontos positivos: progressista, assertivo, alto poder de intuição, corajoso, resoluto, eficiente, capacidade de aconselhar. \nPontos negativos: emoções descontroladas, falta de ordem, egoísmo, vaidade, ambição desmedida e incapacidade de concluir trabalhos ou projetos. Dinâmico e ativo, frequentemente deseja o poder e precisa de desafios constantes. Pode, às vezes, ser crítico, difícil de ser agradado ou ter inclinações para a controvérsia. Com sua personalidade e poder marcantes, pode usar esses predicados para ajudar seus semelhantes, dar bons conselhos e resolver problemas alheios. O dia 18 é o dia dos bruxos, dos magos, dos médiuns famosos e dos religiosos poderosos. É, também, o dia dos seres felizes e infelizes, ou seja, o dia dos extremos. Se agir com dignidade, trabalhar e estudar para o bem da humanidade, com certeza alcançará altos postos e será feliz. No outro extremo, ou seja, se agir fraudulentamente, sempre querendo levar vantagem em tudo e com todos, é certo que se arruinará, mais dia, menos dia. De natureza psíquica, o 18 é profundamente espiritual; tem contato com dimensões superiores, visões, de ampla e irrestrita intuição, é comunicativo e altamente sensível a problemas espirituais. É um vencedor nato! É inteligente, tem mente ativa e desperta e consegue se sobressair em quase todas as atividades.  Os seus maiores inimigos são seus próprios defeitos: vaidade e ambição. No outro extremo, ou seja, caso não seja evoluído espiritualmente, tem tendência a se entregar ao pessimismo, por medo do desconhecido, do futuro e, quase sempre, nesse estado, acaba sendo presa fácil para as adversidades, terminando na ruína completa. Tem no seu lado positivo, a intuição e independência prestativa e desinteressada, intelectual, emotiva e requintada, gostando de discutir sobre os mais diversos assuntos. No negativo, pode às vezes ser crítico, difícil de ser agradado ou ter inclinações para a controvérsia. Porém, quando quer, consegue superar todos os obstáculos, pois é um grande lutador, encarando todos de frente, sem medo de nada nem de ninguém. Tem caráter afetivo e é também dedicado aos outros. Como é muito prestativo, consegue fazer amizades com facilidade e ter relacionamentos duradouros, pois também é muito amoroso, apesar de ser facilmente vulnerável e explosivo. Quando as coisas não correm como gostaria, fica frustrado e tem tendência a desanimar, a ter problemas cardíacos, dores de cabeça e a ter o sistema nervoso abalado. Em vista disso, qualquer vício, principalmente o cigarro e o álcool, lhe são verdadeiros venenos para o organismo.',
    19: 'DIA NATALÍCIO 19 = DIA DO CARÁTER Pontos positivos: dinamismo, criatividade, liderança, progressismo, otimismo, convicções fortes, competitividade, independência, espírito gregário e muita sorte. \nPontos negativos: egocentrismo, preocupação, medo de ser rejeitado, materialismo, impaciência e tendência à depressão quando não consegue atingir os objetivos. Ambição e humanitarismo são algumas das principais características deste excelente dia. Decisivo e cheio de recursos, possui uma visão penetrante, mas o lado sonhador da sua natureza é compassivo, idealista e criativo. Apesar de ser sensível, a necessidade de ser alguém pode levá-lo a ser dramático e a querer ser o centro das atenções. Muitas vezes há um forte desejo de estabelecer uma identidade individual. Para isso, precisa primeiro superar a influência causada pela pressão   externa. Para os outros, os nascidos neste dia aparentam ser uma pessoa confiante, resistente e cheia de recursos, mas as tensões internas podem causar instabilidade emocional. Assim como o (5), quer mudanças, é versátil e está sempre desejando o melhor para si e também para a família. É independente, artístico, original e dotado de espírito de iniciativa e criatividade. 19 é o dia do sucesso, da prosperidade e também da felicidade. Esta vibração altamente positiva tem em si embutido também certa tendência à arrogância à teimosia e à vaidade. Possui grande poder de realização, mas se irrita com certa facilidade, tendo acessos de crises de violência que normalmente afetam sua saúde. Apesar desta negatividade, jamais guarda rancor de quem quer que seja e rapidamente esquece qualquer ofensa de que é vítima. Sendo o seu lema o caráter, na mais pura expressão, o seu possuidor é um reformador, e como tal, sempre pensa em uma maneira de transformar o mundo, conquistando dessa maneira, simpatias e a admiração de todos. Quando quer alguma coisa, é capaz de gestos teatrais e até atitudes extremas e não aceita seguir o tradicional. Amigável e popular não têm dificuldades para fazer amigos e seduzir as pessoas. Em geral atraente, tem muitos admiradores do sexo oposto. Contudo, se deixar levar pela instabilidade, as suas experiências amorosas podem ser como andar de montanha russa, com muitas perdas e sucessos. Pelo seu instinto “paternal”, as decepções (principalmente com amigos), frustrações (ideológicas) e fracassos (profissionais), podem afetar o seu coração, a visão e também o sistema auditivo.',
    20: 'DIA NATALÍCIO 20 = DIA DA SENSIBILIDADE Pontos positivos: boas parcerias, gentileza, tato, receptividade, intuição, consideração, harmonia, presença agradável, embaixador da boa vontade. \nPontos negativos: desconfiança, subserviência, timidez, sensibilidade excessiva, certo egoísmo, tendência a magoar-se com facilidade. O nativo deste dia é sensível, intuitivo, adaptável e compreensivo, e gosta de pertencer a um grupo. Em geral aprecia atividades cooperativas, nas quais possa interagir, compartilhar experiências e aprender com as outras pessoas. É um ser humano caseiro, ordeiro, pacífico, sempre distribuindo amor e paz a todos os conhecidos e sabe como poucos criar atmosferas agradáveis e harmoniosas. Receptivo e intuitivo, gosta de se comunicar com os outros e, muitas vezes, tem talento para lidar com as pessoas, sendo excelente mediador ou negociador. Se desenvolver os seus dons artísticos, pode escolher uma carreira nas áreas de educação, como professor, ou na área editorial como redator ou jornalista. Os relacionamentos afetivos têm muita importância em sua vida e é vital que não caia em situações de dependência. Como precisa de companheirismo, em geral não gosta de ficar sozinho por muito tempo. Felizmente, seu charme, habilidade diplomática e excelente poder de persuasão podem ajudá-lo a ter amigos e amantes. Contudo, os seus sentimentos podem, às vezes, ser extremos e difíceis para as outras pessoas, mesmo quando genuinamente deseja a harmonia. Como é encantador, sociável e divertido, normalmente é um excelente anfitrião. É tremendamente hábil e é capaz de fazer qualquer serviço mais rápido e melhor do que qualquer outra pessoa. Também tem grande habilidade para tratar com o público, sendo dessa maneira, um bom político ou trabalhar em repartições governamentais voltadas para o povo. As decepções, frustrações e contrariedades podem lhe causar graves prejuízos ao sistema nervoso. ',
    21: 'DIA NATALÍCIO 21 = DIA DO IDEALISMO Pontos positivos:  inspiração, criatividade, uniões por amor e relacionamentos duradouros. \nPontos negativos: dependência, temperamental, nervoso, falta de visão, medo de mudanças. Apesar de ser idealista e liberal, o nativo deste dia necessita da companhia de outras pessoas, pois dessa irmandade depende o seu sucesso e também o seu bem-estar. Nasceu para manifestar e expressar seus sentimentos e ideias. É ambicioso, mas dispersivo, e dificilmente consegue acabar o que começa, ou seja, tem muitos começos e poucos fins. Muitas vezes musical e artístico, com queda para o teatro e a oratória, pode ser excelente ator de cinema ou teatro. Espirituoso e divertido, tem uma personalidade brilhante. Como é amigável, tem uma vida social ativa. Atraído por pessoas independentes e bem-sucedidas, quando se envolve em algum relacionamento precisa de liberdade e autoconfiança. Um dos lados da sua natureza afetiva é ser muito dramático, desconfiado e retraído, causando certo desconforto nos relacionamentos. Quando casado, é excelente companheiro e gosta por demais da família, a ponto de grandes sacrifícios para tornar tudo harmonioso. É por demais emotivo, sujeito a extremos, que o leva quase sempre a um estado de depressão. Em vista dessa sua fragilidade e inconstância, encontrará sérios obstáculos na juventude, mas por fim terá sucesso na idade mais madura, pois tem absoluta certeza de que tudo acabará bem. São suas qualidades positivas: amizade, idealismo e capacidade de entretenimento. Após os 40 anos, a determinação e a vontade em conseguir sucesso material se fortalecerá e as privações antes dessa idade servirão como exemplo e   também como um orientador que o conduzirá ao sucesso desejado. O nascido no dia 21 é considerado como “sortudo”, pois, de maneira inexplicável (para os outros números), consegue “tudo” o que deseja. Está sempre procurando um parceiro que seja sensível e compreensivo e que tenha fé nas suas habilidades. Cuidado com as doenças psicossomáticas adquiridas das frustrações, decepções e contrariedades.',
    22: 'DIA NATALÍCIO 22 = DIA DA PRATICIDADE Pontos positivos: intuição elevada, pragmatismo, praticidade, habilidade com as mãos, capacidade de organização, realismo, resolução de problemas e empreendedor. \nPontos negativos: esquemas de enriquecimento rápido, nervosismo, autoritarismo, materialismo, falta de visão, ganância, autopromoção e preguiça. O nativo deste dia, como especificado, é tremendamente prático, adapta-se a qualquer tipo de trabalho e para atingir seus objetivos (caso os tenha), é capaz de feitos heroicos. Como tem visão futurista e perfeccionista ao extremo, inúmeras vezes deixa de aproveitar as ocasiões que se lhe deparam no dia a dia. Sendo um ser totalmente independente, tanto em considerar as coisas como as pessoas, despreza os convencionalismos e as tradições, o que normalmente lhe é prejudicial profissionalmente. Ainda com referência ao futurismo do nativo, normalmente não tem grandes ambições materiais e pouco se importa em ganhar ou acumular fortuna e, dessa maneira, está sujeito a muitos altos e baixos durante a vida. Para viver adequadamente, deve manter o equilíbrio entre as emoções e a praticidade. É de certa maneira nervoso, tenso e necessita muito de repouso. Este seu lado negativo, em certos momentos, mostra um certo desequilíbrio emocional, tendendo à intolerância, impaciência, não se entendendo a si próprio e, assim vivendo em constante conflito com os mais próximos. Sendo um duplo “2”, também na vida real as coisas tendem a acontecer-lhe em dobro, tanto para o bem, como para o mal, portanto, deve fazer um esforço redobrado para viver construtivamente e em harmonia com toda a humanidade. As 22 enxergas longe; em vista disso, deve sempre procurar profissões ou ocupações de caráter mais geral e não aquelas de interesse pessoal. Nasceu para a humanidade e em vista disso, tem enorme responsabilidade com seus semelhantes e para que todos os seus dons (que são muitos) possam se manifestar deve trabalhar como alto executivo, político, escritor de temas universalistas, artista, conferencista ou chanceler. Ainda que tenha visões e convicções fortes, o ser 22 é astuto no que diz respeito aos seus relacionamentos. Como o amor e o companheirismo são tão importantes, muitas vezes cede ou usa as suas habilidades diplomáticas para manter a harmonia nos seus relacionamentos. Ocasionalmente, quando não recebe a atenção ou afeição de que precisa, pode ficar inseguro ou ciumento. Em vista do seu alto grau de sensibilidade, está sujeito    a distúrbios psíquicos, nervosos e também alterações do sistema glandular, principalmente quando reprimem ou lhe frustram seus ideais. Os vícios, principalmente o cigarro e o álcool são verdadeiros venenos para o seu organismo.',
    23: 'DIA NATALÍCIO 23 = DIA DA PERSUASÃO Pontos positivos: lealdade, responsabilidade, adora viajar, comunicativo, intuitivo, criativo, versátil, paciente, persuasivo e confiável. \nPontos negativos: egoísta, inseguro, teimoso, inflexível, crítico, reservado e preconceituoso.  Se quiser levar uma vida sem maiores problemas, tem   de trabalhar ao máximo o seu lado compreensivo. Sensível e emotivo, o 23 é um idealista, audacioso e romântico. Atraído por pessoas vigorosas, usa os seus poderosos sentimentos expressando de forma dramática o amor que sente. No entanto, algumas vezes, seus relacionamentos pessoais podem ser perturbados devido à instabilidade do seu humor ou ao seu excessivo materialismo. Apesar de ser paciente, nunca descansa até conseguir o que quer. É o número do sucesso material, do dinheiro, e o seu portador precisa aprender a seguir caminhos profissionais, de preferência os de alto nível e não enveredar para os negócios, pois sendo intelectual, pertence ao mundo sensível e não está apto a servir ninguém. Tem personalidade marcante, rara inteligência (aprende tudo com grande facilidade), inclinação social (gosta de festas e reuniões) e se dá melhor com o sexo oposto do que com o próprio. Em virtude da sua grande sensibilidade, quando lhe tolhem os objetivos ou reprimem seus ideais, pode sofrer sobremaneira do sistema nervoso. Usualmente é versátil, pensa com rapidez, e é dotado de uma atitude profissional e uma mente cheia de ideias criativas. Gosta de viagens, aventuras e de conhecer pessoas novas. Ainda que esconda os seus sentimentos, a sua natureza idealista e confiável indica que quando fala o que pensa pode ser terrivelmente franco. Mesmo que às vezes possa ser destemido, deve tomar cuidado para que essa franqueza não seja ofensiva às pessoas. Pode ser um excelente diplomata, político, médico, psiquiatra, psicólogo, terapeuta holístico, escritor metafísico ou até viver no meio artístico, conseguindo com esta profissão fama e alta posição social.',
    24: 'DIA NATALÍCIO 24 = DIA DA UNIÃO Pontos positivos: energia, idealismo, habilidades práticas, forte determinação, honestidade, franqueza, justiça, harmonia, diplomacia, alegria, generosidade, amor à casa, ativo e enérgico. \nPontos negativos: materialista, muito econômico, aversão à rotina, pouco confiável, dominador, teimoso, vingativo e ciumento. Quem nasce no dia 24, além dos predicados acima descritos, é também grande amigo, amante da verdade e tolerante com as falhas alheias. Esteja onde estiver e com quem, é quase sempre o centro das atenções, não pela beleza física ou porte, mas sim pela inteligência, calor humano e compreensão de tudo e com todos. É muito hábil em trabalhos manuais, sejam eles mecânicos, eletroeletrônicos, de marcenaria, ou concertos rotineiros os mais variados. Tem, também, gostos gastronômicos, sendo excelente cozinheiro e nas festas que participa é sempre solicitado para fazer os aperitivos ou o churrasco. Em virtude da sua grande sensibilidade, tem tendência a proteger os fracos e oprimidos e a se deixar levar pelo sofrimento alheio. No amor, caso seja demasiadamente sonhador, ardente e romântico, dificilmente se ajustará à vida monótona do romance, pois se sentirá incompreendido e solitário. Como tem presença marcante e cativante, é sempre preferível ir pessoalmente a algum lugar, em vez de telefonar ou escrever. Pode ser grande médico ou psicólogo, pois é muito sensível e compreensível. É confiável e justo, embora às vezes possa ser reservado, e tem a tendência a acreditar que as ações falam mais alto do que as palavras. Com essa maneira pragmática de encarar a vida, desenvolve um bom senso para os negócios e habilidade para superar os obstáculos e ser bem-sucedido. Mesmo que vencer seja muito importante para este nativo, deve evitar a tendência a ser autocentrado ou ditatorial. ',
    25: 'DIA NATALÍCIO 25 = DIA DO PROGRESSO Pontos positivos: altamente intuitivo, perfeccionista, perceptivo, mente criativa, ponderado, talento para lidar com as pessoas e capacidade para ganhar dinheiro. \nPontos negativos: impulsivo, impaciente, excessivamente emotivo, ciumento, reservado, instável, crítico e de certa maneira mal-humorado. O nativo deste dia, além da ambição material inerente ao ser humano, vive constantemente em busca do desejo da moralidade. É um pensador, um estudioso e, em vista disso, profundo conhecedor de vários assuntos, podendo se destacar e ter sucesso nos mais variados segmentos, como ciência, ocultismo, filosofia ou sobre a Natureza na sua mais abrangente expressão. O 25 é perfeccionista, exigente (consigo e com os outros), diplomata, versátil, com grande capacidade intuitiva, senso analítico e perspicaz. É rápido e enérgico, apesar de intuitivo e ponderado. O desejo de perfeição o impele a trabalhar duro e a ser produtivo. No entanto, precisa ser menos impaciente ou crítico quando as coisas não correm de acordo com os seus planos. Como tem dons proféticos e desenvolvida intuição, por vezes é instável e sujeito a vacilações e flutuações na sua personalidade. Tem como seu grande defeito, subestimar as suas qualidades, sendo ao longo da vida subjugado por pessoas muito inferiores a si. Como é honesto, bondoso e leal, julga que os outros principalmente os “amigos” também o são e, assim, vive sendo usado por essas pessoas, que tudo fazem para lhe tirar dinheiro e também para fazê-lo de empregado. Apesar dessas decepções ou frustrações e fracassos ocasionais, enfrenta tudo com muita valentia, mas pode ter problemas estomacais, como úlceras, sofrer de algum mal cardíaco ou pulmonar, na qual é recomendada a total abstinência ao cigarro. Deve a todo custo evitar o álcool, pois seu organismo frágil não suporta tal vício, embriagando-se com certa   facilidade e, dessa forma, metendo-se em confusões que jamais entraria se estivesse sóbrio. Quando se sente inspirado por um ideal ou causa, defende bravamente as suas fortes convicções. Carismático e direto, gosta de pertencer a um grupo maior ou de trabalhar com o público, pois é amigável e sociável, um ser devotado e com fortes laços familiares. ',
    26: 'DIA NATALÍCIO 26 = DIA DA JUSTIÇA Pontos positivos: prático, atencioso, orgulhoso da família, entusiástico, corajoso, justo e perseverante. \nPontos negativos: teimoso, rebelde, falta de entusiasmo, de persistência e com relacionamentos instáveis. A justiça na sua mais pura expressão, a perseverança e a moderação são as principais características do nativo deste dia. Tem, também, grande capacidade de discernimento, competência e organização, jamais desistindo dos seus objetivos e ideais, mesmo em algumas ocasiões parecendo indeciso, não sabendo muito bem o que quer. Tem uma maneira pragmática de abordar a vida, habilidade executiva e bom senso para negócios. Normalmente responsável e íntegro, com senso estético natural e amor à vida doméstica, precisa construir uma boa base sólida ou encontrar uma estabilidade verdadeira. Tem personalidade marcante e certo ar de superioridade, que com certeza lhe garantem certas inimizades e algumas perturbações. Quando é contrariado, torna-se agressivo e mal-humorado. O nativo deste dia é normalmente um ser solitário, de certa forma incompreendido, parecendo frio e calculista; na realidade, é uma extraordinária alma humana, sempre pronto a ajudar os fracos, os amigos e aqueles que necessitam de ajuda humanitária. Nasceu para mandar. É muito organizado, justo, de aspecto intelectual, com grande cultura e senso de responsabilidade. É também elegante no vestir e despreza o modernismo, preferindo o convencional. Ainda que possa ser uma pessoa teimosa e com ideias fixas, quando está apaixonado pode ser responsável e dedicado a fazer grandes sacrifícios. Frustrações e decepções podem lhe causar problemas biliares, dores de cabeça, reumatismo e problemas de circulação sanguínea. ',
    27: 'DIA NATALÍCIAO27 = DIA DA AUDÁCIA Pontos positivos: versátil, imaginativo, criativo, resoluto, corajoso, compreensivo, inventivo, espiritual, audaz e com grande força mental. \nPontos negativos: brigão, inquieto, nervoso, desconfiado e protelador. O nativo deste excelente dia é normalmente conhecedor dos mistérios da vida e pode, se quiser, ir a extremos: para o bem ou para o mal, e normalmente com 18 anos já definiu o caminho que vai percorrer. É um ser idealista e sensível, intuitivo e analítico, com uma mente fértil e criativa, que consegue impressionar as outras pessoas com os seus pensamentos originais. Às vezes pode parecer dissimulado, racional ou distante, mas, na verdade, pode estar ocultando tensões internas. Essas tensões podem incluir tendências impulsivas, indecisão ou desconfiança sobre mudanças. Se desenvolver a sua capacidade de comunicação, pode superar a resistência a expressar os seus sentimentos mais profundos. Por mais que necessite ficar só, deve evitar isolar-se. Pode ter mais paz quando cria uma atmosfera harmônica à sua volta. Tem personalidade audaciosa, liberal, corajosa, combativa e independente; é também grande amante da liberdade e não suporta dar satisfação dos seus atos, preferindo trabalhar só. Quase sempre bem-sucedido no plano material, dificilmente tem maiores problemas de dinheiro, pois sabe como consegui-lo. Tem elevado senso de fraternidade e mente Universal. É afetuoso, emotivo, nervoso e de certa maneira um tanto extravagante, principalmente em se tratando de sua aparência. O amor, a afeição e dedicação ao semelhante representam muito e é capaz de grandes sacrifícios pelos que ama. É por demais pacíficos e jamais procura problemas, o que não quer dizer que seja covarde; muito pelo contrário, pois se pode tornar violento quando atingido por injustiças e ingratidões. Caso não tenha uma existência superior e altruísta, as frustrações, fracassos e decepções podem lhe causar perturbações cardíacas e algum tipo de problema cerebral.',
    28: 'DIA NATALÍCIO 28 = DIA DO QUERER Pontos positivos: compaixão, progressismo, temperamento artístico, ambição, trabalho, vida doméstica estável e voluntarioso. \nPontos negativos: sonhador, falta de compaixão, autoritário, agressividade, falta de confiança, orgulho, vive se queixando e é excessivamente dependente. É muito contraditório, pois nasceu com o dom do querer, mas vive se queixando. Livre dessa face doentia poderá se impor a tudo e a todos, pelo seu admirável senso diplomático e espírito de justiça e compreensão. Será mais bem-sucedido como chefe ou em atividades independentes, pois não gosta de ser mandado nem criticado. Como tem grande vitalidade e energia física, pode se dedicar a várias atividades ao mesmo tempo sem se cansar. Como o 29, tem muitos começos e poucos fins, e por isso deixa de aproveitar as oportunidades de se tornar famoso e rico. Tem tendência a aumentar seus aborrecimentos, embora não aparente e nem concorde com isso e, dessa maneira, se sujeita a muitos desapontamentos. Sempre pronto para a ação e para novos empreendimentos, aceita corajosamente os desafios da vida, e com o seu entusiasmo pode com facilidade inspirar as outras pessoas, se não a se juntar, pelo menos, a apoiá-lo nas suas aventuras. Para superar todos os pontos negativos, deve desenvolver otimismo, autoconfiança e perseverança de propósitos. Estudar a natureza humana, procurando compreendê-la, deixar de ser egoísta e se elevar sempre, seja no plano material como intelectual, são a base de sustentação para que atinja o sucesso e a felicidade tão almejadas.',
    29: 'DIA NATALÍCIO 29 = DIA DA ESPIRITUALIDADE Pontos positivos: inspiração, equilíbrio, paz interior, generosidade, sucesso, criatividade, intuição, misticismo, capacidade de liderança e mundanismo. \nPontos negativos: nervosismo, mau humor, extremismo, falta de consideração, arrogância e orgulho. Quem nasce neste dia e souber direcionar sua vida para o bem, conseguirá tudo o que desejar. O seu sensível intelecto, combinado com a sua poderosa memória e capacidade de liderança, garante que pode dar uma contribuição valiosa em diversas áreas. Com atração por ocupações e atividades relacionadas ao trato com as pessoas, pode se sobressair em carreiras como as de professor, treinador, publicitário, relações públicas, escritor, advogado ou em qualquer ramo holístico, sobretudo na área de autoajuda. Tem grande capacidade auditiva e senso variado de humor, podendo em questão de segundos ir da alegria contagiante a mais negativa das formas: a violência. É um ser altamente espiritualizado e as pessoas que com ele convivem devem também comungar de seus ideais e participarem dos seus projetos, pois caso contrário podem-se tornar seus inimigos. Para conseguir se realizar usa de imaginação, brandura de modos (quando não o contrariam) além de elevado espírito de conciliação. Como é moral e intelectualmente elevado, usa a fé, o idealismo e o conhecimento inspirado para tingir seus objetivos, seus ideais. Sendo extremista, está sujeito a muitas mudanças comportamentais ao longo de sua duradoura vida e, por isso, deve procurar interesses definidos e manter o ânimo calmo e equilibrado, pois a sua normal agitação o torna disperso, provocando na primeira metade da vida (até os 45 anos), muitos começos e poucos fins. Há uma grande tendência a se voltar para a religião ou esoterismo após os 45 anos e deve cuidar para não cair no fanatismo e também não induzir os demais, pois como é inspirado e cativante, tem facilidade de convencer quem quer que seja. Para ter sucesso na vida, necessita de harmonia em tudo e com todos, pois tem muita dificuldade em se situar no meio termo. Normalmente quem nasce neste dia afortunado, tem diversos relacionamentos, sendo considerado por todos os místicos o dia do casamento, dos divórcios e das separações (existem as exceções, é claro!). Pode sofrer inúmeras decepções amorosas e, se encontrar a sua “cara metade”, normalmente casa-se cedo. Porém necessita controlar suas emoções e evitar atitudes apaixonadas, pois as uniões desfeitas causam-lhe imensos sofrimentos e dificuldades para se ajustar a uma nova relação. Muito observador, precisa aprender a ser menos crítico e ter mais consideração pelas pessoas à sua volta. Embora possua uma excelente capacidade de raciocínio, em geral julga as situações segundo os seus sentimentos e precisa estar em contato com suas emoções mais profundas. Depois que desenvolve os seus critérios de julgamento e aprende a pensar por si mesmo, pode se tornar menos dependente da opinião das outras pessoas. Tem tendência à obesidade, a sofrer de males estomacais e hepáticos. Por isso, quando se frustra ou as coisas não correm como deseja, ou é traído e enganado, sofre em demasia, principalmente nos órgãos citados. Uma das características mais incríveis deste número, é a cura pura e simples (sem qualquer tipo de remédio) de moléstias de difícil diagnóstico. Fumar lhe é altamente prejudicial à saúde.',
    30: 'DIA NATALÍCIO 30 = DIA DA REALIZAÇÃO Pontos positivos: amor à diversão, lealdade, amizade, talento com as palavras, criatividade e generosidade. \nPontos negativos: preguiça, obstinação, impaciência, insegurança, indiferença, desperdício de energia e não gosta de ser criticado. Amável e caloroso, gosta de atividades sociais e pode ser excepcionalmente carismático e leal. Estar apaixonado ou emocionalmente satisfeito é um requisito fundamental; na sua busca pela felicidade, deve evitar ser excessivamente indulgente ou impaciente. Detesta ser criticado, pois é altamente sensível e não suporta se ver “despido” de seus princípios e ideais. O 0 (zero) à direita do 3 (três) mostra claramente que existe uma tendência à auto anulação, à autodesvalorização, subestimando-se em demasia. Precisa constantemente se conscientizar de seu grande valor e de sua habilidade em superar dificuldades e, principalmente, impor-se antes a si próprio para depois conquistar o respeito e a admiração dos demais. O poder do amor é a sua maior virtude e com a sua natureza romântica, seu temperamento apaixonado e sua generosidade, as pessoas se sentem atraídas pelo seu carisma e encanto. Embora esteja disposto a fazer grandes sacrifícios pelos    seus entes queridos, precisa superar a tendência a permitir que seus sentimentos controlem a sua mente. Tem personalidade marcante, de certa forma perfeccionista, independente e não suporta ser mandado ou trabalhar em cargos ou funções subalternas. Caso não seja moralmente desenvolvido, poderá tentar atingir seus objetivos de forma ilegal, fraudulentamente, usando de artifícios pouco convencionais (chegando mesmo a ser cruel), e como não consegue disfarçar seus sentimentos, quase sempre é pego e acaba se arruinando e arruinando os parentes, principalmente aqueles mais próximos, como filhos, irmãos ou o cônjuge. Fora desse lado negativo, normalmente é compreensivo e tolerante com tudo e com todos (apesar de às vezes pensar que está sempre certo). Quando lhe tolhem seus ideais ou frustram seus objetivos, o seu sistema nervoso sofre sobremaneira.',
    31: 'DIA NATALÍCIO 31 = DIA DA HABILIDADE Pontos positivos: liderança, criatividade, progressista, vigoroso, otimista, fortes convicções, competitivo, independente e habilidoso. \nPontos negativos: arrogância, ciúme, egoísmo, orgulho, fraqueza de caráter, hesitação e impaciência. Como o número indica, os seus nativos possuem grande habilidade, capacidade, autoridade, e gostam de segurança econômica, e também dão grande valor às suas realizações, esquecendo-se, em muitos casos, de si próprio. Tem ideias originais, um bom-senso de forma e habilidade para ser bem-sucedido nos negócios, se tiver um plano de ação e colocá-lo em prática no seu ritmo. Uma das suas boas características é trabalhar em serviços comunitários, mas que os participantes não lhe tolham os passos ou a iniciativa. Como confia em todos, por causa da sua boa-fé, normalmente são mal interpretados e enfrenta mais obstáculos que os demais, pois pensa que todos são como ele, honestos, retos e competentes, o que não é verdade e, assim, está sempre às voltas com problemas financeiros e também profissionais. O nativo deste dia vive num mundo só seu e a maioria das pessoas tem certa dificuldade em compreendê-lo e para viver bem em sua companhia, a pessoa deve ser leal e compreensiva. Faz amigos e inimigos com a mesma facilidade. São suas características marcantes: trabalhar duro, ser honesto, leal, determinado e econômico. Jamais esquece um favor ou uma ofensa. Com o seu charme e sua habilidade para irradiar calor humano, atrai as pessoas, seja no trato individual ou público. Como é muito sociável, é um excelente anfitrião e tem compaixão pelos problemas alheios. Normalmente se sente atraído por gente determinada, mas deve tomar cuidado para não se envolver em jogos de poder com os seus parceiros. Normalmente as pessoas nascidas neste dia se esforçam muito para manter os relacionamentos harmoniosos, embora ambos os sexos tenham inclinação à inquietude. Se encontrar a sua cara metade, deve se casar cedo, pois a vida de casado acalma-o. Como adora viajar, pura e simplesmente, o parceiro deve comungar desses ideais, caso contrário, a relação ficará tremendamente prejudicada, pois o nativo deste dia não suporta ser contrariado em seus desejos e opiniões.',

}
AnoPessoal = {
    0: 'A vida de todos os seres humanos transcorre em ciclos sucessivos de nove anos. Cada um desses anos — de 1 a 9 — tem suas vibrações próprias que não podem ser evitadas nem alteradas. Cada ano tem o seu conjunto de inﬂuências, oportunidades e obstáculos. Quando se sabe antecipadamente o que nos pode acontecer, podemos nos preparar adequadamente para tirar o máximo proveito dessas oportunidades e evitar os contratempos tão comuns aos menos avisados.',
    1: 'ANO PESSOAL 1 (plantando as sementes) Este é o ano para começar coisas novas, o ano que estabelece o estilo de todo o ciclo de nove anos. É o momento de tomar iniciativa e mostrar coragem e determinação. Para ter sucesso e conquistar a felicidade, a pessoa precisa ser independente, criativa, segura, seletiva e seguir sua própria intuição. A armadilha a ser evitada NUMEROLOGIA CABALÍSTICA: “A ÚLTIMA FRONTEIRA” 175 é a falta de iniciativa, a qual poderá muito bem resultar numa apatia que se estenderá por todo o ciclo.',
    2: 'ANO PESSOAL 2 – (as sementes criam raízes) Este é o ano de agir com discrição e ser paciente, mantendo-se receptivo às ideias dos outros e permanecendo em segundo plano. É uma fase muito boa para amizades e relacionamentos. Para ter sucesso e conquistar a felicidade, a pessoa precisa ser delicada, diplomática e cooperativa no trato com as outras pessoas. As armadilhas a serem evitadas são a hipersensibilidade e a propensão para envolver-se em discussões e/ou ser excessivamente atrevido.',
    3: 'ANO PESSOAL 3 – (surgem os primeiros brotos) Este é um ano de boa saúde e de intensa vida social que, no entanto, poderá trazer tensão emocional. É uma fase boa para divertimentos, viagens, crescimento pessoal e para cultivar novas amizades. O sucesso e a felicidade resultam do fato de se dar vazão à própria criatividade e de expressar-se construtivamente através das palavras — escrevendo, lecionando, representando ou cantando. As armadilhas a serem evitadas são a extravagância (que pode acarretar dificuldades financeiras no ano “4”) e a perda de oportunidades em decorrência da dispersão das energias. O “3” é também um ano propenso a romances, sendo que as pessoas casadas que sucumbirem a relacionamentos extraconjugais nesse ano, é quase certo que estarão preparando o caminho para o divórcio ou separação no ano “4”.',
    4: 'ANO PESSOAL 4 – (cavando e capinando) Este é um ano de restrições, de trabalho duro e de grandes despesas. É o momento de se construir para o futuro. O sucesso e a felicidade resultam da autodisciplina, de ser metódico e de dar forma concreta às suas ideias. Esta é uma boa fase para lidar com propriedades e imóveis. A armadilha a ser evitada é a da negligência com a saúde — poderá haver algum problema com os ossos e/ou dentes.',
    5: 'ANO PESSOAL 5 – (formam-se os botões) Este é um ano para deixar-se levar por condições em constante mudança, para viver o presente sem fazer muitos planos para o futuro, para abandonar-se aos seus impulsos! É um ano em ritmo acelerado, propenso a acidentes e cheio de mudanças inesperadas — do bom para o ruim, do ruim para o bom. É um ano para divertimentos, sem se pensar no futuro. O sucesso e a felicidade resultam de a pessoa ser adaptável, de procurar desenvolver   a própria personalidade e de aproveitar as oportunidades antes que sejam perdidas. As armadilhas a serem evitadas são a dispersão de energias, o excesso de atividades sexuais e o mau uso da liberdade pessoal em prejuízo dos outros.',
    6: 'ANO PESSOAL 6 – (floração) Este é um ano para o casamento e as responsabilidades domésticas. É uma fase em que alguma doença crônica talvez se manifeste para que possa ser tratada. Pode também ser um ano de problemas relativos a viagens, tais como defeitos no carro, bagagem perdida e assim por diante. O sucesso e a felicidade resultam de uma dedicação altruísta à família e à comunidade. As armadilhas a serem evitadas consistem em ser excessivamente idealista ou propenso a discussões e em esperar demasiado dos outros.',
    7: 'ANO PESSOAL 7 – (as plantas dão fruto) Este é um ano para o isolamento e o descanso, um ano para o estudo e o aperfeiçoamento interior. Esta é uma fase de introspecção. O sucesso e a felicidade resultam do estudo dos significados últimos da vida. As atividades materialistas devem ser evitadas. O dinheiro só aparecerá se não se correr atrás dele. Quanto menor for a ambição, melhores serão os resultados e vice-versa. As armadilhas a serem evitadas são as de negligenciar a saúde, forçar decisões, ser excessivamente crítico e permitir que temores e complexos submersos venham à tona.',
    8: 'ANO PESSOAL 8 – (época da colheita) Este é um ano dinâmico e materialista. Os negócios provavelmente deverão prosperar. Poderá haver grandes perdas ou grandes ganhos, dependendo de como se lidou com os outros anos do ciclo. É uma excelente fase para se comprar e vender imóveis. É a hora de pagar e de cobrar as dívidas. O dinheiro pode surgir de fontes inesperadas. O sucesso e a felicidade resultam da coragem de ousar grandes feitos, de utilizar o bom-senso, de preocupar-se com o dinheiro, de ser organizado, prático e eficiente. As armadilhas a serem evitadas são as de ser emotivo e sentimental.',
    9: 'ANO PESSOAL 9 – (época de limpar a terra após a colheita e prepará-la para um novo plantio) Este é um ano de faxina entre o fim de um ciclo e o começo do próximo. É o momento de livrar-se de tudo o que for desnecessário ou estiver gasto pelo uso — especialmente pessoas para as quais não existe mais lugar em sua vida. É uma boa fase para escrever, representar, viajar e para dedicar-se a estudos metafísicos. Haverá alguma espécie de perda um relacionamento poderá terminar. O sucesso e a felicidade resultam da compaixão, das atividades humanitárias, do desapego emocional e de deixar que se vá o que quer que comece a sair da sua vida. As armadilhas a serem evitadas são as de ser ciumento e/ou possessivo.',
}
MêsPessoal = {
    1: 'MÊS PESSOAL 1 – É o mês dos pioneiros, das pessoas influentes, dos inventores e dos planejadores. Porém, é um mês onde também as pessoas tendem a dominar, consciente ou inconscientemente, todos os envolvidos. Não é um bom mês para se fazer novos amigos e é melhor se ater à rotina nesse sentido. Para negócios e especulações, é excelente.',
    2: 'MÊS PESSOAL 2 – É um mês passivo e receptivo. As pessoas devem se manter calmas, gentis, bondosas, organizadas e escrupulosas. Como é um mês de hesitações, aconselhamos que nenhum projeto seja iniciado neste período, sob pena de não dar certo. É também um mês de dúvidas e de resoluções importantes.',
    3: 'MÊS PESSOAL 3 – É o mês dos extrovertidos, das pessoas inteligentes, criativas e espirituosas. É um bom período para se fazer novos amigos, discutir, planejar e empreender ação às suas ideias. É o mês dos ambiciosos, dos orgulhosos e daqueles que procuram prazer pessoal.',
    4: 'MÊS PESSOAL 4 – É o mês dos realistas e dos equilibrados e também daqueles em quem se pode confiar. Deve ser mantida a justiça a todo custo, pois a pessoa estará sujeita a algumas depressões e até acessos de raiva, caso as coisas não lhe corram como deseja. Deve ser organizado, prático, racional e equilibrado em todos os sentidos, para poder driblar as adversidades.',
    5: 'MÊS PESSOAL 5 – É o mês dos inteligentes, espertos, brilhantes e impacientes. Será um mês movimentado e é um excelente período para conhecer e manter novas amizades. Novas sensações e novos projetos devem ser tentados, pois a colheita é certa. Como 5 é o número da sensualidade, é um excelente período para novos relacionamentos e manter os já existentes.',
    6: 'MÊS PESSOAL 6 – Neste período deve se manter calmo, sereno e equilibrado. É um mês em que a pessoa deve manter a rotina, permanecer em casa ou trabalhando sem inovar. Não deve se meter em confusões, discussões ou iniciar novos projetos neste período.',
    7: 'MÊS PESSOAL 7 – É um mês de introspecção; os filósofos, os místicos e os ocultistas têm neste mês o seu apogeu, pois convida-os para que se mantenham isolados, contemplativos, compenetrados e discretos. As pessoas devem se manter calmas, evitar discussões e não decidir definitivamente sobre planos e projetos.',
    8: 'MÊS PESSOAL 8 – É o mês dos empresários, dos políticos e das pessoas que trabalham com o público em geral. Excelente período para iniciar ou concluir negócios, principalmente os que envolvem muito dinheiro. Porém, como os nervos estarão à flor da pele, é aconselhável manter-se a calma e não discutir sob hipótese alguma. A pessoa estará cativante, analítica, diplomática e muito sensível.',
    9: 'MÊS PESSOAL 9 – É o mês da espiritualidade e da intelectualidade. Bom para o romantismo, para poesia, para as artes, filosofia e espiritualidade, os objetivos, neste período, devem ser canalizados para a humanidade como um todo. As pessoas estão confiantes, independentes, corajosas e determinantes.',
    11: 'MÊS PESSOAL 11 – É o mês das grandes realizações e das revelações, sejam elas de cunho material, mental ou espiritual. É o período para se pôr as ideias em ordem, pois as pessoas estarão intuitivas, vibrantes, imaginativas e as ideias estarão aflorando neste período.'
}
DiaPessoal = {0: '',
              1: 'DIA PESSOAL 1 – É um dia para se cultivar a individualidade, a independência e a capacidade de trabalhar sozinho. Deve terminar o que estiver fazendo e somente depois, começar algo novo. Como é o número da individualidade, a pessoa deve agir somente pelos seus pensamentos e não se deixar levar pelos “conselhos” alheios.',
              2: 'DIA PESSOAL 2 – É o dia das associações, da harmonia pessoal e de se procurar seguir os conselhos alheios. A pessoa deve cultivar a paciência, o tato, a cooperação e a lealdade.',
              3: 'DIA PESSOAL 3 – Neste dia a pessoa necessita cultivar a criatividade, os contatos sociais e a expressão de seus ideais. Pode e deve ser ousado e expor suas ideias com clareza, de forma simples e sem rodeios. Faça tudo o que desejar e não abuse do: cigarro, álcool ou qualquer outra droga.',
              4: 'DIA PESSOAL 4 – É um dia para o trabalho duro, para se conquistar aquilo que se deseja. Também, e por isso, deve, neste dia, cultivar a paciência, a confiança e a disposição para servir.',
              5: 'DIA PESSOAL 5 – É o dia da adaptabilidade e da responsabilidade. Há necessidade da pessoa mudar seu ponto de vista e perceber que o mundo não gira em torno de si e que alguma mudança deve ser feita para que haja uma perfeita união entre seus desejos e os desejos alheios.',
              6: 'DIA PESSOAL 6 – É um dia em que as suas ideias devem ser defendidas e de servir a todos com alegria. É um dia de grande responsabilidade profissional, onde deve lutar para manter a liderança que alcançou. Deve também manter a calma, ouvir mais do que falar e, assim, demonstrar seu ponto de vista sem qualquer alarde ou imposição.',
              7: 'DIA PESSOAL 7 – É um dia para ficar só, para meditar e adquirir sabedoria. A pessoa necessita desenvolver seus poderes mentais, estudar, meditar e buscar o significado último da vida e tornar-se um especialista em qualquer assunto. Procure não beber neste dia e fale o mínimo possível.',
              8: 'DIA PESSOAL 8 – É o dia da eficiência e a pessoa deve desenvolver os seus projetos, principalmente aqueles que envolvem muito dinheiro. Neste dia estará eficiente e totalmente cônscio dos seus deveres, devendo falar com calma e pausadamente, não deixando qualquer dúvida quanto os seus planos e projetos.',
              9: 'DIA PESSOAL 9 – A pessoa necessita, neste dia, colocar o interesse dos outros a frente dos seus. É um excelente dia para se pensar na humanidade, no próximo. Bom para se fazer doações, estudar, falar em público, dar e receber conselhos e não começar absolutamente nada.',
              11: 'DIA PESSOAL 11 – É um dia de grande energia e poder. A pessoa deve confiar na própria intuição e não se deixar levar por conselhos de outras pessoas. Deve conservar-se humilde e inspirar os que os rodeiam com seu exemplo. Bom período para viajar, adquirir e conservar amizades, escrever aos amigos, perdoar os inimigos, cultivar a paciência e meditar sobre os novos projetos.',
              22: 'DIA PESSOAL 22 – É o dia em que a pessoa deve pensar altruisticamente, fazer somente o que é bom para a humanidade, ouvir as pessoas e praticar a justiça. Como é um número superpoderoso, este dia deve ser aproveitado ao máximo e a pessoa deve se levantar o mais cedo possível, fazer uma longa oração ao raiar do Sol, comer pouco, meditar muito, não usar qualquer tipo de droga (cigarro, álcool ou outras), ler algo sobre esoterismo e criar objetivos que sejam bons para toda a coletividade.',
              }
PrimeiroCiclo = {
    1: 'O 1 (um) no primeiro Ciclo de Vida indica um período difícil. Quando criança, a pessoa necessita aprender a desenvolver sua individualidade, pois caso contrário, na juventude e adolescência ou mesmo até à entrada do 2º Ciclo, terá problemas emocionais e grande dificuldade de se estabilizar profissionalmente. O ideal é que a criança nesse Ciclo tenha liberdade acima do normal e não frear os seus instintos em hipótese alguma. No caso de pessoa maior de 18 anos e que ainda esteja no primeiro Ciclo e tenha sido reprimida, ou seja, não tenha tido educação condizente, que absorva estes ensinamentos e os coloque em prática imediatamente.',
    2: 'O 2 (dois) no primeiro Ciclo de Vida, indica uma criança extremamente mimada, que possivelmente sofreu grande influência da mãe ou dos avós. É natural que na adolescência, em vista da possessividade familiar, pense em casar-se o mais cedo possível e isso é muito comum, principalmente entre os homens.',
    3: 'O 3 (três) indica uma infância e adolescência feliz, despreocupada e com muitos amigos. Não é um período particularmente favorável ao aprendizado, que deverá ocorrer a partir do segundo Ciclo, mas haverá provavelmente muitas oportunidades para a expressão de ideias e emoções, após os 18 anos (alguns com menos idade), através das artes em geral, da música, do teatro e escrita. Não é um bom período para contrair matrimônio.',
    4: 'O 4 (quatro) é uma indicação de restrições familiares e de trabalho duro na escola. Normalmente tem pais e parentes severos que impõem seu comando não admitindo qualquer contestação. Na juventude e mesmo após os 18 anos, em vista dessa restrição, o jovem é pouco criativo, limitando-se a obedecer, pouco ou nada criando, levando-o a um relacionamento por vezes desastroso, onde normalmente o cônjuge é que manda e dita normas e leis.',
    5: 'O 5 (cinco) no primeiro Ciclo de Vida indica muitas mudanças e uma liberdade que às vezes é demasiado grande para que se possa lidar com ela de maneira construtiva. Sem orientação adequada, o jovem nesse período pode ter problemas causados por envolvimentos precoces com sexo, álcool e drogas. É um péssimo período para o casamento e normalmente quando isso acontece, dura pouco. Também no lado profissional a pessoa tem dificuldade de se assentar, mudando continuamente de emprego ou atividade, que só terá término quando da entrada no segundo Ciclo.',
    6: 'O 6 (seis) indica infância e juventude restritiva, cheia de deveres e responsabilidades e, para fugir dessa restrição, normalmente casa-se cedo e muitas vezes esse casamento é um completo fracasso, pois não é escorado em bases sólidas do amor e sim como uma fuga. Tem, também, dificuldades em se ajustar à sociedade, pois é incompreendido nos seus planos e objetivos.',
    7: 'O 7 (sete) no primeiro Ciclo de Vida indica um período muito difícil. A criança e o jovem conservam-se retraídos e podem sofrer com a falta de compreensão dos pais, professores e amigos. Tal incompreensão leva, invariavelmente, ao isolamento, retraimento e até medo de encarar a vida nessa fase. Na faixa dos 20 anos, em virtude dessa retração, pode desenvolver complexos de culpa e falta de autoconfiança, restringindo o seu progresso pessoal e profissional.',
    8: 'O 8 (oito) no primeiro Ciclo de Vida indica um período de realizações. É extraordinário para o aprendizado acerca dos aspectos materiais da vida. É nesse período que se forjam os homens de negócios, comércio, políticos, advogados e todos aqueles que pensam mais no material do que no espiritual.',
    9: 'O 9 (nove) é o mais difícil dentre todos os Primeiros Ciclos. Quando criança existe boas oportunidades educacionais, mas também muita tensão. Quando adolescente, a mesma se sente confusa, assustada, nervosa e tem grandes dificuldades de concentração. O 9 é muito sensível, espiritualizado (por natureza) e normalmente incompreendido por todos que o cercam. Não deve se casar cedo, e caso o faça, normalmente esse enlace dura pouco. Profissionalmente tem grandes problemas de relacionamento e por isso, permanece inativo longos períodos, até adentrar ao segundo Ciclo, quando então, poderá pôr em prática toda a sua potencialidade. O 2 (dois) no primeiro Ciclo de Vida, indica uma criança extremamente mimada, que possivelmente sofreu grande influência da mãe ou dos avós. É natural que na adolescência, em vista da possessividade familiar, pense em casar-se o mais cedo possível e isso é muito comum, principalmente entre os homens.',
    11: 'O 11(onze) no primeiro Ciclo de Vida é um número demasiadamente enérgico e complicado para qualquer criança ou mesmo adolescente e aconselhamos a que seja reduzido a 2, onde poderá lidar melhor com ele.'
}
SegundoCiclo = {
    1: 'O 1 (um) no segundo Ciclo de Vida mostra um período de ambições, um grande desejo de realizações e também de sucesso relativo. A pessoa necessita desenvolver seus próprios recursos, estudando e se dedicando o máximo possível, além de lutar para tornar-se independente e chegar ao terceiro Ciclo já com definição profissional, social e financeira.',
    2: 'O 2 (dois) neste período é indicador de sociabilidade e receptividade. É necessário cultivar a paciência, o tato, a diplomacia e a capacidade de perceber os sentimentos alheios. Pode indicar ainda, uma carreira diplomática, ser juiz, médico, professor ou consultor.',
    3: 'O 3 (três) nos mostra uma fase agradável na vida, com certa despreocupação. É a fase da sociabilidade, na qual a criatividade e a originalidade podem exteriorizar suas ideias e sentimentos através de algum tipo de arte: pintura, música, teatro, escrita, etc. É um magnífico período para se desenvolver a criatividade, porém, não deve despender demasiada energia, principalmente em coisas fúteis.',
    4: 'O 4 (quatro) é sinônimo de trabalho duro, de produtividade e de construção do alicerce que deverá se apoiar no futuro. É um período em que a pessoa necessita aprender a aceitar a rotina e a trabalhar em algo produtivo, sólido e a fazer grande economia.',
    5: 'O 5 (cinco) é indicativo de um período de expansão de horizontes, época propícia a viagens, mudanças, romances, liberdade, de novas atividades e também novos amigos. Quase sempre, neste período, a pessoa terá de encontrar as suas oportunidades, longe do domicílio. Precisa aprender a se adaptar, a procurar novas maneiras de ver as coisas e a evitar a tendência para fixar-se num determinado lugar. Em resumo, é um período de grande movimentação, de grandes mudanças e de novos horizontes.',
    6: 'O 6 (seis) neste Ciclo nos mostra um período de ajustes e de responsabilidades nos assuntos domésticos em geral. É um bom momento para se contrair matrimônio, ter filhos e solidificar a família. Em suma, é um período familiar, de colocar a casa em ordem, de viver mais para a família, e deixar de ser tanto individualista.',
    7: 'O 7 (sete) indica um período de crescimento tranquilo, de estudos e de meditação. A menos que esteja casado, este não é um bom Ciclo para se contrair matrimônio, pois a pessoa necessita desenvolver seus recursos interiores e a incompreensão quase sempre aparece nesse período.',
    8: 'O 8 (oito) neste Ciclo mostra um período de preocupação com os aspectos materiais da vida. Normalmente a pessoa tem tendência a adquirir riqueza e poder material. Existe, ainda, a grande possibilidade de realizações no mundo dos negócios, a ganhar muito dinheiro com o trabalho e também através de especulações.',
    9: 'O 9 (nove) neste Ciclo traz a possibilidade de sucesso na vida pública. É um período altamente espiritual e a pessoa necessita aprender a cultivar a tolerância, o amor à humanidade, o altruísmo e o controle emocional. Dificilmente um romance é bem-sucedido e os casamentos tendem a pouca duração caso sejam realizados neste período e também é indício de    alguma perda, seja ela material, afetiva ou social.',
    11: 'O 11 (onze) nos mostra um período de ideais, de revelações, de grandeza e, possivelmente, de fama. Aconselha-se que a pessoa se mantenha longe de empreendimentos comerciais ou de especulações, sejam elas financeiras ou imobiliárias. É o momento de desenvolver a mente, de especializar-se em alguma coisa, de estudar, ensinar e também de inspirar as outras pessoas através do seu próprio exemplo.',
    22: 'O 22 (vinte e dois) no segundo Ciclo é indício de grandes realizações e de liderança em algum nível. O objetivo primordial da pessoa neste Ciclo deve ser o de beneficiar a humanidade como um todo. Em virtude do grande poder deste número, os nervos e as emoções serão testados durante todo o período e a pessoa deve se manter o mais calmo possível e seguir a orientação de sua intuição.',
}
TerceiroCiclo = {
    1: 'O 1 (um) nos indica um final de vida solitário. A pessoa precisa permanecer ativa e independente e contar com seus próprios recursos.',
    2: 'O 2 (dois) mostra um período de amor sincero e de amigos íntimos. A pessoa se sentirá impelida a colecionar coisas, tais como selos, moedas, antiguidades ou qualquer coisa extravagante.',
    3: 'O 3 (três) no terceiro Ciclo de Vida indica um período de expressão de ideias e sentimentos através de diversas formas de arte, música, teatro e literatura. A criatividade vai se desenvolver. Haverá muitos amigos e grande atividade social.',
    4: 'O 4 (quatro) neste Ciclo nos mostra que a pessoa, mesmo aposentada, deverá continuar trabalhando, seja por necessidade, seja por escolha, pois o 4 não o deixará levar uma vida monótona e rotineira.',
    5: 'O 5 (cinco) é o período da liberdade pessoal, de viagens, mudanças, de novas atividades e variedade, seja de amigos, de atividades ou de residência.',
    6: 'O 6 (seis) poderá ser o mais agradável de todos os terceiros ciclos de vida uma fase de felicidade e harmonia no lar se a pessoa tiver aprendido a adaptar-se e assumir responsabilidades. Caso não tenha aprendido estas coisas, ela poderá ser sobrecarregada com muitos problemas domésticos.',
    7: 'O 7 (sete) indica um período de isolamento ou de semi-isolamento. Trata-se de uma fase tranquila, apropriada para se estudar em casa e adquirir sabedoria e conhecimento.',
    8: 'O 8 (oito) neste Ciclo mostra que a pessoa precisa agir com sabedoria, trabalhar e estudar duramente nos dois primeiros quando terá grande possibilidade de ficar rico neste e ter poder e sucesso ilimitados no mundo dos negócios.',
    9: 'O 9 (nove) mostra um período de retiro para o estudo e   o aprendizado. A pessoa precisa cultivar a tolerância e o amor pela humanidade. Neste Ciclo geralmente há alguma espécie de perda.',
    11: 'O 11 (onze) é o período de isolamento, de inspiração, de leitura, de alguma incursão na arte de escrever e, possivelmente, de fama.',
    22: 'O 22 (vinte e dois) no terceiro ciclo de vida talvez torne a pessoa tensa e nervosa. Ela deve procurar manter-se ativa durante esse período e dedicar-se a um hobby, tal como a escultura, a pintura, as artes divinatórias, etc.'

}
Desafios = {
    1: 'Desafio 1 – O consulente precisará aprender a se situar num meio termo entre um sentimento excessivo ou insuficiente de sua própria personalidade ou importância. Precisa aprender a ser firme, positivo independente e autoconfiante, sem impor sua vontade às outras pessoas ou esperar que tudo gire em torno de si.',
    2: 'Desafio 2 – Poderá tender a ser tão sensível em relação aos seus próprios sentimentos e a passar tanto tempo pensando neles, que acabará não tomando conhecimento dos sentimentos dos outros. Pequenas coisas são ampliadas fora de qualquer proporção e nunca esquecidas ou perdoadas. O consulente precisa aprender a cuidar de si mesmo, a cultivar uma atitude mais liberal e tolerante em relação à vida e a parar de utilizar seus próprios sentimentos e emoções como ponto de referência para tudo.',
    3: 'Desafio 3 – Precisará aprender a situar-se num meio termo, entre ter medo de contatos sociais e ser por demais festeiro. Tem de aprender a ser sociável e a exprimir suas ideias e sentimentos sem dispersar suas energias ou comportar-se como pessoa fútil.',
    4: 'Desafio 4 – É o mais fácil de todos os desafios, visto que não há nenhum conflito envolvido. Precisa aprender a situar-se num meio termo entre agir como um “burro de carga” ou ser preguiçoso.',
    5: 'Desafio 5 – Precisa aprender a situar-se num meio termo entre desejar uma liberdade excessiva e ter um receio injustiçado dela entre uma ânsia exagerada de experiências sensuais e o medo de tentar coisas novas. Precisa aprender a não buscar sexo, álcool e drogas e o mais difícil de tudo precisa aprender quando e como renunciar a pessoas ou coisas cuja presença na sua vida não tem mais razão de ser.',
    6: 'Desafio 6 – Precisa aprender a situar-se num meio termo entre comportar-se como um “capacho” ou ser demasiado exigente e dominador. Precisa aprender a aceitar as pessoas como elas são sem esperar que elas vivam de acordo com os seus padrões; respeitar os pontos de vista de todos e não estabelecer regras além de você mesmo.',
    7: 'Desafio 7 – Precisará aprender a situar-se num meio termo entre o orgulho excessivo e a modéstia exagerada. Deveria tomar cuidado para não se refugiar dentro de si mesmo e nem tentar escapar das coisas desagradáveis da vida, recorrendo ao álcool e às drogas. É particularmente uma boa educação, aprender a compreender o que se passa no mundo à sua volta e acima de tudo ter fé.',
    8: 'Desafio 8 – Precisará aprender a situar-se num meio termo entre uma preocupação excessiva com as questões materiais, e um desinteresse exagerado em relação a esse assunto. Precisa aprender a utilizar corretamente o dinheiro e o poder e a voltar seu pensamento para outras coisas que não o dinheiro e o que ele poderá fazer por você.',
    0: 'Desafio 0 (zero) – É o desafio da escolha. Caso tenha um desafio 0, poderá ser altamente evoluído e terá de tomar suas próprias decisões. Deve atentar para todos os desafios sem uma ênfase especial em qualquer deles. Espera-se que o consulente decida por si mesmo quais as armadilhas que a vida lhe reserva. O Desafio Zero, não é, como alguns numerólogos supõem, negativo e o consulente castigado por algum devaneio em vidas passadas ou mesmo nesta existência. Muito pelo contrário; o consulente que em seu Mapa Numerologico possui o Desafio Zero é possuidor de dons e atributos que aos outros não é dado esse privilégio, ou seja, de poder “escolher” o Desafio que melhor se adapta à sua existência, ou não ter Desafio algum a vencer. Claro, aparentemente é mais fácil ter de enfrentar um só problema do que nove, porém é bom não esquecer que a Natureza jamais dá ao ser humano fardo maior do que ele pode suportar. Assim sendo, a pessoa com esse Desafio está vibrando em todas as esferas e o “dever aprender” passa a ser uma “obrigação”, um “dever” do mesmo, visto que muitas coisas especiais e elevadas lhe estão sendo projetadas e ele terá, obrigatoriamente, de cumprir com o que é determinado, pois que a humanidade necessita de seus préstimos.'

}
RelaçõesInterValores = {
    0: 'Como já foi dito, muitos números acompanham o ser humano durante toda a sua vida. Alguns são harmônicos, outros conflitantes e, outros ainda, denotam certas capacidades e aptidões especiais. Como tudo na Natureza, a harmonia deve imperar também nos números, e qualquer tendência à desarmonia deve ser analisada mais profundamente. Sendo assim, podemos deduzir que tais excessos indicam talentos e deficiências naturais da pessoa em questão.',
    1: 'RI-1– Indica independência, ambição e interesses próprios. Também mostra egoísmo e possessividade.',
    2: 'RI-2 – Indica que a pessoa é dotada de tato e diplomacia, possui grande amor à música e às artes de um modo geral. É harmônico e tem capacidade de cooperação. Por vezes, também indica insegurança e timidez.',
    3: 'RI-3 – Indica pessoa de rara capacidade de expressão, forte imaginação e senso de humor. Às vezes é sinal de irresponsabilidade e impaciência a uma atitude realista ou materialista.',
    4: 'RI-4 – Indica que a pessoa é econômica, honesta e tem tendência para o trabalho árduo. Porém, tem carência de concentração e julgamento imparcial; possibilidade de obstinação.',
    5: 'RI-5 – Indica pessoa impulsiva e nervosa, com grande desejo de sexo. As viagens e as mudanças lhe são altamente favoráveis.',
    6: 'RI-6 – Indica capacidade de assumir grandes responsabilidades. É de confiança, caseiro, pai e educador nato. Tem tendência a polêmicas, brigas e instabilidade emocional.',
    7: 'RI-7 – Indica poder de análise, agilidade mental, perfeccionismo, equilíbrio e cultura. Grande inclinação pelos assuntos metafísicos e a se retrair.',
    8: 'RI-8 – Indica capacidade para negócios, habilidade executiva, liderança, iniciativa, tato e grande senso de valores materiais. Tem tendência a se mostrar como dono da verdade.',
    9: 'RI-9 – Indica um modo de ver universal. Tem dons artísticos e literários. Adora viajar. Em muitos casos, também indica visão estreita e egocentrismo, ou demasiado desapego e afasta- mento da realidade.'
}
MomentoDecisivo = {0: '',
                   1: 'MOMENTO DECISIVO 1 – Não é um período fácil; exige coragem, determinação e muita força de vontade. É o momento propício para se “cultivar” a individualidade, a independência e a engenhosidade. Inúmeros acasos e situações inesperadas forçarão a pessoa a enfrentar a vida pensando e agindo por si mesma. Um Momento Decisivo 1 no primeiro Ciclo de Vida, indica uma criança agitada, voluntariosa e por vezes complicada, que será difícil controlar e compreender.',
                   2: 'MOMENTO DECISIVO 2 – Traz consigo a oportunidade para “cultivar” o tato e a compreensão. Se for amigo, companheiro e atencioso com seus semelhantes, este será um período de amizades sinceras e de relacionamentos duradouros. Excelente fase para se contrair matrimônio. Se for impaciente e desatencioso, poderá ser uma fase de relacionamentos difíceis, de grandes incompreensões, brigas, discussões, em que você poderá causar graves prejuízos às pessoas que o cercam. Um Momento Decisivo 2 no primeiro Ciclo de Vida, é indício de uma “mãe” forte e dominadora, ou pai ausente (por motivo de trabalho, morte ou separação). A criança, nesse caso, pode se tornar excessivamente sensível e ter reflexos dessa sensibilidade na juventude e adolescência, obstruindo dessa maneira, as possibilidades de progresso.',
                   3: 'MOMENTO DECISIVO 3 – É o momento de expandir a vida social e “cultivar” os próprios talentos. Trata-se de uma fase apropriada para a auto expressão, novas amizades, romance e fertilidade. A manifestação descuidada das emoções poderá ter consequências desagradáveis, pois existe, nesse estado, tendência ao desmando: vícios, brigas, discórdias. Cuidado com os “amigos”, pois apesar de serem necessários, por vezes são más companhias. Um Momento Decisivo 3 no primeiro Ciclo de Vida, geralmente indica uma criança com dificuldade de se adaptar aos estudos. Indica, também, oportunidades artísticas que se não alimentadas e direcionadas condizentemente, poderão ser desperdiçadas, fazendo com que a pessoa já adulta venha a se lamentar dessa negligência dos pais ou educadores.',
                   4: 'MOMENTO DECISIVO 4 – Este Momento Decisivo traz a oportunidade de se construir um sólido alicerce para o futuro. É um período de trabalho duro e até de algumas restrições e é necessário “cultivar” a paciência e os bons hábitos de trabalho. Neste período, poderá haver alguns problemas econômicos, que serão superados com inteligência, trabalho e dedicação ao projeto final. A família e os parentes por afinidade podem se transformar num peso e a pessoa terá de ajudá-los, tanto financeiramente, como prestando ajuda humanitária, em uma doença, por exemplo. As recompensas sempre aparecem a partir da aplicação dos preceitos corretos de vida e do esforço para se obter os resultados positivos. Um Momento Decisivo 4 no primeiro Ciclo de Vida, frequentemente indica que a pessoa poderá começar a trabalhar muito nova e a assumir grandes responsabilidades ainda na juventude.',
                   5: 'MOMENTO DECISIVO 5 – Traz oportunidades para viagens, para experimentar novas sensações, novos empreendimentos e para se livrar de tudo que está obsoleto ou que já não nos faz falta. É uma fase de liberdade, de mudanças e de desenvolvimento pessoal, principalmente se vier após um Momento Decisivo 4 ou 6. Um Momento Decisivo 5 no primeiro Ciclo de Vida, indica uma criança ousada, inquieta, esperta e pouco constante. Geralmente empreende mudanças súbitas, ora gostando disto, ora daquilo, sem esperar as recompensas resultantes de um esforço ou trabalho empreendido.',
                   6: 'MOMENTO DECISIVO 6 – É o momento dos ajustes e das responsabilidades familiares. Caso tenha consciência disso, este é o Momento de grande afetividade, de amor e de felicidade doméstica, além de sucesso e segurança material. Do contrário, ou seja, caso seja dispersivo ou inconstante, poderá ser um período de desgostos, discussões, brigas e graves problemas domésticos e até indício de separação. Um Momento Decisivo 6 no primeiro Ciclo de Vida, geralmente indica casamento precoce ou a responsabilidade de tomar conta dos pais ou de algum familiar. Quando o 6 for o último Momento Decisivo, ele poderá trazer o reconhecimento do trabalho já feito. Caso a pessoa esteja solteira, este Momento trará a oportunidade para um novo amor e para o materialismo.',
                   7: 'MOMENTO DECISIVO 7 – É uma fase de introspecção, de meditação e estudo do significado último da vida. Caso não esteja casado, desaconselhamos o matrimônio nesta fase. Velhos relacionamentos que não produzem mais frutos, podem e devem ser deixados para trás. A pessoa normalmente sente vontade de se retirar para dentro de si mesma, o que de certa forma poderá causar problemas de relacionamento, tanto a nível pessoal como familiar. Um Momento Decisivo 7 no primeiro Ciclo de Vida, nos indica uma criança retraída, solitária, pensativa e muito estudiosa. Quando os pais são excessivamente rígidos e severos, a criança poderá, pela regressão de suas ideias e projetos, contrair algum tipo de doença psicossomática ou mesmo depressão, ser temperamental e desenvolver algum tipo de complexo.',
                   8: 'MOMENTO DECISIVO 8 – É um período de grandes realizações no mundo dos negócios. As despesas são altas, não obstante, é uma excelente fase para se correr ao encontro dos objetivos, de conquistar poder, fama e sucesso material. Com dedicação, estudo e trabalho sistemático, com objetivo definido e com colaboradores aptos e interessados, a pessoa dificilmente deixa de conseguir tudo o que deseja. Um Momento Decisivo 8 no primeiro Ciclo de Vida, indica que a pessoa começará ainda jovem a se dedicar aos negócios, a trabalhar para se sustentar e também sustentar algum membro da família.',
                   9: 'MOMENTO DECISIVO 9 – Traz a oportunidade para se “cultivar” o amor, a solidariedade, o altruísmo e para se viajar para o exterior. Poderá haver algum tipo de perda e até desapontamentos, principalmente entre amigos. Um bom investimento para o consulente é fazer obras humanitárias durante este período, pois os frutos dessa plantação são certos, e o sucesso e a fama se farão presentes. Um Momento Decisivo 9 no primeiro Ciclo de Vida, normalmente não é dos mais afortunados, pois quase sempre a criança é incompreendida por colegas, amigos e até familiares, que por causa dessa incompreensão exigem muito e retribuem pouco, o que faz com que o jovem se retraia e fique tímido e introspectivo.',
                   11: 'MOMENTO DECISIVO 11 – Por ser um número altamente espiritual e elevado, a pessoa nesse período sente-se tensa e muito nervosa. É uma excelente fase para estudar esoterismo, espiritualismo e expandir seus horizontes. Este momento traz inspiração, iluminação e, quase sempre, fama e prestígio nacional e até internacional. Não faça nada nem diga por trás o que não teria coragem de dizer ou fazer na frente das pessoas.',
                   22: 'MOMENTO DECISIVO 22 – É, sem dúvida alguma, o número e o Momento mais poderoso. A pessoa fica altamente criativa e neste estado tornam-se possíveis todas as realizações. É uma fase de interesses pelos problemas mundiais e de grande expansão da consciência.'
                   }
Amor = {
    0: 'O povo hebreu sempre deu grande importância aos relacionamentos. Como se sabe, o divórcio não é bem visto entre os judeus e o convívio a dois, sagrado, é um dos princípios fundamentais e também a base da Cabala. ',
    1: 'NÚMERO UM -Este número é o do idílio (amor poético), mas a atração deve ter uma base intelectual. Como este número também estimula a variedade, não será muito fácil chegar-se a uma união duradoura, pois haverá a tentação de se tratar de duas coisas ao mesmo tempo. Além disso, a tendência intelectual prejudicará o fogo da paixão, que você procura em vão. Os casamentos podem ocorrer de repente, e também podem ocorrer encontros quando você estiver fazendo uma viagem bem longe de casa ou em um centro de estudos. O possuidor deste número deve ser sempre cauteloso na vida matrimonial, pois pode haver o perigo de aventuras extraconjugais provocadas pelo tédio ou pela falta de liberdade mental.',
    2: 'NÚMERO DOIS - Este número é discriminador na escolha de um cônjuge, principalmente por causa do seu acentuado interesse pelo conforto e pela estabilidade. A sugestão pode desempenhar um papel importante na sua vida e deve ter cuidado para impedir críticas indevidas por parte de parentes ou amigos, ou seja, interferência na sua vida amorosa. Pode haver também amizades secretas, de natureza puramente platônica, que devem ser encaradas com cautela, a fim de serem evitadas desavenças conjugais. Você há de querer uma pessoa inteligente, mas que seja acima de tudo prática e capaz de garantir a segurança financeira que você deseja.',
    3: 'NÚMERO TRÊS - Às vezes você tem a infelicidade de conhecer seus parceiros ideais demasiadamente tarde, quando normalmente já estão casados e, se não tiver cuidado, pode tornar-se um lado de um triângulo, em vez de esperar outro parceiro que esteja livre. Você é muito idealista para compartilhar seu casamento, mas o seu espírito de sacrifício pode ser explorado maldosamente por aqueles em quem confia. Possivelmente, haverá um ou mais de um quase erro, porém você tem a proteção que impede de se prender a uma pessoa de gênio insuportável ou de mau caráter. Se deixar agir a sua intuição, a sua escolha será correta, pois você compreende muito bem as motivações das pessoas e, em via de regra, escolherá alguém cujos interesses combinem com os seus, encontrando então o companheiro ideal para toda a vida.',
    4: 'NÚMERO QUATRO - Você sentirá uma forte atração pelo casamento, devido à sua natureza emocional e afetiva, mas nem sempre terá a necessária discrição, e a sua escolha lhe poderá acarretar dificuldades. Seu caminho será ainda mais dificultado pela inveja e despeito de outras pessoas, e também por sua tendência de julgar excessivamente os outros pela aparência. No casamento, embora procure um parceiro ativo e dominador, poderá haver atritos conjugais, com certa frequência, se houver muitas exigências de sua parte. Seu parceiro terá interesse pelo ocultismo, e poderá fazer com que você se interesse por sua superioridade intelectual. ',
    5: 'NÚMERO CINCO - Este número acarreta viagens, mudanças e variedades. É quase certo que a pessoa se casará e se separará mais de uma vez, devido à inquietação e gosto pela mudança. Por outro lado, os dois cônjuges, mesmo separados, podem e normalmente mantêm negócios juntos, pois existe uma tendência à amizade. Não costuma ser muito seletivo em seus relacionamentos, principalmente até os 25 anos e com isso vive trocando de parceiro constantemente, sempre à procura do “ideal”. Obstáculos e atrasos serão comuns na juventude, e o aconselhável é o casamento depois de trinta anos, e com uma pessoa da mesma idade ou até um pouco mais velha. ',
    6: 'NÚMERO SEIS - Não sendo muito precoces, as pessoas possuidoras deste número, embora ardentes no íntimo, são tímidas e pouco expressivas. No entanto, experimentam, provocados pelo ardor interior, períodos de exaltação, e se não houver cuidado, tais ardores poderão acarretar a formação de alianças de natureza não muito satisfatória que, se terminarem em casamento, podem provocar infelicidade. É um número não muito forte, apresentando certa vacilação e variedade, e as pessoas regidas por ele devem ter muito cuidado para não desejarem novas companhias capazes de prejudicarem o estado conjugal, em particular quando elas se encontram afastadas do outro cônjuge. Se o casamento ocorrer depois de 33 anos de idade, haverá mais esperança de seu sucesso e estabilidade. ',
    7: 'NÚMERO SETE - As pessoas regidas por este número são muitas vezes consideradas namoradeiras e até volúveis. Não é verdade. Embora não devam se casar muito jovens, ou seja, antes que o seu caráter e os seus gostos estejam bem pronunciados, quase sempre isso acontece e, em consequência, normalmente esses relacionamentos duram pouco. São tão ativas, quer mental, quer fisicamente, que a sua existência é muito variada e costumam conhecer um número muito maior de pessoas do que aqueles regidos pelos outros números, com exceção do cinco. Frequentemente se casa mais de uma vez, e têm muitos casos amorosos, sem que procure aventuras, principalmente durante a vigência do casamento. Haverá muitas separações de pessoas amadas, amigas, e até mesmo mortes prematuras de pessoas queridas. Em via de regra têm muita responsabilidade desde a juventude. ',
    8: 'NÚMERO OITO - As pessoas governadas por este número se casam para toda a vida, pois são sinceras e honradas, mas curiosamente, costumam ocorrer muitas vezes idílios desfeitos antes do casamento, por falta de sinceridade da outra parte. Se sofrer tal coisa, a pessoa pode custar a se recuperar e, se enviuvar depois de um casamento feliz, a pessoa regida por este número raramente se casa de novo. A religião e a filosofa desempenham um papel importante em sua vida e embora o impulso sexual seja forte, há um sentimento de escrúpulo que evita a indulgência com o sexo antes do casamento. O companheirismo é uma questão de importância vital para essas pessoas e, como o namoro, geralmente conduzido com muito critério e bom senso, o desempenho de atividades conjuntas, fará feliz e duradouro o casamento ',
    9: 'NÚMERO NOVE - O casamento será a aspiração suprema e constante, e você a ele sacrificará todos os seus pensamentos, esperanças e aspirações, desde muito jovem. Na juventude, a tendência é de gostar e até namorar com pessoas bem mais velhas do que você, o que poderá lhe acarretar alguns problemas, particularmente quando, com o correr dos tempos, você passe a se interessar por pessoas mais jovens. Deve tomar cuidado para que os seus objetivos e atividades profissionais sejam compatíveis com os do seu cônjuge, pois, do contrário, terá muitas decepções e aborrecimentos. '
}
AConseq1ano = {
    1: 'Plantando as sementes.',
    2: 'As sementes criam raízes.',
    3: 'surgem os primeiros brotos.',
    4: 'Cavanto e capinando.',
    5: 'Formam-se os botões.',
    6: 'Floração',
    7: 'As plantas dão frutos',
    8: 'Época da colheita.',
    9: 'Época de limpar e aterrar após a colheita e prepará-la para um novo plantio.'
}
AConseq2objetivo = {
    1: 'Ano para começar coisas novas',
    2: 'Ano para agir com discrição e ser paciente e receptivo',
    3: 'Ano de crescimento pessoal e cultivo de novas amizades',
    4: 'Ano de restrições, trabalho duro e autodisciplina',
    5: 'Ano para abandonar-se aos impulsos da vida e viver o presente.',
    6: 'Ano para dedicação altruísta à família e à comunidade',
    7: 'Ano para isolamento, descanso e aperfeiçoamento interior',
    8: 'Ano dinâmico e materialista Hora de pagar e cobrar dívidas.',
    9: 'Ano de faxina entre o fim de um ciclo e o começo do próximo'
}
AConseq3oquedevefazer = {
    1: 'Ser independente, criativo, seguro, seletivo e seguir a própria intuição.',
    2: 'Ser diplomata, delicado e cooperativo no trato com as outras pessoas.',
    3: 'Dar vazão à criatividade, evitar extravagâncias e dispersão de energia.',
    4: 'Ser autodisciplinado, metódico e dar forma concreta às ideias.',
    5: 'Desenvolver a própria personalidade, tornar-se adaptável e aproveitar as oportunidades.',
    6: 'Amar e dedicar-se mais à família em todos os seus aspectos.',
    7: 'Estudar os significados últimos da vida e evitar atividades materialistas.',
    8: 'Ter coragem de ousar grandes feitos, usar o bom senso, preocupar-se com o dinheiro e ser organizado.',
    9: 'Excelente fase para escrever, representar, viajar e para dedicar-se a estudos metafísicos.'
}
AConseq4perigo = {
    1: 'Falta de iniciativa que poderá influenciar em todo ciclo de 9 anos.',
    2: 'Propensão para discussões ou ser excessivamente atrevido.',
    3: 'Perda de grandes oportunidades em razão da dispersão das energias.',
    4: 'Negligência com a saúde e acomodamento profissional.',
    5: 'Dispersão de energias, excesso de atividades sexuais e não acompanhar o ritmo normal da vida.',
    6: 'Ser excessivamente idealista e esperar demasiado dos outros.',
    7: 'Negligenciar a saúde, forçar decisões, ser crítico e permitir que complexos aflorem.',
    8: 'Ser emotivo, sentimental e gastar mais do que ganha.',
    9: 'Ser ciumento e possessivo.'
}


########################################################################################################################
# Funções para impressão em word
def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row
#######################################################################################################################
# Funções para calcular os valores


def calcular(palavra):
    # essa beleza aqui reduz a palavra em número. Mas ainda não é o número final,
    #  pois como temos que verificar os 11, 22 e 33 ele reduz em 3 ou 2 algarismos.
    palavra = list(palavra)
    total = 0
    for letra in palavra:
        value = cabalacod[letra]
        total = total + value
    return (total)
def reduzirmes(numero):
    # este aqui pega o numero formado das letras e reduz para um algarismo ou para 11 e 22.
    if numero == 11:  # números 11, 22 não são reduzidos.
        total = numero
    else:
        total = numero
        tamanho = len(str(total))
        if tamanho > 1:
            while tamanho > 1:  # repetindo até o número seja reduzido em um dígito.
                if total == 11:
                    break
                palavranum = str(total)
                palavranum = list(palavranum)
                total = 0
                for letra in palavranum:
                    total = total + int(letra)
                tamanho = len(str(total))
        else:
            total = numero
    return (total)
def reduzir(numero):
    # este aqui pega o numero formado das letras e reduz para um algarismo ou para 11 e 22.
    if numero == 11 or numero == 22:  # números 11, 22 não são reduzidos.
        total = numero
    else:
        total = numero
        tamanho = len(str(total))
        if tamanho > 1:
            while tamanho > 1:  # repetindo até o número seja reduzido em um dígito.
                if total == 11 or total == 22:
                    break
                palavranum = str(total)
                palavranum = list(palavranum)
                total = 0
                for letra in palavranum:
                    total = total + int(letra)
                tamanho = len(str(total))
        else:
            total = numero
    return (total)
def reduzirtotal(numero):
    # este aqui pega o numero formado das letras e reduz para um algarismo ou para 11 e 22.
    total = numero
    tamanho = len(str(total))
    if tamanho > 1:
        while tamanho > 1:  # repetindo até o número seja reduzido em um dígito.
            palavranum = str(total)
            palavranum = list(palavranum)
            total = 0
            for letra in palavranum:
                total = total + int(letra)
            tamanho = len(str(total))
    else:
        total = numero
    return (total)
def reduzirdata(numero):
    # Lindo, reduzimos aqui a data em um número de até duas casas decimais.
    dia, mes, ano = numero.split('/')
    datanasc = numero.replace("/", "")
    # este aqui pega o numero formado das letras e reduz para um algarismo ou para 11 e 22.
    if numero == 11 or numero == 22:  # números 11, 22 não são reduzidos.
        total = numero
    else:
        total = numero
        tamanho = len(str(total))
        if tamanho > 1:
            while tamanho > 1:  # repetindo até o número seja reduzido em um dígito.
                if total == 11 or total == 22:
                    break
                else:
                    total == 2 or total == 4
                    break
                palavranum = str(total)
                palavranum = list(palavranum)
                total = 0
                for letra in palavranum:
                    total = total + int(letra)
                tamanho = len(str(total))
        else:
            total = numero
    total = 0
    dia1 = reduzir(int(dia))
    mes1 = reduzir(int(mes))
    ano1 = reduzir(int(ano))
    total = dia1 + mes1 + ano1
    return (total)
def letnum(palavra):
    # essa beleza aqui encontra os valores das letras da tabela de cabala.
    palavra = list(palavra)
    letranum = list()
    for letra in palavra:
        value = cabalacod[letra]
        letranum.append(value)

    return (letranum)
def listanum(listaent):
    lista = list(listaent)
    tabelaletnum = list()
    pri = seg = ter = quar = quin = sext = seti = oita = nona = 0
    for c, v in enumerate(lista):
        if v == 1:
            pri += 1
        elif v == 2:
            seg += 1
        elif v == 3:
            ter += 1
        elif v == 4:
            quar += 1
        elif v == 5:
            quin += 1
        elif v == 6:
            sext += 1
        elif v == 7:
            seti += 1
        elif v == 8:
            oita += 1
        elif v == 9:
            nona += 1
    tabelaletnum.insert(0, pri)
    tabelaletnum.insert(1, seg)
    tabelaletnum.insert(2, ter)
    tabelaletnum.insert(3, quar)
    tabelaletnum.insert(4, quin)
    tabelaletnum.insert(5, sext)
    tabelaletnum.insert(6, seti)
    tabelaletnum.insert(7, oita)
    tabelaletnum.insert(8, nona)
    return (tabelaletnum)
def retornacarma(lista):
    # vai adicionando os valores do carma em uma lista de carmas
    lista = list(lista)
    listacarma = list()
    for c in enumerate(lista):
        aux = c[1]
        if aux == 0:
            valor = c[0] + 1
            listacarma.append(valor)
    return (listacarma)
def tendenciaoculta(lista):
    lista = list(lista)
    tendenocult = list()
    for c in enumerate(lista):
        if c[1] > 3:
            aux = c[0] + 1
            tendenocult.append(aux)
    return (tendenocult)
def compararlistas(desafio, comparacoes):
    # a função começa aqui e a entrada são as duas listas preenchidas a primeiro de desafio a segunda dos comparadores
    doencas = list()
    for i in desafio:
        for c in comparacoes:
            if i == c:
                valor = c
                doencas.append(valor)
    return (doencas)
def relacinterval(listareznum):
    relinval = list()
    m = max(listareznum)
    for i, j in enumerate(listareznum):
        if j == m:
            relinval.append(int(i) + 1)

    return (relinval)
def tabelaano(dia, mes, ano, ):
    # essa def aqui pega a data de nascimento e coloca dentro de um dicionário, o índice é o ano,e o valor é o resultado numerologico
    anolist = list()
    tabano = dict()
    anox = int(ano)
    y = 0
    while y < 117:
        y += 1
        anolist.append(anox)
        anox += 1
    for i in anolist:
        anopessoal = '%s/%s/%s' % (dia, mes, i)
        anopessoal = reduzir(reduzirdata(anopessoal))
        tabano.update({int(i): int(reduzirtotal(reduzir(anopessoal)))})
    return (tabano)
def contarep(listaaver, indice):
    N = int(0)
    for k in range(0, 9):
        if k != indice:
            if listaaver[k] == listaaver[indice]:
                N = N + 1
    return N
def lisPosMaiores(listaaver):
    maior = listaaver[0]
    lista = list()
    for k in range(1, 9):
        if maior < listaaver[k]:
            maior = listaaver[k]
    for k in range(0, 9):
        if listaaver[k] == maior:
            lista.append(int(k))
    return lista
def retiradup(listasaida):
    listaret = list()
    for cada in listasaida:
        if cada not in listaret:
            listaret.append(cada)
    return listaret
def relacinterval(listareznum):
    listapos = list()
    listamaiores = lisPosMaiores(listareznum)

    for cada in listamaiores:
        for j in range(0, 9):
            if cada != j:
                N = contarep(listareznum, j)

                if listareznum[cada] >= 2 * N and listareznum[cada] != 0:
                    listapos.append(int(cada + 1))
    listapos = retiradup(listapos)
    return (listapos)
def matrixnum(palavra):
    palavra = list(palavra)
    numeros = list()
    for letra in palavra:
        value = cabalacod[letra]
        numeros.append(value)
    return (numeros)
def giromatrix(matrixnum):
    temp = list()
    cont = 0
    while cont < (len(matrixnum)-1):
        value = matrixnum[cont] + matrixnum[cont + 1]
        if len(matrixnum) == 2:
            value = reduzir(value)
            temp.append(value)
            break

        value = reduzirtotal(value)
        temp.append(value)
        cont += 1
    return (temp)
def piramidnum(matrix):
    cont = 0
    temp = list()
    resultado = list()
    matrixtemp = giromatrix(matrix)
    while cont <= len(matrix):
        temp.append(matrixtemp)
        resultado.append((temp[:]))
        matrixtemp = giromatrix(matrixtemp)
        temp.clear()
        cont += 1
    resultado.pop()
    resultado.pop()
    return(resultado)





extenso = str(input(' Digite seu nome completo :'))  # ================ Inserir dados
#extenso = 'Rodrigo Rodrigues e Silva'
nome = extenso
extenso = extenso.upper()
nomecompleto = extenso.replace(' ', '')
extenso = extenso.split(" ")
primeironome = extenso[0]
namelist = list()
namelist = nomecompleto

nascimento = input(' Digite sua data de nascimento (dd/mm/aaaa) :')  # ================ Inserir dados
#nascimento = "22/04/1986"
datadenascimento = nascimento
dia, mes, ano = nascimento.split('/')
datanasc = nascimento.replace("/", "")

# Pegando vogais e consoantes
vogaisnome = re.findall(r"[AIYÕÊÁÍÛÃEÔUOÉÂÎÚÓ]", primeironome)
vogaiscompleto = re.findall(r"[AIYÕÊÁÍÛÃEÔUOÉÂÎÚÓ]", nomecompleto)

consnome = re.findall(r"[QJBKR'CGLSDMTHNVWXÇZFP]", primeironome)
conscompleto = re.findall(r"[QJBKR'CGLSDMTHNVWXÇZFP]", nomecompleto)

# numerologia do primeiro nome
numerolpnome = reduzir(calcular(primeironome))

# calculando o número de impressão ( soma das consoantes)
impress = reduzir(calcular(conscompleto))
aparencia = reduzirtotal(impress)

# calculando o numero de expressão ( reduzir nome inteiro)
reqexpr = calcular(nomecompleto)
if reqexpr == 20 or reqexpr == 31 or reqexpr == 40:
    contecons = 0
    reqexpr = 0
    while contecons < len(extenso):
        reqexpr = reqexpr + reduzir(extenso(contecons))
        contecons += 1
    else:
        expressao = reduzir(reqexpr)
else:
    expressao = reduzir(calcular(nomecompleto))

# calculando número de motivação ( soma das vogais)
# Sempre que a soma dos números das vogáis que correspondem ao número de motivação for, por, exemplo 20, 31 ou 40, eles podem ou não se transformar em 2 ou 4 dependendo da somatória de cada nome em separado.
reqmotiv = calcular(vogaiscompleto)
if reqmotiv == 20 or reqmotiv == 31 or reqmotiv == 40:
    contvog = 0
    reqmotiv = 0
    while contvog < len(extenso):
        wogel = re.findall(r'[AIYÕÊÁÍÛÃEÔUOÉÂÎÚÓ]', extenso[contvog])
        reqmotiv = reqmotiv + reduzir(calcular(wogel))
        contvog += 1
    else:
        motivac = reduzir(reqmotiv)

else:
    motivac = reduzir(calcular(vogaiscompleto))

# calculando o número do destino ( soma de todos os números da data de nascimento)
destinonum = int(reduzirdata(nascimento))
destinonum = int(reduzir(destinonum))

# Lições cármicas
licaocarm = list()
letranumero = letnum(nomecompleto)
listanumsnome = listanum(letranumero)
licaocarm = retornacarma(listanumsnome)

# TENDÊNCIAS OCULTAS
tendenciasocultas = list()
tendocult = tendenciaoculta(listanumsnome)

# RELAÇÕES INTERVALORES

ltno = letnum(primeironome)
lnn = listanum(ltno)
relinterval = relacinterval(lnn)

# RESPOSTA SUBCONSCIENTE
respsub = 9 - len(licaocarm)

# Dividas cármicas, para se saber se uma pessoa carrega Dívidas Cármicas, deve-se em principio observar o doa do seu nascimento:

divida = list()
# dívidas de acordo com os dias
if int(dia) == 13:
    divida.append(13)
elif int(dia) == 14:
    divida.append(14)
elif int(dia) == 16:
    divida.append(16)
elif int(dia) == 19:
    divida.append(19)
# Dívidas de acordo com destino, motivação e expressão

if destinonum == 4:
    divida.append(13)
elif motivac == 4:
    divida.append(13)
elif expressao == 4:
    divida.append(13)

if destinonum == 5:
    divida.append(14)

elif motivac == 5:
    divida.append(14)

elif expressao == 5:
    divida.append(14)

if destinonum == 7:
    divida.append(16)
elif motivac == 7:
    divida.append(16)
elif expressao == 7:
    divida.append(16)

if destinonum == 1:
    divida.append(19)
elif motivac == 1:
    divida.append(19)
elif expressao == 1:
    divida.append(19)

# Número da missão : o número de destino mais o número de Expressão somados e reduzidos.
# observar os números 11 e 22 não se reduzem.
missao = destinonum + expressao
missao = reduzir(missao)

# Ano pessoal : soma do dia, do mês de nascimento e do ano em curso no momento em que estiver fazendo o cálculo.
# calculo para aniversário, caso antes reduz em 1 o ano da analise numerológica.
if int(mes) >= mesatual:
    if int(dia) < diaatual:
        anoatual = anoatual - 1

anopessoal = '%s/%s/%s' % (dia, mes, anoatual)
anopessoal = reduzirtotal(reduzirdata(anopessoal))

# Mês pessoal
mespessoal = mesatual + anopessoal
mespessoal = reduzirmes(mespessoal)
#totalmeses = mesespessoal()


# Tabela de Ano Pessoal
tabelano = []
tabelano = tabelaano(dia, mes, ano)

# CICLOS DE VIDA
# 1º ciclo - tem o número reduzido do mês do nascimento, salvo o mês 11 que não se reduz.

if int(mes) == 11:
    numciclo1 = 11
else:
    numciclo1 = reduzir(int(mes))
idadeciclo1 = (37 - destinonum)
periodociclo1 = (37 - destinonum) + int(ano)

# 2º ciclo - tem o número reduzido do dia do nascimento, salvo 11 e 22.

if int(dia) == 11:
    numciclo2 = 11
elif int(dia) == 22:
    numciclo2 = 22
else:
    numciclo2 = reduzir(int(dia))
idadeciclo2 = 27

# 3º ciclo - tem o número reduzido do ano do nascimento salvo os números 11 e 22;

if int(ano) == 11:
    numciclo3 = 11
elif int(ano) == 22:
    numciclo3 = 22
else:
    numciclo3 = reduzir(int(ano))

# RELAÇÃO DE CICLO DE VIDA COM AS LIÇÕES CARMICAS
listanumciclo = [numciclo1, numciclo2, numciclo3]
relacaovidacarma = list()
relacaovidacarma = compararlistas(listanumciclo, licaocarm)

# DESAFIOS
# 1º desafio - reduzir o dia e o mes do nascimento então subtrair o menor do maior.
diad = reduzirtotal(int(dia))
mesd = reduzirtotal(int(mes))
if mesd < diad:
    desafio1 = diad - mesd
else:
    desafio1 = mesd - diad
desafio1 = reduzirtotal(desafio1)
# 2º desafio - reduzir o dia e o ano de nascimento então subtrair o menor do maior.
anod = reduzirtotal(int(ano))
if anod < diad:
    desafio2 = diad - anod
else:
    desafio2 = anod - diad
desafio2 = reduzirtotal(desafio2)
# 3º desafio ou desafio principal - deve-se subtrair o menor do maior entre os desafios.
if desafio1 < desafio2:
    desafioprinc = desafio2 - desafio1
else:
    desafioprinc = desafio1 - desafio2
desafioprinc = reduzirtotal(desafioprinc)

# MOMENTOS DECISIVOS
# 1º Momento Decisivo -  é a soma reduzida do dia e do mês do nascimento. Sua duração é igual à do 1º ciclo de vida
# não se reduzem os números 11 e 22.

md1 = reduzir(int(dia)) + reduzir(int(mes))
md1 = reduzir(md1)
# 2º Momento Decisivo - é a soma reduzida do dia e do ano do nascimento.
md2 = reduzir(int(dia)) + reduzir(int(ano))
md2 = reduzir(md2)
# 3º Momento Decisivo - é a soma reduzida do 1º e do 2º momento decisivo.
md3 = reduzir(md1) + reduzir(md2)
md3 = reduzir(md3)
# 4º Momento Decisivo - é a soma reduzida do 1º e do 2º momento decisivo.
md4 = reduzir(int(mes)) + reduzir(int(ano))
md4 = reduzir(md4)
mdmotivac = 0
mdexpressao = 0
mdestino = 0
mdlista = list()
mdlista.append(md1)
mdlista.append(md2)
mdlista.append(md3)
mdlista.append(md4)

# verificação de igualdade sobre fatos do MD.
if md1 == motivac or md2 == motivac or md3 == motivac or md4 == motivac:
    mdmotivac = 1
if md1 == expressao or md2 == expressao or md3 == expressao or md4 == expressao:
    mdexpressao = 1
if md1 == destinonum or md2 == destinonum or md3 == destinonum or md4 == destinonum:
    mdestino = 1

# DIA DO MÊS FAVORÁVEL
diaux = (numerosfavoraveis[int(mes) - 1][int(dia)][1])
diafavoravel = list()
diafavoravel.append(diaux[0])
diafavoravel.append(diaux[1])
auxnum = 0
auxnum = diaux[1] * 2
diafavoravel.append(auxnum)
operador = 0
while auxnum < 30:
    auxnum = auxnum + diaux[0]
    diafavoravel.append(auxnum)
    operador = 1
    if operador == True:
        auxnum = auxnum + diaux[1]
        diafavoravel.append(auxnum)
        operador = 0
examinador = 31
examinador2 = 0
while examinador < 41:
    examinador += 1
    if (examinador in diafavoravel):
        examinador2 = examinador
if examinador2 in diafavoravel:
    diafavoravel.remove(examinador2)

# GRAUe DE ASCENSÃO
vogais = reduzirtotal(calcular(vogaiscompleto))
consoantes = reduzirtotal(calcular(conscompleto))
if vogais == consoantes:
    ascensao = "Estamos diante de um espírito elevado que veio a este planeta (neste momento) para iluminar outras almas."
elif vogais > consoantes:
    ascensao = 'Estamos diante de um espírito que alcançou um elevado grau de honrarias e,' \
               ' transgredindo as Leis Naturais, foi "rebaixado", voltando agora em um meio inferior' \
               ' ao que viveu antes.'
else:
    ascensao = 'Estamos diante de um espírito em ascensão.'

# NÚMEROS DO AMOR E RELACIONAMENTOS -  Soma do número de expressão mais Destino.
numamor = expressao + destinonum
numamor = reduzirtotal(numamor)

# ANO UNIVERSAL, MES UNIVERSAL E DIA UNIVERSAL
anouniversal = reduzirtotal(ano)
mesuniversal = reduzirtotal(mes) + anouniversal
mesuniversal = reduzirtotal(mesuniversal)
diauniversal = reduzirtotal(dia) + mesuniversal
diauniversal = reduzirtotal(diauniversal)

# ANO PESSOAL, MES PESSOAL E DIA PESSOAL
# POSSIVELMENTE AQUI ENTRA A IDEIA DE FAZER UMA TABELA PARA A PESSOA ACOMPANHAR OU O DIÁRIO DE VARIOS DIAS E ANOS.
# COMO ANOS PASSADOS, E ANOS FUTUROS E COMO FOI... FAZER UMA HISTÓRIA.


# NÚMERO QUE SE HARMONIZAM DE ACORDO COM O DIA DO NASCIMENTO
harmoniadia = reduzirtotal(dia)
incompativelmes = reduzirtotal(mes)
# falta a saber como faz com o mes incompativel


# Harmonia numérica. entre duas pessoas


# SAÚDE PELO NÚMERO DE DESAFIO
listadesafio = [desafio1, desafio2, desafioprinc]
listacomparacao = [numciclo1, numciclo2, numciclo3, md1, md2, md3, md4, destinonum]
doencasi = compararlistas(listadesafio, listacomparacao)

#Piramide numérica
mtx = matrixnum(nomecompleto)
result = piramidnum(mtx)


#########################################################################################################################################
# print de tela

CRED = '\033[91m'
CEND = '\033[0m'
div = '-' * 60
div2 = '-' * 35
print(div)
print(f'{"NUMEROLOGIA CABALÍSTICA":^40}')
print(div)
print(CRED + '\n {}\n {}\n'.format(nome, datadenascimento) + CEND)
print('A numerologia do seu primeiro nome {} é {} '.format(primeironome, numerolpnome))

print('\nO Seu número de Motivação é', motivac)
print(Motivação[0])
print('A motivação {}. {}'.format(motivac, Motivação[motivac]))

print('\n\nO Número de Impressão ou Aparência é', aparencia)
print(Impressão[0])
print('A impressão {}. {}'.format(aparencia, Impressão[aparencia]))

print('\n\nO Número de Expressão é', format(expressao))
print(Expressão[0])
print('A expressão {}. {}'.format(expressao, Expressão[expressao]))

print('\n\nO Número de Destino é', destinonum)
print(Destino[0])
print('O destino {}. {}'.format(destinonum, Destino[destinonum]))

if licaocarm:
    print('\n\nAs suas lições cármicas são:', licaocarm)
    print(Liçõescarmicas[0])
    for lic in licaocarm:
        print(Liçõescarmicas[lic])
else:
    print('Você não possui lições cármicas!')

print('\n\na tendência oculta são:', tendocult)
print(Tendênciasocultas[0])
for to in tendocult:
    print(Tendênciasocultas[to])

print('\n\nA resposta do subconsciente: ', respsub)
print(RespostaSubconsciente[0])
print(RespostaSubconsciente[respsub])

if divida:
    print('\n\nAs Dívidas Cármicas são:', divida)
    print(DívidasCármicas[0])
    for d in divida:
        print(DívidasCármicas[d])
else:
    print('\n\nVocê não possui Dívidas Cármicas!')

print('\n\nO Número da Missão é: ', missao)
print(Missão[0])
print(Missão[missao])

print('\n\nO seu dia de Nascimento é: ', dia)
print(
    'O dia do nascimento é um número único, onde cada dia do mês possui suas próprias características, diferenciando-se em cada uma delas.')
print(DiadoNascimento[int(dia)])

print('\n\nA vibração de cada ano influencia a pessoa,\n'
      'Observe que se o dia do aniversário ainda não chegou\n'
      'prevalece o ano anterior como número base\n'
      'o Seu ano pessoal é o ano: {} no ano atual de {}. E o seu mês pessoal {}.\n'.format(anopessoal, anoatual,
                                                                                           mespessoal))
print(AnoPessoal[anopessoal])
print(MêsPessoal[mespessoal])
print('\n')

print(f'{"O SEU CICLO DE VIDA":^30}')
print(
    '1º Ciclo: Numerologia {} - Período: {} - {}, até aos {} anos.'.format(numciclo1, ano, periodociclo1, idadeciclo1))
print('2º Ciclo: Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(numciclo2, periodociclo1,
                                                                                    periodociclo1 + 27, idadeciclo1,
                                                                                    idadeciclo1 + idadeciclo2))
print(
    '3º Ciclo: Numerologia {} - Período: {} - ****, dos {} até ao final da vida.'.format(numciclo3, periodociclo1 + 27,

                                                                                         idadeciclo1 + idadeciclo2))
if relacaovidacarma:
    print(
        '\nQuando um ciclo de vida tiver o mesmo número que uma das lições cármicas,'
        '\nesse período poderá ser um tanto conturbado... '
        '\nAté que a lição cármica seja aprendida e eliminada:\n'
        ' Atenção ao(s) ciclo(s) cuja numerologia seja:', relacaovidacarma)
print('\n')
print(f'{"DESCRIÇÃO DO SEU CICLO DE VIDA":^15}')
print(
    '1º Ciclo: Numerologia {} - Período: {} - {}, até aos {} anos.'.format(numciclo1, ano, periodociclo1, idadeciclo1))
print(PrimeiroCiclo[numciclo1])
print('2º Ciclo: Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(numciclo2, periodociclo1,
                                                                                    periodociclo1 + 27, idadeciclo1,
                                                                                    idadeciclo1 + idadeciclo2))
print(SegundoCiclo[numciclo2])
print(
    '3º Ciclo: Numerologia {} - Período: {} - ****, dos {} até ao final da vida.'.format(numciclo3, periodociclo1 + 27,
                                                                                         idadeciclo1 + idadeciclo2))
print(TerceiroCiclo[numciclo3])

print(f'\n{"DESAFIOS":^30}')
print(
    'Paralelos aos Ciclos de Vida, existem os Desafios, que nada mais são que certas fraquezas, obstáculos que se entrepõem no caminho de cada um de nós. Os Desafios representam etapas em que a pessoa tem certa dificuldade para encontrar o seu rumo e que tende a agir de forma extremada e precipitada.')
print('\n1º Desafio - Numerologia {}'.format(desafio1))
print(Desafios[desafio1])
print('2º Desafio - Numerologia {}'.format(desafio2))
print(Desafios[desafio2])
print('DESAFIO PRINCIPAL - Numerologia {}'.format(desafioprinc))
print(Desafios[desafioprinc])
print('')

for item in saude:
    for item1 in doencasi:
        if item == item1:
            print(
                '\nPoderá haver certos problemas de saúde durante esse período do desafio {} com uma das seguintes partes do corpo: {}'.format(
                    item, saude[item]))

print(f'\n{"MOMENTOS DECISIVOS":^30}')
print('1º Momento Decisivo - Numerologia {} - Período: {} - {}, até aos {} anos.'.format(md1, ano, periodociclo1,
                                                                                         idadeciclo1))
print('2º Momento Decisivo - Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(md2, periodociclo1,
                                                                                                periodociclo1 + 9,
                                                                                                idadeciclo1,
                                                                                                idadeciclo1 + 9))
print('3º Momento Decisivo - Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(md3, periodociclo1 + 9,
                                                                                                periodociclo1 + 18,
                                                                                                idadeciclo1 + 9,
                                                                                                idadeciclo1 + 18))
print('4º Momento Decisivo - Numerologia {} - Período: {} - ****, dos {} até ao final da vida.'.format(md4,
                                                                                                       periodociclo1 + 18,
                                                                                                       idadeciclo1 + 18))
if mdmotivac == True:
    print('Durante os anos do MD {} ,\na pessoa terá os seus desejos e necessidades atendidas.'.format(motivac))
if mdexpressao == True:
    print('Durante os anos do MD {} ,\na pessoa será ajudada na sua vocação (profissão ideal).'.format(expressao))
if mdestino == True:
    print('Durante os anos do MD {} em vigor,\n terá uma grande oportunidade de realização dos sonhos (destino)'.format(
        destinonum))
if len(licaocarm) > 0:
    for q in mdlista:
        for j in licaocarm:
            if q == j:
                print(
                    'O período do {}º Momento Decisivo conterá muitas dificuldades no que relaciona ao aprendizado da lição cármica {}, enquanto esta não for aprendida.'.format(
                        (mdlista.index(q) + 1), q))
print('\nDESCRIÇÃO DOS SEUS MOMENTOS DECISIVOS')
print(MomentoDecisivo[md1])
print(MomentoDecisivo[md2])
print(MomentoDecisivo[md3])
print(MomentoDecisivo[md4])

print('\nDias do mês favoráveis para fechar negócios, encontros, decisões e bons resultados:', diafavoravel)

if relinterval:
    print('\nRelações intervalores', relinterval)
    print(RelaçõesInterValores[0])
    for riv in relinterval:
        print(RelaçõesInterValores[riv])

print('\nGrau de Ascensão:', ascensao)
print('\nO número do Amor e Relacionamento é:', numamor)
print('\n', Amor[numamor])
print(('\nNUMEROLOGIA DO AMOR: \n'))
print('Se o número estiver repetido, é por que essa ligação é extremamente forte.\n')
print('Vibra com {}. É sinal de forte atração sexual (paixão) que pode,\n'
      'caso não se transforme em amor, levar à separação em virtude de ciúmes\n'
      'exagerados, inconstância sexual, arrogância de um ou ambos os parceiros.\n'.format(
    harmoniaconjugal[numamor - 1][1]))

print('Atrai {}. São totalmente compatíveis entre si e o amor e sexo se mesclam\n'
      'ardentemente. Tendem para a amabilidade, cordialidade, delicadeza e compreensão\n'
      'mútuas. O relacionamento com esta característica tem tudo para ser bem-sucedido,\n'
      'e se não for o for, ou será por interferência de terceiros ou alteração sem análise do nome\n'
      '(assinatura) de um dos parceiros ou mesmo de ambos.\n'.format(harmoniaconjugal[numamor - 1][2]))

print('É oposto ao {}. Por incrível que possa parecer, este não é o pior dos relacionamentos,\n'
      'quando um dos envolvidos tem consciência do fato e são intelectualmente desenvolvidos.\n'
      'Para que dure "eternamente", ambos os parceiros devem usar diplomacia constantemente e ceder\n'
      '(os dois) em muitos pontos, a fim de impedir o fracasso dessa união. Jamais, em hipótese alguma,\n'
      'um deve querer adaptar o outro aos seus ideias ou desejos. Neste caso a separação será\n'
      'inevitável e traumática para ambos.\n'.format(harmoniaconjugal[numamor - 1][3]))

print('É passivo ao {}. Sofrem influências, sejam elas negativas ou positivas.\n'
      'Este tipo de relacionamento transforma os parceiros em "amigos", em vez de \n'
      '"amantes". Por vezes, em virtude de uma forte vibração de um dos parceiros,\n'
      'o relacionamento tende a dar certo por longo tempo, ou seja, quando um ama\n'
      'demasiadamente e ou outro se mostra um tanto indiferente. O mais normal é que\n'
      'nesta configuração o relacionamento tem tendência a se " arrastar por toda a vida",\n'
      'ou então até que apareça a um dos envolvidos "alguém" mais interessante.\n'.format(
    harmoniaconjugal[numamor - 1][4]))

print(f'\n{"TABELA I - NÍVEL UNIVERSAL":^30}')
print(div2)
print('Ano Universal =', anouniversal)
print('Mês Universal =', mesuniversal)
print('Dia Universal =', diauniversal)
print('\nNúmeros que se harmonizam de acordo com o dia do nascimento:', harmoniadatanasc[harmoniadia - 1][1])
print('Números que são neutros de acordo com o dia do nascimento:', neutronasc[harmoniadia - 1][1])
print('Números que são incompatíveis de acordo com o dia do nascimento:', incompativelnasc[harmoniadia - 1][1])

print(
    '\nO branco simboliza a luz, a alegria, o poder de Deus e a pureza; o amarelo significa o amor divino, a compreensão e a caridade; azul é a cor da eternidade, da sabedoria e da vontade; o ver de significa a vida, o crescimento e a esperança; o violeta a dignidade e o saber; o vermelho é a cor das emoções; o preto simboliza a tristeza e a escuridão. A mudança do número de Expressão, não altera a cores favoráveis originais.')
print('Cores Favoráveis, de acordo com o número de expressão: \n', cores[expressao])


print(div2*3)
print('==============================================================================================')
print('{}'.format(namelist.replace("", "  ")[1: -1]))
print('{}'.format(mtx))
for i in range(len(result)):
    c = ' '
    print('{}{}'.format(c*(int(i)),result[i]))

print(div)
print(f'{"NUMEROLOGIA CABALÍSTICA":^40}')
print(div)

#######################################################################################################################
##############################################################################################################################
# Print de texto


document = Document()

document.add_heading('Numerologia Cabalística', 0)

p = document.add_paragraph('Bem vindo ao seu melhor e mais completo mapa numerológico ')
document.add_page_break()


p = document.add_paragraph(' ')
p.add_run('{}'.format(namelist.replace('','  ')[1: -1])).bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph('{}'.format(mtx))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i in range(len(result)):
    c = ' '
    b = document.add_paragraph('{}{}'.format(c*(int(i)),result[i]))
    b.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_page_break()


# document.add_heading('{} {}.'.format(nome,datanasc), level=1)
document.add_paragraph('{}, {}.'.format(nome, datadenascimento), style='Intense Quote')

document.add_heading('Motivação')
document.add_paragraph('\n{}'.format(Motivação[0]))
document.add_paragraph('O seu número de Motivação é {}.'.format(motivac))
document.add_paragraph('{}'.format(Motivação[motivac]), )
document.add_page_break()

document.add_heading('Impressão ou Aparência')

document.add_paragraph('\n{}'.format(Impressão[0]))
document.add_paragraph('O seu número de Impressão é {}.'.format(aparencia))
document.add_paragraph('{}'.format(Impressão[aparencia]))
document.add_page_break()

document.add_heading('Expressão')
document.add_paragraph('\n{}'.format(Expressão[0]))
document.add_paragraph('O seu número de Expressão é {}.'.format(expressao))
document.add_paragraph('{}'.format(Expressão[expressao]))
document.add_page_break()

document.add_heading('Número de Destino')
document.add_paragraph('\n{}'.format(Destino[0]))
document.add_paragraph('O seu número de Destino é {}.'.format(destinonum))
document.add_paragraph('{}'.format(Destino[destinonum]))
document.add_page_break()

document.add_heading('Lição Cármica')
if licaocarm:
    document.add_paragraph('\n{}'.format(Liçõescarmicas[0]))
    document.add_paragraph('As suas lições cármicas são: {}'.format(licaocarm))
    for lic in licaocarm:
        document.add_paragraph('{}'.format(Liçõescarmicas[lic]))
    document.add_page_break()
else:
    document.add_paragraph('Você não possui lições cármicas!')

document.add_heading('Tendência Oculta')
if tendocult:
    document.add_paragraph('\n{}'.format(Tendênciasocultas[0]))
    document.add_paragraph('\nA tendência oculta são:{}'.format(tendocult))
    for to in tendocult:
        document.add_paragraph(Tendênciasocultas[to])
    document.add_page_break()

else:
    document.add_paragraph('Você não possui Tendências Ocultas!')
    document.add_page_break()

document.add_heading('A resposta do Subconsciente')
document.add_paragraph('\n{}'.format(RespostaSubconsciente[0]))
document.add_paragraph('\nSua resposta do subconsciente é {}.'.format(respsub))
document.add_paragraph(RespostaSubconsciente[respsub])
document.add_page_break()

document.add_heading('Dívidas Cárimas')
if divida:
    document.add_paragraph('\nAs Dívidas Cármicas são:{}'.format(divida))
    document.add_paragraph(DívidasCármicas[0])
    for d in divida:
        document.add_paragraph(DívidasCármicas[d])
    document.add_page_break()
else:
    document.add_paragraph('\nVocê não possui Dívidas Cármicas!')
    document.add_page_break()

document.add_heading('Missão')
document.add_paragraph('\n{}'.format(Missão[0]))
document.add_paragraph('O número da sua Missão é: {}'.format(missao))
document.add_paragraph('{}'.format(Missão[missao]))
document.add_page_break()

document.add_heading('Análise do Dia de Nascimento')
document.add_paragraph(
    '\nO dia do nascimento é um número único, onde cada dia do mês possui suas próprias características, diferenciando-se em cada uma delas.')
document.add_paragraph('\nO seu dia de Nascimento é:{}'.format(dia))
document.add_paragraph(DiadoNascimento[int(dia)])
document.add_page_break()



document.add_heading('Cilo de Vida')
document.add_paragraph(
    '\n1º Ciclo: Numerologia {} - Período: {} - {}, até aos {} anos.'.format(numciclo1, ano, periodociclo1,
                                                                             idadeciclo1))
document.add_paragraph(
    '2º Ciclo: Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(numciclo2, periodociclo1,
                                                                                  periodociclo1 + 27, idadeciclo1,
                                                                                  idadeciclo1 + idadeciclo2))
document.add_paragraph(
    '3º Ciclo: Numerologia {} - Período: {} - ****, dos {} até ao final da vida.\n\n'.format(numciclo3, periodociclo1 + 27,
                                                                                         idadeciclo1 + idadeciclo2))
if relacaovidacarma:
    document.add_paragraph(
        '\nQuando um ciclo de vida tiver o mesmo número que uma das lições cármicas, esse período poderá ser um tanto conturbado. '
        '\nAté que a lição cármica seja aprendida e eliminada:\n'
        ' Atenção ao(s) ciclo(s) cuja numerologia seja:{}'.format(relacaovidacarma))

document.add_paragraph('DESCRIÇÃO DO SEU CICLO DE VIDA')

document.add_paragraph(
    '1º Ciclo: Numerologia {} - Período: {} - {}, até aos {} anos.'.format(numciclo1, ano, periodociclo1, idadeciclo1))
document.add_paragraph(PrimeiroCiclo[numciclo1])
document.add_paragraph(
    '2º Ciclo: Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(numciclo2, periodociclo1,
                                                                                  periodociclo1 + 27, idadeciclo1,
                                                                                  idadeciclo1 + idadeciclo2))
document.add_paragraph(SegundoCiclo[numciclo2])
document.add_paragraph(
    '3º Ciclo: Numerologia {} - Período: {} - ****, dos {} até ao final da vida.'.format(numciclo3, periodociclo1 + 27,
                                                                                         idadeciclo1 + idadeciclo2))
document.add_paragraph(TerceiroCiclo[numciclo3])
document.add_page_break()

document.add_heading('DESAFIOS')
document.add_paragraph(
    '\nParalelos aos Ciclos de Vida, existem os Desafios, que nada mais são que certas fraquezas, obstáculos que se entrepõem no caminho de cada um de nós. Os Desafios representam etapas em que a pessoa tem certa dificuldade para encontrar o seu rumo e que tende a agir de forma extremada e precipitada.')
document.add_paragraph('\n1º Desafio - Numerologia {}'.format(desafio1))
document.add_paragraph('{}'.format(Desafios[desafio1]))
document.add_paragraph('2º Desafio - Numerologia {}'.format(desafio2))
document.add_paragraph('{}'.format(Desafios[desafio2]))
document.add_paragraph('DESAFIO PRINCIPAL - Numerologia {}'.format(desafioprinc))
document.add_paragraph(Desafios[desafioprinc])

for item in saude:
    for item1 in doencasi:
        if item == item1:
            document.add_paragraph(
                '\nPoderá haver certos problemas de saúde durante esse período do desafio {} com uma das seguintes partes do corpo: {}'.format(
                    item, saude[item]))
document.add_page_break()

document.add_heading('Momentos Decisivos')

document.add_paragraph(
    '\n1º Momento Decisivo - Numerologia {} - Período: {} - {}, até aos {} anos.'.format(md1, ano, periodociclo1,
                                                                                         idadeciclo1))
document.add_paragraph(
    '2º Momento Decisivo - Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(md2, periodociclo1,
                                                                                              periodociclo1 + 9,
                                                                                              idadeciclo1,
                                                                                              idadeciclo1 + 9))
document.add_paragraph(
    '3º Momento Decisivo - Numerologia {} - Período: {} - {}, dos {} até aos {} anos.'.format(md3, periodociclo1 + 9,
                                                                                              periodociclo1 + 18,
                                                                                              idadeciclo1 + 9,
                                                                                              idadeciclo1 + 18))
document.add_paragraph(
    '4º Momento Decisivo - Numerologia {} - Período: {} - ****, dos {} até ao final da vida.\n\n'.format(md4,
                                                                                                         periodociclo1 + 18,
                                                                                                         idadeciclo1 + 18))
if mdmotivac == True:
    document.add_paragraph(
        'Durante os anos do MD {} ,\na pessoa terá os seus desejos e necessidades atendidas.'.format(motivac))
if mdexpressao == True:
    document.add_paragraph(
        'Durante os anos do MD {} ,\na pessoa será ajudada na sua vocação (profissão ideal).'.format(expressao))
if mdestino == True:
    document.add_paragraph(
        'Durante os anos do MD {} em vigor,\n terá uma grande oportunidade de realização dos sonhos (destino)'.format(
            destinonum))
if len(licaocarm) > 0:
    for q in mdlista:
        for j in licaocarm:
            if q == j:
                document.add_paragraph(
                    'O período do {}º Momento Decisivo conterá muitas dificuldades no que relaciona ao aprendizado da lição cármica {}, enquanto esta não for aprendida.'.format(
                        (mdlista.index(q) + 1), q))
document.add_paragraph('\nDESCRIÇÃO DOS SEUS MOMENTOS DECISIVOS')
document.add_paragraph('{}'.format(MomentoDecisivo[md1]))
document.add_paragraph('{}'.format(MomentoDecisivo[md2]))
document.add_paragraph('{}'.format(MomentoDecisivo[md3]))
document.add_paragraph('{}'.format(MomentoDecisivo[md4]))
document.add_page_break()

document.add_heading('Dias favoráveis')
document.add_paragraph(
    '\nDias do mês favoráveis para fechar negócios, encontros, decisões e bons resultados:\n{}'.format(diafavoravel))

if relinterval:
    document.add_heading('Relações Intervalores')
    #    document.add_paragraph('\n{}'.format(relinterval))
    document.add_paragraph('{}'.format(RelaçõesInterValores[0]))
    for riv in relinterval:
        document.add_paragraph('{}'.format(RelaçõesInterValores[riv]))

document.add_heading('Grau de Ascensão')
document.add_paragraph('{}'.format(ascensao))
document.add_page_break()
document.add_heading('Número do Amor e Relacionamento')
document.add_paragraph('\nO número do Amor e Relacionamento é: {}'.format(numamor))
document.add_paragraph('{}'.format(Amor[numamor]))
document.add_paragraph('Se o número estiver repetido, é por que essa ligação é extremamente forte.\n')
document.add_paragraph('Vibra com {}. É sinal de forte atração sexual (paixão) que pode,\n'
                       'caso não se transforme em amor, levar à separação em virtude de ciúmes exagerados, inconstância sexual, arrogância de um ou ambos os parceiros.\n'.format(
    harmoniaconjugal[numamor - 1][1]))

document.add_paragraph('Atrai {}. São totalmente compatíveis entre si e o amor e sexo se mesclam '
                       'ardentemente. Tendem para a amabilidade, cordialidade, delicadeza e compreensão '
                       'mútuas. O relacionamento com esta característica tem tudo para ser bem-sucedido, '
                       'e se não for o for, ou será por interferência de terceiros ou alteração sem análise do nome'
                       '(assinatura) de um dos parceiros ou mesmo de ambos.\n'.format(harmoniaconjugal[numamor - 1][2]))

document.add_paragraph('É oposto ao {}. Por incrível que possa parecer, este não é o pior dos relacionamentos, '
                       'quando um dos envolvidos tem consciência do fato e são intelectualmente desenvolvidos. '
                       'Para que dure "eternamente", ambos os parceiros devem usar diplomacia constantemente e ceder'
                       '(os dois) em muitos pontos, a fim de impedir o fracasso dessa união. Jamais, em hipótese alguma, '
                       'um deve querer adaptar o outro aos seus ideias ou desejos. Neste caso a separação será\n'
                       'inevitável e traumática para ambos.\n'.format(harmoniaconjugal[numamor - 1][3]))

document.add_paragraph('É passivo ao {}. Sofrem influências, sejam elas negativas ou positivas. '
                       'Este tipo de relacionamento transforma os parceiros em "amigos", em vez de '
                       '"amantes". Por vezes, em virtude de uma forte vibração de um dos parceiros, '
                       'o relacionamento tende a dar certo por longo tempo, ou seja, quando um ama '
                       'demasiadamente e ou outro se mostra um tanto indiferente. O mais normal é que '
                       'nesta configuração o relacionamento tem tendência a se " arrastar por toda a vida", '
                       'ou então até que apareça a um dos envolvidos "alguém" mais interessante.\n'.format(
    harmoniaconjugal[numamor - 1][4]))

document.add_heading('TABELA I - NÍVEL UNIVERSAL')
document.add_paragraph(div2)
document.add_paragraph('Ano Universal = {}'.format(anouniversal))
document.add_paragraph('Mês Universal = {}'.format(mesuniversal))
document.add_paragraph('Dia Universal = {}'.format(diauniversal))
document.add_paragraph(
    '\nNúmeros que se harmonizam de acordo com o dia do nascimento: {}'.format(harmoniadatanasc[harmoniadia - 1][1]))
document.add_paragraph(
    'Números que são neutros de acordo com o dia do nascimento:'.format(neutronasc[harmoniadia - 1][1]))
document.add_paragraph(
    'Números que são incompatíveis de acordo com o dia do nascimento:'.format(incompativelnasc[harmoniadia - 1][1]))

document.add_heading('Cores e seus Significados')

document.add_paragraph(
    '\nO branco simboliza a luz, a alegria, o poder de Deus e a pureza; o amarelo significa o amor divino, a compreensão e a caridade; azul é a cor da eternidade, da sabedoria e da vontade; o ver de significa a vida, o crescimento e a esperança; o violeta a dignidade e o saber; o vermelho é a cor das emoções; o preto simboliza a tristeza e a escuridão. A mudança do número de Expressão, não altera a cores favoráveis originais.')
document.add_paragraph('Cores Favoráveis, de acordo com o número de expressão: {}'.format(cores[expressao]))
document.add_page_break()


document.add_heading('Vibração do ano')
document.add_paragraph('\nA vibração de cada ano influencia a pessoa')
document.add_paragraph(
    ' O seu ano pessoal é {} e mês pessoal {}. Sob a influência do ano {}.'.format(anopessoal, mespessoal, anoatual))
document.add_page_break()


temp =list()
resultadodia = list() #anno, mês, dia e numerologia do dia
listapordia = list()
sep = list()
for y in range(1, 10):
    mesespessoal = []
    aux = reduzirtotal(int(y)) #ano
    document.add_heading('Ano Pessoal {}'.format(aux))
    document.add_paragraph('\n{}'.format(AnoPessoal[aux]))
    document.add_page_break()

    for k in range(1, 13):
        aux2 = int(k) + aux
        aux2 = reduzirmes(aux2)
        mesespessoal.append(aux2)
        document.add_heading('{} do Ano Pessoal {}\n'.format(Descmes[k],y),level=2) #mes
        document.add_paragraph('{}, Com a Vibração Numerológica {}\n'.format(Descmes[k],MêsPessoal[aux2]))
        for d in range(1, 32): #verificação do dia
            aux3 = int(d) + aux2
            aux3 = reduzir(aux3)
            temp.append(int(aux)) #Registro do Ano
            temp.append(str(Descmes[k])) #Registro do mês
            temp.append(int(d)) #REgistro do dia do mês
            temp.append(int(aux3)) # Registro da numerologia do dia
            resultadodia.append(temp[:])
            temp.clear()

        aux2 = 0
        sep.append(resultadodia[:])
        resultadodia.clear()
    document.add_page_break()
print('lista dos dias separados {}:'.format(sep))
print('tamaho da lista separados {}'.format(len(sep)))
#
# for year in range(1,10):
#     print(year)
#     for month in range(1,13):
#         print(month)
#         for day in range(0,31):
#             print(f'[{resultadodia[month][day]}]', end='')
#     print()





document.add_heading('Seus Dias Pessoais')
document.add_paragraph('\nUm período de 24 oferece tempo suficiente para você trabalhar com a vibração, para ser usada em sua vantagem.')
for i in DiaPessoal:
    document.add_paragraph(DiaPessoal[i])

document.add_page_break()


document.add_heading('Numerologia Para Cada Ano\n')

cont = 0
for ead, cas in enumerate(tabelano):
    document.add_paragraph('Ano de {}   Numerologia: {}. {} Objetivo: {} Perigos: {}'.format(cas,tabelano[cas], AConseq2objetivo[tabelano[cas]], AConseq3oquedevefazer[tabelano[cas]],AConseq4perigo[tabelano[cas]]))
    cont = cont +1
    if cont %10 == 0:
        document.add_page_break()


document.add_page_break()

document.save('{} {} {}.docx'.format(nome, datanasc, date.today()))








#
#
#
# print('Ano :',sep[0][0][0])
# print('Mês :',sep[0][0][1])
#
# print('Dia :',sep[0][0][2])
# print('Numerol:',sep[0][0][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][1][2])
# print('Numerologia :',sep[0][1][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][2][2])
# print('Numerologia :',sep[0][2][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][3][2])
# print('Numerologia :',sep[0][3][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][4][2])
# print('Numerologia :',sep[0][4][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][5][2])
# print('Numerologia :',sep[0][5][3])
# print('---------------------------------------------------')
# print('Dia :',sep[0][6][2])
# print('Numerologia :',sep[0][6][3])
#
#
# print('-==-=-=-=-=-=-=-=-=-=-=--=-=-=-=-=-=-=-=-=-=-=-=-=-=--=-=-=-=-=-=-=-=-=-=-')
# for i in range(0,32):
#     print('Valor de i : ',i)
#     print('Dia',sep[0][i][2],end='')
#     print('/N:',sep[0][i][3])






















# table = document.add_table(rows=1, cols=7)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id in separadopormes:
#    print('qty {} , id {}'.format(qty,id))


# i, f = 0, 1
# tam = len(separadopormes)
#
# while f < tam:
#     print(f'{separadopormes[i]:.<31}R$ {separadopormes[f]:>6.2f}', end='\n')
#     i += 2
#     f += 2
# print(div)













    # col_names = ('Ano','Num.','Objetivos','Deve ser feiro','Perigos')
# table = document.add_table(rows=1, cols=len(col_names))
# hdr_cells = table.rows[0].cells
# set_repeat_table_header(table.rows[0])
#
# def set_col_widths(table):
#     widths = (Inches(1), Inches(2), Inches(1.5))
#     for row in table.rows:
#         for idx, width in enumerate(widths):
#             row.cells[idx].width = width
#
# set_col_widths(table)
#
# for idx, name in enumerate(col_names):
#     paragraph = hdr_cells[idx].paragraphs[0]
#     run = paragraph.add_run(name)
#     run.bold = True
# cont = 0
# for anoid, id in enumerate(tabelano):

#    row_cells = table.add_row().cells
#    row_cells[0].text = str(id)
#    row_cells[1].text = str(tabelano[id])
#    row_cells[2].text = str(AConseq2objetivo[tabelano[id]])
#    row_cells[3].text = str(AConseq3oquedevefazer[tabelano[id]])
#    row_cells[4].text = str(AConseq4perigo[tabelano[id]])
#    cont += 1
